"""
index_rag.py — Индексатор файлов DOCX / XLSX / XLS / PDF для RAG-системы.

Запуск:
    python index_rag.py                  # использует config.json
    python index_rag.py --recreate       # пересоздать коллекцию с нуля
    python index_rag.py --catalog "D:\\docs"  # другая папка
"""

import argparse
import json
import logging
import os
import sys
import time
import uuid
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

# ── Workaround: Python 3.14 platform._wmi_query зависает на Windows ──────────
# torch.__init__ вызывает platform.uname() → win32_ver() → _wmi_query() которая
# блокируется при загрузке WMI-провайдера. Заменяем правильным фейковым ответом
# (5-кортеж: version, product_type, build_type, spmajor, spminor).
import platform as _platform_mod
if hasattr(_platform_mod, '_wmi_query'):
    def _fast_wmi_query(namespace_class, *args, **kwargs):
        # Возвращаем разумные значения для Win32_OperatingSystem
        return ('10.0.19041', '1', 'Multiprocessor Free', '0', '0')
    _platform_mod._wmi_query = _fast_wmi_query
# ─────────────────────────────────────────────────────────────────────────────

from docx import Document
from openpyxl import load_workbook
from qdrant_client import QdrantClient
from qdrant_client.models import (
    Distance,
    FieldCondition,
    Filter,
    FilterSelector,
    MatchValue,
    PointStruct,
    VectorParams,
)
from sentence_transformers import SentenceTransformer
from tqdm import tqdm

from rag_core import load_config

# ─────────────────────────── logging ───────────────────────────────────
# Явно задаём UTF-8 для FileHandler и StreamHandler, чтобы кириллица
# не превращалась в кракозябры на Windows (cmd/PowerShell cp866/cp1251).
_stream_handler = logging.StreamHandler(sys.stdout)
_stream_handler.setFormatter(logging.Formatter("%(asctime)s - %(levelname)s - %(message)s"))
try:
    _stream_handler.stream.reconfigure(encoding="utf-8")  # Python 3.7+
except AttributeError:
    pass  # не все объекты stream поддерживают reconfigure

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[_stream_handler],
)
logger = logging.getLogger(__name__)

# Поддерживаемые расширения
SUPPORTED_EXTENSIONS = {".docx", ".xlsx", ".xls", ".pdf"}

# ─────────────────────────── этапы индексирования ─────────────────────────────
# Индексирование разбито на этапы: сначала быстрое наполнение индекса, затем
# прогрессивное «догонение» содержимого. Это даёт пользователю работоспособный
# поиск через минуты, а не сутки.
#
#   1. metadata — индексируется ТОЛЬКО имя/путь/размер/mtime всех файлов.
#      58 тыс. файлов за ~5 минут. Поиск по именам работает сразу.
#   2. small    — полное содержимое быстрых файлов: docx/xlsx/xls любого
#      размера + PDF < 2 МБ. Обычно 10-20 минут на 50k файлов.
#   3. large    — содержимое тяжёлых файлов: крупные xlsx/docx + большие PDF
#      + сканированные PDF (медленно, часы). Запускается отдельным проходом
#      когда поиск по метаданным уже работает.
#
# Порядок важен: чем меньше число, тем «старше» этап (больше информации).
STAGES = ("metadata", "small", "large")
STAGE_RANK = {name: i for i, name in enumerate(STAGES)}

# Пороги для категоризации файлов между small/large
SMALL_OFFICE_MB = 20.0  # docx/xlsx/xls < 20 МБ считаем «быстрыми»
SMALL_PDF_MB = 2.0      # PDF < 2 МБ — обычно текстовые, читаются быстро


def _file_category(filepath: Path) -> str:
    """
    Возвращает «small» или «large» для данного файла.
    Используется для разделения файлов между этапами small и large.
    """
    try:
        size_mb = filepath.stat().st_size / 1_048_576
    except OSError:
        return "large"  # не можем прочитать stat — кидаем в «медленный» этап
    ext = filepath.suffix.lower()
    if ext in (".docx", ".xlsx", ".xls") and size_mb < SMALL_OFFICE_MB:
        return "small"
    if ext == ".pdf" and size_mb < SMALL_PDF_MB:
        return "small"
    return "large"


# ═══════════════════════════ RAGIndexer ════════════════════════════════

class RAGIndexer:
    """
    Индексирует файлы DOCX / XLSX / XLS / PDF в векторную базу Qdrant.

    Особенности:
    - Инкрементальное индексирование (пропускает неизменённые файлы).
    - Атомарная запись state.json (через временный файл + os.replace).
    - При --recreate очищает state.json вместе с коллекцией.
    - При изменении файла — сначала удаляет старые векторы из Qdrant.
    - При удалении файла — чистит векторы из Qdrant и из state.json.
    - .xls читается через xlrd (старый формат Excel).
    - PDF: сначала pdfplumber (текстовый слой), затем OCR (если слой пуст).
    """

    def __init__(
        self,
        catalog_path: str,
        qdrant_db_path: str,
        embedding_model: str,
        collection_name: str,
        vector_size: int,
        chunk_size: int,
        chunk_overlap: int,
        batch_size: int,
        recreate_collection: bool = False,
        skip_ocr: bool = False,
        max_chunks_per_file: int = 0,
        read_workers: int = 4,
        use_onnx: bool = False,
        qdrant_url: str = "",
        metadata_only_extensions: Optional[set] = None,
    ) -> None:
        # current_stage выставляется при каждом запуске index_directory(stage=...)
        # и определяет поведение skip-логики и экстракции содержимого.
        self.current_stage: str = "content"  # legacy-совместимый дефолт
        self.catalog_path = Path(catalog_path)
        if not self.catalog_path.exists():
            raise FileNotFoundError(f"Папка каталога не найдена: {catalog_path}")

        self.qdrant_db_path = Path(qdrant_db_path)
        self.collection_name = collection_name
        self.vector_size = vector_size
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.batch_size = batch_size
        self.recreate = recreate_collection
        self.skip_ocr = skip_ocr
        self.max_chunks_per_file = max_chunks_per_file  # 0 = без ограничений
        self.read_workers = read_workers
        # Расширения, для которых индексируется только метадата (без чтения содержимого).
        # Полезно для быстрого первого прохода по скан-PDF: имя/путь/размер — за секунды.
        self.metadata_only_extensions = {
            e.lower() if e.startswith(".") else f".{e.lower()}"
            for e in (metadata_only_extensions or set())
        }

        if use_onnx:
            logger.info("Загрузка модели эмбеддинга: %s (backend=onnx)", embedding_model)
            try:
                self.embedder = SentenceTransformer(embedding_model, backend="onnx")
                logger.info("ONNX backend загружен успешно")
            except Exception as exc:
                logger.warning("ONNX backend недоступен (%s), использую PyTorch", exc)
                self.embedder = SentenceTransformer(embedding_model)
        else:
            logger.info("Загрузка модели эмбеддинга: %s", embedding_model)
            self.embedder = SentenceTransformer(embedding_model)

        # Режим подключения: сервер (Docker) или локальный SQLite
        if qdrant_url:
            logger.info("Подключение к Qdrant серверу: %s", qdrant_url)
            self.qdrant = QdrantClient(url=qdrant_url)
            # state_file рядом с db-папкой (или в qdrant_db_path как запасной путь)
            self.qdrant_db_path.mkdir(parents=True, exist_ok=True)
        else:
            logger.info("Qdrant локальный режим: %s", qdrant_db_path)
            self.qdrant = QdrantClient(path=str(self.qdrant_db_path))
        self._setup_collection()

        self.state_file = self.qdrant_db_path / "index_state.json"
        self.state = self._load_state()

        self._points_buffer: List[PointStruct] = []
        self.point_count = 0

    # ── collection setup ───────────────────────────────────────────────

    def _setup_collection(self) -> None:
        existing = [c.name for c in self.qdrant.get_collections().collections]
        if self.collection_name in existing:
            if self.recreate:
                logger.info("Пересоздание коллекции %s…", self.collection_name)
                self.qdrant.delete_collection(self.collection_name)
                self._create_collection()
                # ВАЖНО: очищаем state.json, чтобы он соответствовал пустой коллекции
                state_file = self.qdrant_db_path / "index_state.json"
                if state_file.exists():
                    state_file.unlink()
                    logger.info("state.json очищен (--recreate)")
            else:
                logger.info("Коллекция %s уже существует.", self.collection_name)
        else:
            self._create_collection()

    def _create_collection(self) -> None:
        logger.info("Создание коллекции %s…", self.collection_name)
        self.qdrant.create_collection(
            collection_name=self.collection_name,
            vectors_config=VectorParams(size=self.vector_size, distance=Distance.COSINE),
        )

    # ── state management ───────────────────────────────────────────────

    def _load_state(self) -> Dict[str, Any]:
        if self.state_file.exists():
            try:
                with open(self.state_file, "r", encoding="utf-8") as fh:
                    return json.load(fh)
            except Exception as exc:
                logger.warning("Не удалось загрузить state.json: %s. Начинаем заново.", exc)
        return {"files": {}}

    def _save_state(self) -> None:
        """Атомарная запись: пишем во временный файл, затем os.replace."""
        tmp = self.state_file.with_suffix(".tmp")
        try:
            with open(tmp, "w", encoding="utf-8") as fh:
                json.dump(self.state, fh, indent=2, ensure_ascii=False)
            os.replace(tmp, self.state_file)
        except Exception as exc:
            logger.error("Не удалось сохранить state.json: %s", exc)
            if tmp.exists():
                tmp.unlink(missing_ok=True)

    # ── fingerprint ────────────────────────────────────────────────────

    def _get_file_fingerprint(self, filepath: Path) -> Tuple[str, float]:
        """Быстрый fingerprint файла: размер + mtime."""
        stat = filepath.stat()
        fingerprint = f"{stat.st_size}_{stat.st_mtime}"
        return fingerprint, stat.st_mtime

    # ── stage helpers ──────────────────────────────────────────────────

    @staticmethod
    def _effective_stage(target_stage: str) -> str:
        """
        Какой stage записать в state для файла, обработанного в рамках `target_stage`.
        На этапе metadata мы НЕ читаем содержимое, поэтому записываем "metadata".
        На этапах small/large мы читаем содержимое → записываем "content".
        """
        return "metadata" if target_stage == "metadata" else "content"

    def _should_skip_for_stage(self, file_key: str, fingerprint: str) -> bool:
        """
        True если файл уже покрыт текущим или более полным этапом.

        Правило: stage "metadata" покрывается любым current_stage != metadata
        только если файл УЖЕ имеет полное содержимое (т.е. state.stage == content).
        Если state.stage == metadata и мы сейчас на small/large — нужно ПРОАПГРЕЙДИТЬ.
        """
        existing = self.state["files"].get(file_key)
        if not existing:
            return False
        if existing.get("fingerprint") != fingerprint:
            return False  # файл изменился — обязательно переиндексируем

        existing_stage = existing.get("stage", "content")  # backward compat
        # Если текущий этап = metadata, а файл уже проиндексирован (любым этапом) —
        # можем пропустить: мета-запись у файла уже есть.
        if self.current_stage == "metadata":
            return True
        # Для small/large пропускаем только если у файла уже есть "content"
        return existing_stage == "content"

    # ── Qdrant vector deletion ─────────────────────────────────────────

    def _delete_file_vectors(self, filepath: Path) -> None:
        """Удалить все векторы в Qdrant, связанные с данным файлом."""
        try:
            self.qdrant.delete(
                collection_name=self.collection_name,
                points_selector=FilterSelector(
                    filter=Filter(
                        must=[
                            FieldCondition(
                                key="full_path",
                                match=MatchValue(value=str(filepath)),
                            )
                        ]
                    )
                ),
            )
            logger.debug("Удалены старые векторы для: %s", filepath)
        except Exception as exc:
            logger.warning("Не удалось удалить векторы для %s: %s", filepath, exc)

    # ── text extraction ────────────────────────────────────────────────

    def _extract_docx(self, filepath: Path) -> str:
        """Извлечь текст из DOCX (параграфы + таблицы)."""
        try:
            doc = Document(filepath)
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            return "\n".join(parts)
        except Exception as exc:
            logger.warning("Ошибка чтения DOCX %s: %s", filepath, exc)
            return ""

    def _extract_xlsx(self, filepath: Path) -> str:
        """
        Извлечь текст из XLSX через openpyxl.
        Вызывается ТОЛЬКО для .xlsx-файлов — маршрутизация происходит
        в process_file/_extract_spreadsheet, НЕ здесь.
        Ранняя остановка: прекращает чтение строк, когда накоплено
        достаточно текста для max_chunks_per_file чанков.
        """
        try:
            wb = load_workbook(filepath, read_only=True, data_only=True)
            parts: List[str] = []
            # Лимит символов: если задан max_chunks_per_file, останавливаемся раньше.
            # Эффективный шаг чанка = chunk_size - chunk_overlap, поэтому
            # для N чанков нужно ~N*(chunk_size-chunk_overlap)+chunk_overlap символов.
            # Берём chunk_size * max_chunks_per_file с запасом.
            max_chars = (
                self.max_chunks_per_file * self.chunk_size
                if self.max_chunks_per_file
                else 0
            )
            total_chars = 0
            done = False
            for sheet_name in wb.sheetnames:
                if done:
                    break
                ws = wb[sheet_name]
                parts.append(f"Лист: {sheet_name}")
                for row in ws.iter_rows(values_only=True):
                    row_text = " | ".join(str(c) if c is not None else "" for c in row)
                    if row_text.strip():
                        parts.append(row_text)
                        total_chars += len(row_text)
                        if max_chars and total_chars >= max_chars:
                            done = True
                            break
            return "\n".join(parts)
        except Exception as exc:
            logger.warning("Ошибка чтения XLSX %s: %s", filepath, exc)
            return ""

    def _extract_spreadsheet(self, filepath: Path) -> str:
        """
        Точка входа для всех табличных форматов.
        Явно маршрутизирует .xls → xlrd, .xlsx → openpyxl.
        Раньше маршрутизация была внутри _extract_xlsx, что приводило
        к неочевидным варнингам при смешении форматов.
        """
        ext = filepath.suffix.lower()
        if ext == ".xls":
            logger.debug("Формат XLS — использую xlrd: %s", filepath.name)
            return self._extract_xls(filepath)
        elif ext == ".xlsx":
            logger.debug("Формат XLSX — использую openpyxl: %s", filepath.name)
            return self._extract_xlsx(filepath)
        else:
            logger.warning("Неизвестное табличное расширение: %s", ext)
            return ""

    def _extract_xls(self, filepath: Path) -> str:
        """Извлечь текст из старого формата XLS через xlrd.
        Ранняя остановка аналогична _extract_xlsx."""
        try:
            import xlrd  # type: ignore
        except ImportError:
            logger.warning("xlrd не установлен. Установите: pip install xlrd")
            return ""
        try:
            wb = xlrd.open_workbook(str(filepath))
            parts: List[str] = []
            max_chars = (
                self.max_chunks_per_file * self.chunk_size
                if self.max_chunks_per_file
                else 0
            )
            total_chars = 0
            done = False
            for sheet in wb.sheets():
                if done:
                    break
                parts.append(f"Лист: {sheet.name}")
                for row_idx in range(sheet.nrows):
                    row = sheet.row_values(row_idx)
                    row_text = " | ".join(
                        str(v) if v not in ("", None) else "" for v in row
                    )
                    if row_text.strip():
                        parts.append(row_text)
                        total_chars += len(row_text)
                        if max_chars and total_chars >= max_chars:
                            done = True
                            break
            return "\n".join(parts)
        except Exception as exc:
            logger.warning("Ошибка чтения XLS %s: %s", filepath, exc)
            return ""

    def _extract_pdf(self, filepath: Path) -> str:
        """
        Извлечь текст из PDF.
        Использует pymupdf (fitz) — в 3-5x быстрее pdfplumber.
        Fallback на pdfplumber если pymupdf не установлен.
        При пустом текстовом слое — OCR (если не --no-ocr).
        """
        # ── Попытка pymupdf (быстрый путь) ────────────────────────────
        try:
            import fitz  # pymupdf
            parts: List[str] = []
            with fitz.open(str(filepath)) as doc:
                for page in doc:
                    text = page.get_text()
                    if text and text.strip():
                        parts.append(text)
            full_text = "\n".join(parts).strip()
            if full_text:
                return full_text
            # Текстового слоя нет
            if self.skip_ocr:
                logger.debug("Нет текстового слоя, OCR пропущен (--no-ocr): %s", filepath.name)
                return ""
            logger.info("Нет текстового слоя в %s — запуск OCR…", filepath.name)
            return self._ocr_pdf(filepath)
        except ImportError:
            logger.debug("pymupdf не установлен, использую pdfplumber")
        except Exception as exc:
            logger.warning("pymupdf: ошибка чтения %s: %s", filepath.name, exc)
            return ""

        # ── Fallback: pdfplumber ───────────────────────────────────────
        try:
            import pdfplumber  # type: ignore
        except ImportError:
            logger.warning("Ни pymupdf, ни pdfplumber не установлены. pip install pymupdf")
            return ""
        try:
            parts = []
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text and text.strip():
                        parts.append(text)
            full_text = "\n".join(parts).strip()
            if full_text:
                return full_text
            if self.skip_ocr:
                logger.debug("Нет текстового слоя, OCR пропущен (--no-ocr): %s", filepath.name)
                return ""
            logger.info("Нет текстового слоя в %s — запуск OCR…", filepath.name)
            return self._ocr_pdf(filepath)
        except Exception as exc:
            logger.warning("pdfplumber: ошибка чтения %s: %s", filepath.name, exc)
            return ""

    def _ocr_pdf(self, filepath: Path) -> str:
        """OCR сканированного PDF через pytesseract + pdf2image."""
        try:
            import pytesseract  # type: ignore
            from pdf2image import convert_from_path  # type: ignore
        except ImportError:
            logger.warning(
                "pytesseract/pdf2image не установлены. "
                "Установите: pip install pytesseract pdf2image"
            )
            return ""

        try:
            pages = convert_from_path(str(filepath), dpi=200)
            parts: List[str] = []
            for i, page_img in enumerate(pages):
                text = pytesseract.image_to_string(page_img, lang="rus+eng")
                if text.strip():
                    parts.append(text)
                logger.debug("OCR страница %d/%d — %s", i + 1, len(pages), filepath.name)
            return "\n".join(parts)
        except Exception as exc:
            logger.warning("OCR не удался для %s: %s", filepath, exc)
            return ""

    # ── chunking ───────────────────────────────────────────────────────

    def _chunk_text(self, text: str) -> List[str]:
        """Разбить текст на перекрывающиеся чанки."""
        if not text:
            return []
        chunks: List[str] = []
        start = 0
        while start < len(text):
            end = start + self.chunk_size
            chunks.append(text[start:end])
            if end >= len(text):
                break
            start = end - self.chunk_overlap
        return chunks

    # ── indexing helpers ───────────────────────────────────────────────

    def _add_points(self, points: List[PointStruct]) -> None:
        self._points_buffer.extend(points)
        self.point_count += len(points)
        if len(self._points_buffer) >= self.batch_size:
            self._flush_buffer()

    def _flush_buffer(self) -> None:
        if self._points_buffer:
            self.qdrant.upsert(self.collection_name, points=self._points_buffer)
            logger.info(
                "Загружен батч: %d точек (итого %d)", len(self._points_buffer), self.point_count
            )
            self._points_buffer = []

    def _index_metadata(self, filepath: Path, relative_path: Path) -> None:
        stat = filepath.stat()
        text = (
            f"Файл: {filepath.name} | Путь: {relative_path} | Расширение: {filepath.suffix}"
        )
        vector = self.embedder.encode(text, normalize_embeddings=True).tolist()
        point = PointStruct(
            id=str(uuid.uuid4()),
            vector=vector,
            payload={
                "type": "file_metadata",
                "text": text,
                "filename": filepath.name,
                "extension": filepath.suffix.lower(),
                "size_mb": round(stat.st_size / (1024 * 1024), 2),
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "path": str(relative_path),
                "full_path": str(filepath),
            },
        )
        self._add_points([point])

    def _index_content(
        self,
        filepath: Path,
        relative_path: Path,
        file_type: str,
        full_text: str,
    ) -> None:
        if not full_text.strip():
            return
        chunks = self._chunk_text(full_text)
        # Ограничение: не более max_chunks_per_file чанков с одного файла
        if self.max_chunks_per_file and len(chunks) > self.max_chunks_per_file:
            logger.warning(
                "Файл %s: %d чанков, обрезано до %d (--max-chunks-per-file)",
                filepath.name, len(chunks), self.max_chunks_per_file,
            )
            chunks = chunks[: self.max_chunks_per_file]
        # Batch-кодирование: encode сразу все чанки файла — 5-10x быстрее чем поштучно
        vectors = self.embedder.encode(
            chunks, normalize_embeddings=True, batch_size=64, show_progress_bar=False
        )
        ext = filepath.suffix.lower()
        points = [
            PointStruct(
                id=str(uuid.uuid4()),
                vector=v.tolist(),
                payload={
                    "type": f"{file_type}_content",
                    "text": chunk,
                    "filename": filepath.name,
                    "extension": ext,
                    "path": str(relative_path),
                    "full_path": str(filepath),
                    "chunk_index": idx,
                },
            )
            for idx, (v, chunk) in enumerate(zip(vectors, chunks))
        ]
        self._add_points(points)

    # ── single file ────────────────────────────────────────────────────

    def process_file(self, filepath: Path) -> None:
        """Обработать один файл: извлечь текст, нарезать, проиндексировать."""
        relative_path = filepath.relative_to(self.catalog_path)
        fingerprint, mtime = self._get_file_fingerprint(filepath)
        file_key = str(filepath)

        if file_key in self.state["files"]:
            if self.state["files"][file_key]["fingerprint"] == fingerprint:
                logger.debug("Файл не изменился, пропуск: %s", filepath)
                return
            logger.info("Файл изменился, удаляю старые векторы: %s", filepath)
            self._delete_file_vectors(filepath)

        logger.info("Индексирование: %s", filepath)
        self._index_metadata(filepath, relative_path)

        ext = filepath.suffix.lower()
        full_text = ""
        file_type = ""

        if ext == ".docx":
            full_text = self._extract_docx(filepath)
            file_type = "docx"
        elif ext in (".xlsx", ".xls"):
            full_text = self._extract_spreadsheet(filepath)
            file_type = "xlsx"
        elif ext == ".pdf":
            full_text = self._extract_pdf(filepath)
            file_type = "pdf"
        else:
            logger.debug("Неподдерживаемый формат (только метаданные): %s", ext)

        if full_text:
            self._index_content(filepath, relative_path, file_type, full_text)

        self.state["files"][file_key] = {"fingerprint": fingerprint, "mtime": mtime}
        self._save_state()

    # ── directory scan ─────────────────────────────────────────────────

    def index_directory(self, stage: str = "content") -> None:
        """
        Pipeline-индексирование на указанном этапе.

          stage="metadata" — только имя/путь/размер/mtime (не читает файлы);
          stage="small"    — полное содержимое docx/xlsx/xls и мелких PDF;
          stage="large"    — полное содержимое крупных и сканированных PDF;
          stage="content"  — legacy: полное содержимое для всех файлов за один проход.

        Pipeline:
          - ThreadPoolExecutor читает файлы параллельно (I/O-bound)
          - Главный поток batch-кодирует накопленные чанки и пишет в Qdrant

        Прирост производительности:
          - Batch encode: 5-10x быстрее чем поштучно
          - Pipeline: чтение следующих файлов идёт пока GPU/CPU кодирует предыдущие
        """
        if stage not in (*STAGES, "content"):
            raise ValueError(f"Неизвестный stage: {stage!r}. Допустимо: {STAGES} или 'content'")
        self.current_stage = stage
        effective_stage = self._effective_stage(stage)

        ENCODE_BATCH = 256  # сколько чанков накапливать перед одним вызовом encode()
                            # 256 оптимально для CPU (OpenBLAS матричные операции)

        # Семафор ограничивает число одновременно «зависших» daemon-потоков.
        # При массовых SMB-таймаутах без ограничения они накапливаются до OOM.
        # Лимит = 2 * read_workers: за один проход воркеров может зависнуть
        # не более read_workers файлов, запас ×2 на перекрытие таймаутов.
        _reader_sem = __import__("threading").Semaphore(self.read_workers * 2)

        logger.info(
            "════════ Этап '%s' (pipeline, workers=%d): %s ════════",
            stage, self.read_workers, self.catalog_path,
        )

        all_files = [
            f
            for f in self.catalog_path.rglob("*")
            if f.is_file()
            and f.suffix.lower() in SUPPORTED_EXTENSIONS
            and not f.name.startswith("~$")  # пропускать временные файлы Office
        ]
        logger.info("Найдено файлов на диске: %d (DOCX/XLSX/XLS/PDF)", len(all_files))

        # Партиция по этапам: metadata берёт всё, small/large — свою категорию.
        if stage == "small":
            scope_files = [f for f in all_files if _file_category(f) == "small"]
            logger.info("Отфильтровано для этапа 'small': %d файлов (docx/xlsx/xls + PDF < %g МБ)",
                        len(scope_files), SMALL_PDF_MB)
        elif stage == "large":
            scope_files = [f for f in all_files if _file_category(f) == "large"]
            logger.info("Отфильтровано для этапа 'large': %d файлов (крупные Office + большие/сканированные PDF)",
                        len(scope_files))
        else:
            # metadata или legacy "content" — работаем со всем
            scope_files = all_files

        # ── буферы для batch-encode ──────────────────────────────────
        pending_texts: List[str] = []
        pending_payloads: List[Dict[str, Any]] = []
        # (file_key, fingerprint, mtime) — для обновления state после записи
        pending_states: List[Tuple[str, str, float]] = []

        def flush() -> None:
            """
            Batch-encode накопленных текстов и запись в Qdrant.
            Разбивает большой список на куски по ENCODE_BATCH,
            чтобы один вызов encode() не блокировал главный поток надолго.
            """
            if not pending_texts:
                return
            # Нарезаем на мини-батчи — encode каждого занимает < ~1 сек
            for i in range(0, len(pending_texts), ENCODE_BATCH):
                chunk_texts    = pending_texts[i : i + ENCODE_BATCH]
                chunk_payloads = pending_payloads[i : i + ENCODE_BATCH]
                vectors = self.embedder.encode(
                    chunk_texts, normalize_embeddings=True,
                    batch_size=256, show_progress_bar=False,
                )
                points = [
                    PointStruct(id=str(uuid.uuid4()), vector=v.tolist(), payload=p)
                    for v, p in zip(vectors, chunk_payloads)
                ]
                self.qdrant.upsert(self.collection_name, points=points)
                self.point_count += len(points)
            for file_key, fingerprint, mtime in pending_states:
                self.state["files"][file_key] = {
                    "fingerprint": fingerprint,
                    "mtime": mtime,
                    "stage": effective_stage,
                }
            self._save_state()
            logger.info(
                "Записан батч: %d точек (итого %d)", len(pending_texts), self.point_count
            )
            pending_texts.clear()
            pending_payloads.clear()
            pending_states.clear()

        # ── I/O-worker: читает один файл, возвращает тексты+payload ─
        def extract_one(filepath: Path):
            """
            Выполняется в потоке-воркере.
            Не кодирует векторы (encode — в главном потоке).
            Возвращает None если файл не изменился.
            """
            relative_path = filepath.relative_to(self.catalog_path)
            fingerprint, mtime = self._get_file_fingerprint(filepath)
            file_key = str(filepath)

            # Stage-aware skip:
            #  - на этапе metadata пропускаем любой уже проиндексированный файл;
            #  - на этапах small/large пропускаем только те, что уже дошли до "content".
            if self._should_skip_for_stage(file_key, fingerprint):
                return None

            ext = filepath.suffix.lower()
            size_mb = round(filepath.stat().st_size / 1_048_576, 1)

            # Таймаут на извлечение: если воркер завис на сетевом I/O (SMB stall) —
            # пропускаем файл через 5 минут и продолжаем индексирование.
            # Используем daemon-поток: он умрёт сам когда процесс завершится,
            # даже если заблокирован в ядре на SMB read().
            FILE_TIMEOUT = 45  # секунд (быстро бросаем зависшие SMB-файлы)

            t_start = time.monotonic()
            full_text = ""
            file_type = ""

            _buf: list = [None, None]  # [result_text, exception]

            # Режим «только метадата»: либо текущий этап = metadata,
            # либо расширение явно в metadata_only_extensions (legacy флаг).
            # Содержимое не читается — файл попадает в индекс по имени/пути/размеру.
            if self.current_stage == "metadata" or ext in self.metadata_only_extensions:
                file_type = ext.lstrip(".") or "file"
                _fn = None
            elif ext == ".docx":
                file_type = "docx"
                _fn = self._extract_docx
            elif ext in (".xlsx", ".xls"):
                file_type = "xlsx"
                _fn = self._extract_spreadsheet
            elif ext == ".pdf":
                file_type = "pdf"
                _fn = self._extract_pdf
            else:
                _fn = None

            if _fn is not None:
                # Логируем тяжёлые файлы заранее — только для тех, что реально читаем
                if size_mb >= 5:
                    logger.info("Читаю крупный файл (%.1f МБ): %s", size_mb, filepath.name)

                import threading as _th

                # Проверяем семафор без блокировки: если лимит зависших потоков
                # исчерпан — пропускаем файл, не создавая новый поток.
                if not _reader_sem.acquire(blocking=False):
                    logger.warning(
                        "Лимит daemon-потоков исчерпан (%d): пропускаю %s",
                        self.read_workers * 2, filepath.name,
                    )
                    full_text = ""
                else:
                    def _reader():
                        try:
                            _buf[0] = _fn(filepath)
                        except Exception as _e:
                            _buf[1] = _e

                    _t = _th.Thread(target=_reader, daemon=True)
                    _t.start()
                    _t.join(timeout=FILE_TIMEOUT)
                    if _t.is_alive():
                        # Поток завис (SMB stall) — освобождаем семафор только
                        # после его завершения через отдельный cleanup-поток.
                        logger.warning(
                            "ТАЙМАУТ SMB (>%dс): пропускаю %s — воркер остался в фоне",
                            FILE_TIMEOUT, filepath.name,
                        )
                        def _cleanup(_t=_t, _sem=_reader_sem):
                            _t.join()        # ждём в фоне сколько потребуется
                            _sem.release()   # освобождаем слот
                        _th.Thread(target=_cleanup, daemon=True).start()
                        full_text = ""
                    else:
                        _reader_sem.release()   # поток завершился штатно
                        if _buf[1] is not None:
                            logger.warning("Ошибка чтения %s: %s", filepath.name, _buf[1])
                            full_text = ""
                        else:
                            full_text = _buf[0] or ""

            elapsed = time.monotonic() - t_start
            if elapsed >= 30:
                logger.warning(
                    "Долгое извлечение (%.0fс, %.1f МБ): %s",
                    elapsed, size_mb, filepath.name,
                )

            chunks = self._chunk_text(full_text) if full_text.strip() else []
            if self.max_chunks_per_file and len(chunks) > self.max_chunks_per_file:
                logger.warning(
                    "Файл %s: %d чанков → обрезано до %d",
                    filepath.name, len(chunks), self.max_chunks_per_file,
                )
                chunks = chunks[: self.max_chunks_per_file]

            stat = filepath.stat()
            meta_text = (
                f"Файл: {filepath.name} | Путь: {relative_path}"
                f" | Расширение: {filepath.suffix}"
            )
            meta_payload: Dict[str, Any] = {
                "type": "file_metadata",
                "text": meta_text,
                "filename": filepath.name,
                "extension": ext,
                "size_mb": round(stat.st_size / (1024 * 1024), 2),
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "path": str(relative_path),
                "full_path": str(filepath),
            }
            content_payloads = [
                {
                    "type": f"{file_type}_content",
                    "text": chunk,
                    "filename": filepath.name,
                    "extension": ext,
                    "path": str(relative_path),
                    "full_path": str(filepath),
                    "chunk_index": idx,
                }
                for idx, chunk in enumerate(chunks)
            ]
            return {
                "filepath": filepath,
                "file_key": file_key,
                "fingerprint": fingerprint,
                "mtime": mtime,
                "was_indexed": file_key in self.state["files"],
                "meta_text": meta_text,
                "meta_payload": meta_payload,
                "chunks": chunks,
                "content_payloads": content_payloads,
            }

        # ── основной pipeline ────────────────────────────────────────
        with ThreadPoolExecutor(max_workers=self.read_workers) as pool:
            futures = {pool.submit(extract_one, f): f for f in scope_files}
            for future in tqdm(as_completed(futures), total=len(scope_files),
                                desc=f"Этап {stage}"):
                try:
                    result = future.result()
                except Exception as exc:
                    fp = futures[future]
                    logger.error("Ошибка обработки %s: %s", fp, exc, exc_info=True)
                    continue

                if result is None:
                    continue  # файл не изменился

                # Если файл уже был в индексе — удалить старые векторы
                # (нужно и при изменении файла, и при апгрейде metadata→content)
                if result["was_indexed"]:
                    self._delete_file_vectors(result["filepath"])

                # Добавить метаданные и контентные чанки в буфер
                pending_texts.append(result["meta_text"])
                pending_payloads.append(result["meta_payload"])
                for chunk, cpayload in zip(result["chunks"], result["content_payloads"]):
                    pending_texts.append(chunk)
                    pending_payloads.append(cpayload)
                pending_states.append((result["file_key"], result["fingerprint"], result["mtime"]))

                # Достигли порога — кодируем и пишем в Qdrant
                if len(pending_texts) >= ENCODE_BATCH:
                    flush()

        flush()  # финальный батч (остаток)

        logger.info("Этап '%s' завершён. Добавлено точек за сессию: %d",
                    stage, self.point_count)

        # Чистим «фантомы» только когда имеем полный список всех файлов на диске
        # (т.е. на этапах metadata и content). На small/large мы видим только
        # часть файлов и не должны по этому основанию удалять других.
        if stage in ("metadata", "content"):
            self._cleanup_deleted_files(all_files)

        info = self.qdrant.get_collection(self.collection_name)
        logger.info("Коллекция '%s': %d точек", self.collection_name, info.points_count)

    def index_all_stages(self, stages: Optional[List[str]] = None) -> None:
        """
        Прогоняет индексирование последовательно по всем этапам.

        По умолчанию: metadata → small → large.
        После каждого этапа индекс УЖЕ пригоден для поиска, качество растёт
        прогрессивно. Если процесс прерывать/перезапускать — продолжит с того
        этапа, на котором остановился (благодаря полю `stage` в state.json).
        """
        stages = list(stages) if stages else list(STAGES)
        logger.info("▶ Многоэтапная индексация: %s", " → ".join(stages))
        for stage in stages:
            # point_count — счётчик на сессию, сбрасываем между этапами чтобы
            # логи «итого» были понятнее.
            self.point_count = 0
            self.index_directory(stage=stage)
        logger.info("✔ Все этапы завершены: %s", ", ".join(stages))

    def _cleanup_deleted_files(self, existing_files: List[Path]) -> None:
        """Удалить из state.json и Qdrant файлы, которых больше нет на диске."""
        existing_paths = {str(f) for f in existing_files}
        deleted_keys = [k for k in self.state["files"] if k not in existing_paths]
        if not deleted_keys:
            return
        logger.info("Удаление %d удалённых файлов из индекса…", len(deleted_keys))
        for key in deleted_keys:
            self._delete_file_vectors(Path(key))
            del self.state["files"][key]
        self._save_state()


# ─────────────────────────── CLI entry point ───────────────────────────

def main() -> None:
    cfg = load_config()

    # Добавляем FileHandler с UTF-8 (путь к логу из конфига)
    log_file = cfg.get("log_file")
    if log_file:
        try:
            fh = logging.FileHandler(log_file, encoding="utf-8")
            fh.setFormatter(logging.Formatter("%(asctime)s - %(levelname)s - %(message)s"))
            logging.getLogger().addHandler(fh)
        except Exception as exc:
            logger.warning("Не удалось открыть лог-файл %s: %s", log_file, exc)

    parser = argparse.ArgumentParser(
        description="RAG Indexer для DOCX/XLSX/XLS/PDF файлов",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Режимы работы:\n"
            "  (по умолчанию)        — многоэтапно: metadata → small → large\n"
            "  --stage metadata      — только имена/пути/размеры всех файлов (минуты)\n"
            "  --stage small         — содержимое docx/xlsx/xls + мелких PDF (десятки минут)\n"
            "  --stage large         — содержимое больших файлов + сканированные PDF (часы)\n"
            "  --stage all           — все этапы подряд (поведение по умолчанию)\n"
            "  --recreate            — пересоздать коллекцию и очистить state\n"
            "  --cleanup             — только удалить из индекса файлы, которых нет на диске\n"
        ),
    )
    parser.add_argument("--catalog", default=cfg["catalog_path"], help="Папка для индексирования")
    parser.add_argument("--db", default=cfg["qdrant_db_path"], help="Путь к локальной базе Qdrant (SQLite режим)")
    parser.add_argument("--url", default="", dest="qdrant_url",
                        help="URL Qdrant-сервера (например http://localhost:6333). Если указан — используется вместо --db")
    parser.add_argument("--model", default=cfg["embedding_model"], help="Модель эмбеддинга")
    parser.add_argument("--collection", default=cfg["collection_name"], help="Имя коллекции")
    parser.add_argument("--recreate", action="store_true", help="Пересоздать коллекцию и очистить state")
    parser.add_argument("--no-ocr", action="store_true", dest="no_ocr",
                        help="Пропускать OCR для сканированных PDF (быстрее, текст не извлекается)")
    parser.add_argument("--max-chunks", type=int, default=2000, dest="max_chunks",
                        help="Максимум чанков с одного файла (по умолчанию 2000; 0 = без ограничений)")
    parser.add_argument("--workers", type=int, default=4, dest="workers",
                        help="Число параллельных потоков для чтения файлов (по умолчанию 4)")
    parser.add_argument("--onnx", action="store_true", dest="use_onnx",
                        help="Использовать ONNX Runtime для encode (быстрее, но может не работать на Python 3.14)")
    parser.add_argument("--stage", default="all", choices=("all", *STAGES),
                        help="Этап индексирования. По умолчанию 'all' — прогон всех этапов "
                             "(metadata → small → large). Можно запустить отдельный этап для "
                             "дробного прогресса или для тонкой настройки фоновых задач.")
    parser.add_argument("--metadata-only-for", default="", dest="metadata_only_for",
                        help="[legacy] Список расширений через запятую для индексирования ТОЛЬКО "
                             "метаданных (например: .pdf). Используется в рамках одного stage. "
                             "Современный эквивалент: --stage metadata.")
    parser.add_argument("--metadata-only", action="store_true", dest="metadata_only",
                        help="[legacy] Псевдоним для --stage metadata.")
    parser.add_argument(
        "--cleanup",
        action="store_true",
        help=(
            "Только очистить индекс от файлов, которые удалены с диска, "
            "без полного сканирования. Быстро (~1 мин)."
        ),
    )
    parser.add_argument(
        "--mark-stage-metadata-for", default="", dest="mark_stage_metadata_for",
        help="Перед индексированием пометить уже имеющиеся в state записи указанных расширений "
             "как stage=metadata (чтобы они переиндексировались на этапах small/large). "
             "Пример: --mark-stage-metadata-for .pdf. Пригодится после legacy-прохода --metadata-only-for.")
    args = parser.parse_args()

    # Разбор stage и legacy-флагов
    metadata_only_extensions: set = set()
    stage = args.stage
    if args.metadata_only:
        # legacy: --metadata-only → --stage metadata
        stage = "metadata"
        logger.info("Legacy --metadata-only: использую --stage metadata")
    elif args.metadata_only_for:
        metadata_only_extensions = {
            e.strip().lower() if e.strip().startswith(".") else "." + e.strip().lower()
            for e in args.metadata_only_for.split(",") if e.strip()
        }
        logger.info("Legacy --metadata-only-for для расширений: %s (в рамках --stage %s)",
                    sorted(metadata_only_extensions), stage)

    indexer = RAGIndexer(
        catalog_path=args.catalog,
        qdrant_db_path=args.db,
        embedding_model=args.model,
        collection_name=args.collection,
        vector_size=cfg["vector_size"],
        chunk_size=cfg["chunk_size"],
        chunk_overlap=cfg["chunk_overlap"],
        batch_size=cfg["batch_size"],
        recreate_collection=args.recreate,
        skip_ocr=args.no_ocr,
        max_chunks_per_file=args.max_chunks,
        read_workers=args.workers,
        use_onnx=args.use_onnx,
        qdrant_url=args.qdrant_url,
        metadata_only_extensions=metadata_only_extensions,
    )

    # Миграция: пометить указанные расширения в state как stage=metadata
    # (это нужно, если раньше был прогон --metadata-only-for без поддержки stage,
    # и мы хотим, чтобы на этапах small/large эти файлы переиндексировались).
    if args.mark_stage_metadata_for:
        exts = {
            e.strip().lower() if e.strip().startswith(".") else "." + e.strip().lower()
            for e in args.mark_stage_metadata_for.split(",") if e.strip()
        }
        changed = 0
        for key, meta in indexer.state["files"].items():
            if Path(key).suffix.lower() in exts:
                if meta.get("stage") != "metadata":
                    meta["stage"] = "metadata"
                    changed += 1
        if changed:
            indexer._save_state()
        logger.info("Миграция state: %d записей с расширениями %s помечены stage=metadata",
                    changed, sorted(exts))

    if args.cleanup:
        # Режим только очистки: сканируем диск, удаляем «фантомы» и выходим
        logger.info("Режим --cleanup: поиск и удаление удалённых файлов из индекса…")
        all_files = [
            f
            for f in indexer.catalog_path.rglob("*")
            if f.is_file()
            and f.suffix.lower() in SUPPORTED_EXTENSIONS
            and not f.name.startswith("~$")
        ]
        logger.info("Файлов на диске: %d", len(all_files))
        indexer._cleanup_deleted_files(all_files)
        logger.info("Очистка завершена.")
    else:
        if stage == "all":
            indexer.index_all_stages()
        else:
            indexer.index_directory(stage=stage)


if __name__ == "__main__":
    main()
