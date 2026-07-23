"""
helpers.py — File, search, format, and explorer utility functions.

Depends on: .system, .state, core modules, nicegui.
Imported by: api.py, nice_app.py.
"""

from __future__ import annotations

import html
import json
import re
import sqlite3
import subprocess
import threading
import time
import urllib.error
import urllib.request
from collections import Counter
from datetime import datetime
from pathlib import Path, PureWindowsPath
from typing import Any, Dict, List, Optional
from urllib.parse import quote

from nicegui import ui

from rag_catalog.core.cloud_drive import CloudDriveService
from rag_catalog.core.exact_tokens import numeric_query_has_trusted_context
from rag_catalog.core.index_state_db import IndexStateDB
from rag_catalog.core.log_history import (
    iter_history_texts,
    list_log_segments,
    read_history_tail,
    read_history_tail_lines,
)
from rag_catalog.core.rag_core import RAGSearcher
from rag_catalog.core.user_auth_db import UserAuthDB

from .state import (
    PageState,
    _get_auth_db,
    _get_telemetry,
    _log_app_event,
    _username,
    _users_db_path,
)
from .system import (
    _STAGE_LABELS,
    PROJECT_ROOT,
    _find_module_process_pids,
    _process_matches_module,
    _telemetry_db_path,
)

# ─────────────────────────── constants ──────────────────────────────────────

_CADENCE_LABELS: Dict[str, str] = {
    "hourly": "Каждый час",
    "daily": "Ежедневно",
    "weekly": "Еженедельно",
}
_DAY_LABELS = ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]
_DAY_RU = {"Mon": "Пн", "Tue": "Вт", "Wed": "Ср", "Thu": "Чт", "Fri": "Пт", "Sat": "Сб", "Sun": "Вс"}

FILE_PREVIEW_EXTENSIONS = {".txt", ".log", ".csv", ".json", ".md", ".py", ".ps1", ".xml", ".html", ".css"}
INLINE_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".svg"}
OCR_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".tif", ".tiff", ".bmp", ".webp"}
OCR_CAPABLE_EXTENSIONS = {".pdf", *OCR_IMAGE_EXTENSIONS}
OCR_INVENTORY_CACHE_TTL_SECONDS = 60.0
_OCR_INVENTORY_CACHE: Dict[str, Any] = {}
OFFICE_PREVIEW_EXTENSIONS = {".docx", ".xlsx", ".xls", ".rtf"}
PAGE_SIZE = 80
SYSTEM_FILE_EXTENSIONS = {
    ".dll",
    ".exe",
    ".msi",
    ".sys",
    ".inf",
    ".cat",
    ".pnf",
    ".bak",
    ".tmp",
    ".temp",
    ".dat",
    ".db",
    ".db3",
    ".sqlite",
    ".sqlite-journal",
    ".db-shm",
    ".db-wal",
    ".lock",
    ".cfg",
    ".ini",
    ".config",
    ".lnk",
    ".chm",
    ".toc",
    ".cldbin",
}


# ─────────────────────────── URL / path helpers ─────────────────────────────

def _file_url(full_path: str) -> str:
    try:
        p = PureWindowsPath(full_path)
        if not p.parts:
            return ""
        drive = p.drive
        encoded = "/".join(quote(part, safe="") for part in p.parts[1:])
        return "file:///" + drive + "/" + encoded
    except Exception:
        return ""


def _folder_url(full_path: str) -> str:
    try:
        p = PureWindowsPath(full_path).parent
        if not p.parts:
            return ""
        drive = p.drive
        encoded = "/".join(quote(part, safe="") for part in p.parts[1:])
        return "file:///" + drive + "/" + encoded
    except Exception:
        return ""


def _viewer_file_url(full_path: str) -> str:
    value = str(full_path or "").strip()
    if not value:
        return ""
    return f"/api/view-file?path={quote(value, safe='')}"


def _schedule_display_label(sched: Dict[str, Any]) -> str:
    label = str(sched.get("label") or "").strip()
    if label:
        return label
    stage_label = _STAGE_LABELS.get(str(sched.get("stage") or "all"), str(sched.get("stage") or "Все этапы"))
    cadence_label = _CADENCE_LABELS.get(str(sched.get("cadence") or "daily"), str(sched.get("cadence") or "ежедневно"))
    return f"{stage_label} · {cadence_label}"


def _resolve_catalog_file(cfg: Dict[str, Any], raw_path: str) -> Optional[Path]:
    try:
        catalog = Path(str(cfg.get("catalog_path") or "")).resolve()
        if not catalog.exists() or not catalog.is_dir():
            return None
        candidate = Path(str(raw_path or "")).resolve()
        candidate.relative_to(catalog)
    except Exception:
        return None
    if not candidate.exists() or not candidate.is_file():
        return None
    return candidate


# ─────────────────────────── file preview ───────────────────────────────────

def _preview_file(path: Path, limit: int = 6000) -> str:
    if not path.exists() or not path.is_file():
        return "Файл недоступен."
    if path.suffix.lower() not in FILE_PREVIEW_EXTENSIONS:
        return "Для этого типа файла доступно открытие или скачивание."
    try:
        return path.read_text(encoding="utf-8", errors="replace")[:limit]
    except Exception as exc:
        return f"Не удалось прочитать файл: {exc}"


def _preview_office_file(path: Path, limit: int = 12000) -> str:
    ext = path.suffix.lower()
    if ext == ".docx":
        try:
            from docx import Document  # noqa: PLC0415

            doc = Document(path)
            parts: List[str] = [p.text for p in doc.paragraphs if p.text]
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [str(cell.text or "").strip() for cell in row.cells]
                    line = " | ".join(item for item in row_cells if item)
                    if line:
                        parts.append(line)
            return "\n".join(parts)[:limit] or "Документ пуст."
        except Exception as exc:
            return f"Не удалось прочитать DOCX: {exc}"
    if ext == ".xlsx":
        try:
            from openpyxl import load_workbook  # noqa: PLC0415

            wb = load_workbook(path, read_only=True, data_only=True)
            parts: List[str] = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                parts.append(f"Лист: {sheet_name}")
                for row in ws.iter_rows(values_only=True):
                    row_text = " | ".join(str(c) if c is not None else "" for c in row)
                    if row_text.strip():
                        parts.append(row_text)
                    if len("\n".join(parts)) >= limit:
                        return "\n".join(parts)[:limit]
            return "\n".join(parts)[:limit] or "Таблица пуста."
        except Exception as exc:
            return f"Не удалось прочитать XLSX: {exc}"
    if ext == ".xls":
        try:
            import xlrd  # type: ignore  # noqa: PLC0415

            book = xlrd.open_workbook(str(path))
            parts = []
            for sheet in book.sheets():
                parts.append(f"Лист: {sheet.name}")
                for row_idx in range(sheet.nrows):
                    row = sheet.row_values(row_idx)
                    row_text = " | ".join(str(v) if v not in ("", None) else "" for v in row)
                    if row_text.strip():
                        parts.append(row_text)
                    if len("\n".join(parts)) >= limit:
                        return "\n".join(parts)[:limit]
            return "\n".join(parts)[:limit] or "Таблица пуста."
        except Exception as exc:
            return f"Не удалось прочитать XLS: {exc}"
    if ext == ".rtf":
        try:
            raw = path.read_text(encoding="utf-8", errors="replace")
            no_ctrl = re.sub(r"\\[a-z]+\d* ?", " ", raw)
            no_braces = re.sub(r"[{}]", "", no_ctrl)
            clean = re.sub(r"\s+", " ", no_braces).strip()
            return clean[:limit] or "Документ пуст."
        except Exception as exc:
            return f"Не удалось прочитать RTF: {exc}"
    return "Для этого типа файла доступно открытие или скачивание."


# ─────────────────────────── DB query helper ────────────────────────────────

def _db_query_dicts(db_path: Path, query: str, params: Optional[tuple] = None) -> List[Dict[str, Any]]:
    if not db_path.exists():
        return []
    conn = sqlite3.connect(str(db_path), timeout=30.0)
    conn.row_factory = sqlite3.Row
    try:
        conn.execute("PRAGMA busy_timeout=30000;")
        cur = conn.execute(query, params or ())
        return [dict(row) for row in cur.fetchall()]
    except Exception:
        return []
    finally:
        conn.close()


# ─────────────────────────── log readers ────────────────────────────────────

def _read_log_tail(path: Path, *, max_chars: int = 12000) -> str:
    try:
        if not path.exists() and not list_log_segments(path):
            return "Лог-файл не найден."
        text = read_history_tail(path, max_chars=max_chars)
        if not text:
            return "Лог-файл пуст."
        return text
    except Exception as exc:
        return f"Не удалось прочитать лог: {exc}"


def _read_log_tail_lines(path: Path, *, max_lines: int = 200, max_chars: int = 200_000) -> str:
    try:
        if not path.exists() and not list_log_segments(path):
            return "Лог-файл не найден."
        text = read_history_tail_lines(path, max_lines=max_lines, max_chars=max_chars)
        if not text:
            return "Лог-файл пуст."
        return text
    except Exception as exc:
        return f"Не удалось прочитать лог: {exc}"


def _filter_log_text(text: str, level: str) -> str:
    level_key = str(level or "all").strip().lower()
    if level_key in {"", "all"}:
        return text
    token = f" - {level_key.upper()} - "
    lines = [line for line in str(text or "").splitlines() if token in line.upper()]
    return "\n".join(lines) if lines else "Записей для выбранного уровня не найдено."


_LOG_LINE_RE = re.compile(
    r"^(\d{4}-\d{2}-\d{2}) (\d{2}:\d{2}:\d{2}[.,]\d+) - (DEBUG|INFO|WARNING|ERROR|CRITICAL) - (.*)$"
)


def _parse_log_lines(text: str) -> List[Dict[str, Any]]:
    entries: List[Dict[str, Any]] = []
    current: Optional[Dict[str, Any]] = None
    for line in text.splitlines():
        m = _LOG_LINE_RE.match(line)
        if m:
            if current is not None:
                entries.append(current)
            current = {
                "date": m.group(1),
                "time": m.group(2),
                "level": m.group(3),
                "message": m.group(4),
            }
        elif current is not None:
            current["message"] += "\n" + line
    if current is not None:
        entries.append(current)
    return entries


def _read_log_entries(
    path: Path,
    *,
    max_entries: int = 200,
    level: str = "all",
    query: str = "",
    date_from: str = "",
    date_to: str = "",
    max_read_chars: int = 4_000_000,
) -> List[Dict[str, Any]]:
    """Read log entries from end of file, scan back until max_entries matching entries found."""
    try:
        if not path.exists() and not list_log_segments(path):
            return []
        level_key = str(level or "all").strip().upper()
        query_key = str(query or "").strip().lower()
        date_from = str(date_from or "").strip()
        date_to = str(date_to or "").strip()
        collected: List[Dict[str, Any]] = []
        for text in iter_history_texts(path, newest_first=True, max_chars_per_file=max_read_chars):
            if not text:
                continue
            all_entries = _parse_log_lines(text)
            for entry in reversed(all_entries):
                if level_key not in {"", "ALL"} and entry["level"] != level_key:
                    continue
                if date_from and entry["date"] < date_from:
                    continue
                if date_to and entry["date"] > date_to:
                    continue
                if query_key and query_key not in (
                    f"{entry.get('date', '')} {entry.get('time', '')} "
                    f"{entry.get('level', '')} {entry.get('message', '')}"
                ).lower():
                    continue
                collected.append(entry)
                if len(collected) >= max_entries:
                    break
            if len(collected) >= max_entries:
                break
        return collected
    except Exception:
        return []


_LEVEL_ROW_STYLE: Dict[str, str] = {
    "DEBUG":    "color:#6b7280",
    "INFO":     "color:#1e293b",
    "WARNING":  "color:#92400e;background:#fffbeb",
    "ERROR":    "color:#991b1b;background:#fef2f2",
    "CRITICAL": "color:#7c2d12;background:#fee2e2;font-weight:700",
}
_LEVEL_BADGE_COLOR: Dict[str, str] = {
    "DEBUG":    "#9ca3af",
    "INFO":     "#6366f1",
    "WARNING":  "#d97706",
    "ERROR":    "#dc2626",
    "CRITICAL": "#7c2d12",
}


def _format_log_entries_html(entries: List[Dict[str, Any]]) -> str:
    if not entries:
        return '<div style="color:#9ca3af;padding:16px;text-align:center">Нет записей для выбранных фильтров.</div>'
    parts: List[str] = []
    last_date = ""
    for entry in entries:
        date = entry.get("date", "")
        if date != last_date:
            last_date = date
            esc_date = html.escape(date)
            parts.append(
                f'<div style="display:flex;align-items:center;gap:8px;margin:8px 2px 4px">'
                f'<div style="flex:1;height:1px;background:#e5e7eb"></div>'
                f'<span style="font-size:10px;color:#9ca3af;white-space:nowrap">{esc_date}</span>'
                f'<div style="flex:1;height:1px;background:#e5e7eb"></div>'
                f'</div>'
            )
        level = entry.get("level", "INFO")
        row_style = _LEVEL_ROW_STYLE.get(level, "color:#1e293b")
        badge_color = _LEVEL_BADGE_COLOR.get(level, "#6b7280")
        time_str = html.escape(entry.get("time", ""))
        msg = html.escape(entry.get("message", "")).replace("\n", "<br>&nbsp;&nbsp;")
        level_esc = html.escape(level)
        parts.append(
            f'<div style="padding:2px 6px;border-radius:3px;margin-bottom:1px;line-height:1.5;{row_style}">'
            f'<span style="color:#9ca3af;font-size:10px;user-select:none">{time_str}</span> '
            f'<span style="display:inline-block;padding:0 4px;border-radius:3px;font-size:10px;'
            f'font-weight:600;color:#fff;background:{badge_color};user-select:none">{level_esc}</span> '
            f'<span>{msg}</span>'
            f'</div>'
        )
    return "".join(parts)


def _format_log_entries_text(entries: List[Dict[str, Any]]) -> str:
    lines: List[str] = []
    last_date = ""
    for entry in entries:
        date = entry.get("date", "")
        if date != last_date:
            last_date = date
            lines.append(f"──── {date} ────")
        lines.append(f"{entry.get('time', '')} [{entry.get('level', '')}] {entry.get('message', '')}")
    return "\n".join(lines)


# ─────────────────────────── query history helpers ──────────────────────────

def _dedupe_queries(values: List[str], limit: int = 12) -> List[str]:
    out: List[str] = []
    seen = set()
    for value in values:
        query = re.sub(r"\s+", " ", str(value or "")).strip()
        key = query.lower()
        if not query or key in seen:
            continue
        seen.add(key)
        out.append(query)
        if len(out) >= limit:
            break
    return out


def _cloud_query_set(cfg: Dict[str, Any], username: str = "") -> "set[str]":
    """Return set of lowercase query strings that previously returned ≥1 Cloud Drive result for this user."""
    try:
        from rag_catalog.core.telemetry_db import TelemetryDB
        path = _telemetry_db_path(cfg)
        telemetry = TelemetryDB(str(path))
        events = telemetry.list_app_events(feature="search", action="search", username=username or None, limit=200)
        out: set[str] = set()
        for ev in events:
            details = ev.get("details") or {}
            if int(details.get("cloud_results") or 0) > 0:
                q = str(details.get("query_original") or details.get("query") or "").strip().lower()
                if q:
                    out.add(q)
        return out
    except Exception:
        return set()


def _my_recent_queries(cfg: Dict[str, Any], username: str = "", limit: int = 12) -> List[str]:
    rows = _db_query_dicts(
        _telemetry_db_path(cfg),
        """
        SELECT COALESCE(NULLIF(query_original, ''), query) AS query
        FROM search_logs
        WHERE query <> '' AND lower(username) = ?
        ORDER BY id DESC
        LIMIT 80
        """,
        (username.strip().lower(),),
    )
    return _dedupe_queries([str(row.get("query") or "") for row in rows], limit=limit)


def _popular_queries(cfg: Dict[str, Any], exclude_username: str = "", limit: int = 10) -> List[str]:
    rows = _db_query_dicts(
        _telemetry_db_path(cfg),
        """
        SELECT COALESCE(NULLIF(query_original, ''), query) AS query, COUNT(*) AS cnt
        FROM search_logs
        WHERE query <> '' AND lower(username) != ?
        GROUP BY lower(COALESCE(NULLIF(query_original, ''), query))
        ORDER BY cnt DESC
        LIMIT 40
        """,
        (exclude_username.strip().lower(),),
    )
    return _dedupe_queries([str(row.get("query") or "") for row in rows], limit=limit)


_POPULAR_QUERY_STOPWORDS = {
    "and",
    "the",
    "или",
    "для",
    "как",
    "что",
    "это",
    "где",
    "при",
    "над",
    "под",
    "без",
    "все",
    "всех",
    "найти",
    "нужен",
    "нужно",
    "документ",
    "документы",
    "файл",
    "файлы",
    "скан",
    "type",
    "path",
    "after",
    "before",
    "creator",
    "editor",
    "from",
}


def _query_keyword_candidates(query: str) -> List[str]:
    text = str(query or "").replace("ё", "е").lower()
    words = re.findall(r"[a-zа-я0-9][a-zа-я0-9\-]{2,}", text, flags=re.IGNORECASE)
    out: List[str] = []
    for word in words:
        token = word.strip("-")
        if not token or token in _POPULAR_QUERY_STOPWORDS:
            continue
        if token.isdigit() or ":" in token:
            continue
        if token in {"doc", "docx", "xls", "xlsx", "pdf", "jpg", "png"}:
            continue
        out.append(token)
    return out


def _popular_query_terms(cfg: Dict[str, Any], exclude_username: str = "", limit: int = 6) -> List[str]:
    rows = _db_query_dicts(
        _telemetry_db_path(cfg),
        """
        SELECT COALESCE(NULLIF(query_original, ''), query) AS query, COUNT(*) AS cnt
        FROM search_logs
        WHERE query <> '' AND lower(username) != ?
        GROUP BY lower(COALESCE(NULLIF(query_original, ''), query))
        ORDER BY cnt DESC
        LIMIT 120
        """,
        (exclude_username.strip().lower(),),
    )
    counter: Counter[str] = Counter()
    for row in rows:
        weight = max(1, int(row.get("cnt") or 1))
        for token in dict.fromkeys(_query_keyword_candidates(str(row.get("query") or ""))):
            counter[token] += weight
    return [word for word, _count in counter.most_common(max(1, int(limit)))]


def _search_suggestions(state: PageState, typed: str = "") -> List[str]:
    username = _username(state)
    personal = _dedupe_queries([*state.history, *_my_recent_queries(state.cfg, username, limit=12)], limit=24)
    needle = typed.strip().lower()
    if not needle:
        return personal[:12]
    starts = [item for item in personal if item.lower().startswith(needle)]
    contains = [item for item in personal if needle in item.lower() and item not in starts]
    return [*starts, *contains][:12]


def _telegram_deeplink(bot_link: str, purpose: str, token: str) -> str:
    base = str(bot_link or "").strip()
    value = str(token or "").strip()
    if not base or not value:
        return ""
    joiner = "&" if "?" in base else "?"
    return f"{base}{joiner}start={purpose}_{quote(value, safe='')}"


def _remember_query(state: PageState, query: str) -> None:
    clean = re.sub(r"\s+", " ", str(query or "")).strip()
    if clean:
        state.history = _dedupe_queries([clean, *state.history], limit=24)


# ─────────────────────────── query parser ────────────────────────────────────

def _parse_search_query(raw: str) -> Dict[str, Any]:
    """Parse search operators out of raw query string.

    Order matters: phrases and filters are removed first so они don't pollute
    boolean splitting; wildcards are extracted last so дог* или акт preserves OR.
    """
    # 1. Quoted phrases → must match exactly
    must_phrases: List[str] = re.findall(r'"([^"]+)"', raw)
    q = re.sub(r'"[^"]+"', ' ', raw)

    # 2. Excluded words (-слово)
    tokens = q.split()
    excluded_words: List[str] = [t[1:].lower() for t in tokens if t.startswith('-') and len(t) > 1]
    tokens = [t for t in tokens if not (t.startswith('-') and len(t) > 1)]
    q = ' '.join(tokens)

    # 3. Structured filters (consume them before boolean splitting)
    file_type_filter: Optional[str] = None
    m = re.search(r'\btype:(\.?\w+)', q, re.IGNORECASE)
    if m:
        ft = m.group(1).lower()
        file_type_filter = ft if ft.startswith('.') else '.' + ft
        q = (q[:m.start()] + q[m.end():]).strip()

    date_from: Optional[str] = None
    m = re.search(r'\bafter:(\d{4}-\d{2}-\d{2})', q, re.IGNORECASE)
    if m:
        date_from = m.group(1)
        q = (q[:m.start()] + q[m.end():]).strip()

    date_to: Optional[str] = None
    m = re.search(r'\bbefore:(\d{4}-\d{2}-\d{2})', q, re.IGNORECASE)
    if m:
        date_to = m.group(1)
        q = (q[:m.start()] + q[m.end():]).strip()

    path_filter: Optional[str] = None
    m = re.search(r'\bpath:(\S+)', q, re.IGNORECASE)
    if m:
        path_filter = m.group(1).lower()
        q = (q[:m.start()] + q[m.end():]).strip()

    from_filter: Optional[str] = None
    m = re.search(r'\bfrom:(\S+)', q, re.IGNORECASE)
    if m:
        from_filter = m.group(1).lower()
        q = (q[:m.start()] + q[m.end():]).strip()

    creator_filter: Optional[str] = None
    m = re.search(r'\bcreator:(\S+)', q, re.IGNORECASE)
    if m:
        creator_filter = m.group(1).lower()
        q = (q[:m.start()] + q[m.end():]).strip()

    editor_filter: Optional[str] = None
    m = re.search(r'\beditor:(\S+)', q, re.IGNORECASE)
    if m:
        editor_filter = m.group(1).lower()
        q = (q[:m.start()] + q[m.end():]).strip()

    # 4. Boolean OR/AND splitting — wildcards intentionally kept so дог* или акт
    #    produces two OR branches, not one prefix term and a broken OR.
    or_branches = re.split(r'(?<!\w)или(?!\w)', q, flags=re.IGNORECASE)
    raw_bool_groups: List[List[str]] = []
    for branch in or_branches:
        and_terms = re.split(r'(?<!\w)и(?!\w)', branch, flags=re.IGNORECASE)
        group = [t.strip() for t in and_terms if t.strip()]
        if group:
            raw_bool_groups.append(group)

    # 5. Extract wildcards from bool_group terms (keeps them in groups for matching)
    prefix_terms: List[str] = []
    bool_groups: List[List[str]] = []
    for group in raw_bool_groups:
        clean_group: List[str] = []
        for term in group:
            words = term.split()
            for w in words:
                if w.endswith('*') and len(w) > 1:
                    prefix_terms.append(w[:-1].lower())
            # Keep the original term (with *) in bool_groups so _apply can use wildcard matching
            clean_group.append(term)
        bool_groups.append(clean_group)

    has_bool = len(bool_groups) > 1 or any(len(g) > 1 for g in bool_groups)

    def _strip_star(t: str) -> str:
        return ' '.join(w[:-1] if w.endswith('*') and len(w) > 1 else w for w in t.split())

    clean_terms = [_strip_star(t) for group in bool_groups for t in group]
    semantic_query = ' '.join(clean_terms).strip() or ' '.join(must_phrases)

    has_operators = bool(
        must_phrases or excluded_words or prefix_terms
        or file_type_filter or date_from or date_to or path_filter
        or from_filter or creator_filter or editor_filter
        or has_bool
    )
    return {
        "semantic_query": semantic_query,
        "must_phrases": must_phrases,
        "excluded_words": excluded_words,
        "prefix_terms": prefix_terms,
        "file_type_filter": file_type_filter,
        "date_from": date_from,
        "date_to": date_to,
        "path_filter": path_filter,
        "from_filter": from_filter,
        "creator_filter": creator_filter,
        "editor_filter": editor_filter,
        "bool_groups": bool_groups,
        "has_bool": has_bool,
        "has_operators": has_operators,
    }


def _apply_query_operators(results: List[Dict[str, Any]], parsed: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Post-filter search results based on parsed query operators."""
    if not parsed.get("has_operators"):
        return results

    must_phrases = [p.lower() for p in parsed.get("must_phrases", [])]
    excluded_words = [w.lower() for w in parsed.get("excluded_words", [])]
    prefix_terms = [p.lower() for p in parsed.get("prefix_terms", [])]
    file_type_filter: Optional[str] = parsed.get("file_type_filter")
    date_from: Optional[str] = parsed.get("date_from")
    date_to: Optional[str] = parsed.get("date_to")
    path_filter: Optional[str] = parsed.get("path_filter")
    from_filter: Optional[str] = parsed.get("from_filter")
    creator_filter: Optional[str] = parsed.get("creator_filter")
    editor_filter: Optional[str] = parsed.get("editor_filter")
    bool_groups: List[List[str]] = parsed.get("bool_groups", [])
    has_bool: bool = parsed.get("has_bool", False)

    def _term_in(term: str, text: str, words: List[str]) -> bool:
        """Match a single term against text; handles trailing * as prefix wildcard."""
        if term.endswith('*'):
            pfx = term[:-1].lower()
            return bool(pfx) and any(w.startswith(pfx) for w in words)
        return term.lower() in text

    filtered = []
    for item in results:
        text = ' '.join([
            str(item.get("text") or ""),
            str(item.get("chunk_text") or ""),
            str(item.get("filename") or ""),
            str(item.get("path") or ""),
        ]).lower()
        words = text.split()
        path = str(item.get("path") or item.get("full_path") or item.get("cloud_path") or "").lower()
        fname = str(item.get("filename") or "").lower()
        ext = str(item.get("extension") or "").lower()
        if ext and not ext.startswith('.'):
            ext = '.' + ext
        doc_author = str(item.get("doc_author") or "").lower()
        doc_last_editor = str(item.get("doc_last_editor") or "").lower()
        doc_top_editor = str(item.get("doc_top_editor") or "").lower()

        if file_type_filter:
            if ext != file_type_filter and not fname.endswith(file_type_filter):
                continue

        d = str(item.get("modified") or "")[:10]
        if date_from and d and d < date_from:
            continue
        if date_to and d and d > date_to:
            continue

        if path_filter and path_filter not in path:
            continue

        # from: searches path + all author fields
        if from_filter and not (
            from_filter in path
            or from_filter in doc_author
            or from_filter in doc_last_editor
            or from_filter in doc_top_editor
        ):
            continue

        # creator: matches only the original document author
        if creator_filter and creator_filter not in doc_author:
            continue

        # editor: matches last editor or top tracked-change editor
        if editor_filter and not (
            editor_filter in doc_last_editor or editor_filter in doc_top_editor
        ):
            continue

        if any(ew in text for ew in excluded_words):
            continue

        if must_phrases and not all(ph in text for ph in must_phrases):
            continue

        # Standalone prefix check only when no boolean groups control the logic
        if not has_bool and prefix_terms:
            if not all(any(w.startswith(pt) for w in words) for pt in prefix_terms):
                continue

        # Boolean OR-of-ANDs: pass if ANY OR-branch has ALL its AND-terms matching
        if has_bool and bool_groups:
            if not any(
                all(_term_in(term, text, words) for term in grp)
                for grp in bool_groups
            ):
                continue

        filtered.append(item)

    return filtered


# ─────────────────────────── search runners ─────────────────────────────────

def _normalize_search_results(results: Any) -> List[Dict[str, Any]]:
    if results is None:
        return []
    if not isinstance(results, list):
        return []
    return [item for item in results if isinstance(item, dict)]


def _relevance_gate_enabled(searcher: RAGSearcher) -> bool:
    config = getattr(searcher, "config", {})
    return isinstance(config, dict) and bool(config.get("retrieval_relevance_gate_enabled"))


def _run_catalog_search(
    searcher: RAGSearcher,
    *,
    query: str,
    query_original: str,
    query_used: str,
    limit: int,
    file_type: Optional[str],
    content_only: bool,
    title_only: bool,
    username: str = "",
) -> List[Dict[str, Any]]:
    results = _normalize_search_results(
        searcher.search(
            query,
            limit=limit,
            file_type=file_type,
            content_only=content_only,
            title_only=title_only,
            source="nicegui",
            username=username,
            query_original=query_original,
        )
    )
    if results or content_only or title_only:
        return results
    if _relevance_gate_enabled(searcher):
        return []

    try:
        fallback = searcher._lexical_catalog_search(  # noqa: SLF001
            query=query_used,
            limit=max(limit, 10),
            file_type=file_type,
            content_only=False,
            title_only=title_only,
        )
    except Exception:
        return results
    return _normalize_search_results(fallback)[:limit]


def _run_quick_name_search(
    searcher: RAGSearcher,
    *,
    query: str,
    limit: int,
    file_type: Optional[str],
) -> List[Dict[str, Any]]:
    try:
        exact = _normalize_search_results(
            searcher._numeric_exact_search(  # noqa: SLF001
                query=query,
                limit=limit,
                file_type=file_type,
                content_only=False,
                title_only=False,
            )
        )
    except Exception:
        exact = []
    if not numeric_query_has_trusted_context(query):
        exact = []

    quick = _normalize_search_results(
        searcher._lexical_catalog_search(  # noqa: SLF001
            query=query,
            limit=max(limit, 40),
            file_type=file_type,
            content_only=False,
            title_only=True,
        )
    )
    needle = re.sub(r"\s+", " ", str(query or "")).strip().lower().replace("ё", "е")

    def sort_key(item: Dict[str, Any]) -> tuple:
        name = str(item.get("filename") or "").lower().replace("ё", "е")
        path = str(item.get("path") or "").lower().replace("ё", "е")
        exact_name = 1 if needle and needle in name else 0
        exact_path = 1 if needle and needle in path else 0
        score = float(item.get("score") or 0.0)
        return (exact_name, exact_path, score, -len(path))

    quick.sort(key=sort_key, reverse=True)
    merged = _merge_search_results(exact, quick, limit=limit)
    if not _relevance_gate_enabled(searcher):
        return merged

    confident: List[Dict[str, Any]] = []
    for item in merged:
        source = str(item.get("retrieval_source") or "")
        if source in {"numeric_fs_exact", "numeric_exact"} and bool(
            item.get("numeric_query_trusted_context", numeric_query_has_trusted_context(query))
        ):
            confident.append(item)
            continue
        name = str(item.get("filename") or "").lower().replace("ё", "е")
        path = str(item.get("path") or item.get("full_path") or "").lower().replace("ё", "е")
        if needle and (needle in name or needle in path):
            confident.append(item)
    return confident[:limit]


def _count_exact_name_matches(query: str, results: List[Dict[str, Any]]) -> int:
    needle = re.sub(r"\s+", " ", str(query or "")).strip().lower().replace("ё", "е")
    if not needle:
        return 0
    count = 0
    for item in results:
        if (
            str(item.get("retrieval_source") or "") in {"numeric_fs_exact", "numeric_exact"}
            and bool(item.get("numeric_query_trusted_context", numeric_query_has_trusted_context(query)))
        ):
            count += 1
            continue
        name = str(item.get("filename") or "").lower().replace("ё", "е")
        path = str(item.get("path") or "").lower().replace("ё", "е")
        if needle in name or needle in path:
            count += 1
    return count


def _has_confident_numeric_exact_match(query: str, results: List[Dict[str, Any]]) -> bool:
    trusted_query = numeric_query_has_trusted_context(query)
    return any(
        str(item.get("retrieval_source") or "") in {"numeric_fs_exact", "numeric_exact"}
        and bool(item.get("numeric_query_trusted_context", trusted_query))
        for item in results
    )


def _merge_search_results(
    primary: List[Dict[str, Any]],
    secondary: List[Dict[str, Any]],
    *,
    limit: int,
) -> List[Dict[str, Any]]:
    merged_by_key: Dict[str, Dict[str, Any]] = {}

    def key_of(item: Dict[str, Any]) -> str:
        identity = (
            item.get("cloud_file_id")
            or item.get("full_path")
            or item.get("cloud_path")
            or item.get("path")
            or item.get("filename")
            or ""
        )
        return f"{identity}::{item.get('chunk_index')}::{item.get('type')}"

    def rank_key(item: Dict[str, Any]) -> tuple:
        rank_score = float(item.get("rank_score", item.get("score") or 0) or 0)
        score = float(item.get("score") or 0)
        path = str(item.get("path") or item.get("full_path") or "")
        return (rank_score, score, -len(path))

    for item in [*primary, *secondary]:
        key = key_of(item)
        existing = merged_by_key.get(key)
        if existing is None or rank_key(item) > rank_key(existing):
            merged_by_key[key] = item

    ranked = sorted(merged_by_key.values(), key=rank_key, reverse=True)
    return ranked[:limit]


# ─────────────────────────── formatters ─────────────────────────────────────

def _format_file_size(size_b: int) -> str:
    if size_b >= 1_048_576:
        return f"{size_b / 1_048_576:.1f} МБ"
    if size_b >= 1024:
        return f"{size_b / 1024:.1f} КБ"
    return f"{size_b} Б"


def _clean_text(value: Any) -> str:
    raw = str(value or "")
    raw = re.sub(r"<[^>]{1,200}>", " ", raw)
    raw = html.unescape(raw)
    raw = re.sub(r"[ \t]+", " ", raw)
    raw = re.sub(r"\s*\n\s*", "\n", raw)
    return raw.strip()


def _format_bytes(value: Any) -> str:
    try:
        size = float(value or 0)
    except (TypeError, ValueError):
        size = 0.0
    units = ["Б", "КБ", "МБ", "ГБ", "ТБ"]
    idx = 0
    while size >= 1024 and idx < len(units) - 1:
        size /= 1024
        idx += 1
    if idx == 0:
        return f"{int(size)} {units[idx]}"
    return f"{size:.1f} {units[idx]}"


def _format_duration_seconds(value: Any) -> str:
    try:
        seconds = int(float(value or 0))
    except (TypeError, ValueError):
        seconds = 0
    if seconds <= 0:
        return "0 сек"
    hours, rem = divmod(seconds, 3600)
    minutes, secs = divmod(rem, 60)
    if hours:
        return f"{hours} ч {minutes:02d} мин"
    if minutes:
        return f"{minutes} мин {secs:02d} сек"
    return f"{secs} сек"


def _format_relative_time(value: Any) -> str:
    raw = str(value or "").strip()
    if not raw:
        return "не запускался"
    try:
        dt = datetime.fromisoformat(raw.replace("Z", "+00:00"))
    except ValueError:
        try:
            dt = datetime.strptime(raw[:19], "%Y-%m-%d %H:%M:%S")
        except ValueError:
            return raw[:16]
    if dt.tzinfo is None:
        now = datetime.now()
    else:
        now = datetime.now(dt.tzinfo)
    seconds = max(0, int((now - dt).total_seconds()))
    if seconds < 60:
        return "только что"
    minutes = seconds // 60
    if minutes < 60:
        return f"{minutes} мин назад"
    hours = minutes // 60
    if hours < 24:
        return f"{hours} ч назад"
    days = hours // 24
    if days < 30:
        return f"{days} дн назад"
    months = days // 30
    if months < 12:
        return f"{months} мес назад"
    years = months // 12
    return f"{years} г назад"


def _duration_between(start: Any, finish: Any = None) -> int:
    start_s = str(start or "")
    finish_s = str(finish or "")
    if not start_s:
        return 0
    try:
        start_dt = datetime.fromisoformat(start_s)
        finish_dt = datetime.fromisoformat(finish_s) if finish_s else datetime.now(start_dt.tzinfo)
        return max(0, int((finish_dt - start_dt).total_seconds()))
    except ValueError:
        return 0


# ─────────────────────────── file / directory helpers ───────────────────────

def _directory_children(path: str, limit: int = 60) -> Dict[str, Any]:
    p = Path(path)
    out: Dict[str, Any] = {"exists": p.exists(), "is_dir": p.is_dir(), "dirs": [], "files": []}
    if not p.exists() or not p.is_dir():
        return out
    try:
        entries = sorted(
            [x for x in p.iterdir() if not x.name.startswith(".") and not x.name.startswith("~$")],
            key=lambda x: (not x.is_dir(), x.name.lower()),
        )
    except Exception as exc:
        out["error"] = str(exc)
        return out
    for child in entries[:limit]:
        item = {"name": child.name, "path": str(child)}
        if child.is_dir():
            out["dirs"].append(item)
        elif child.is_file():
            try:
                item["size"] = _format_file_size(child.stat().st_size)
            except Exception:
                item["size"] = ""
            out["files"].append(item)
    out["truncated"] = len(entries) > limit
    return out


def _open_os_path(path: str) -> None:
    value = str(path or "").strip()
    if not value:
        return
    try:
        subprocess.Popen(["explorer", value])
    except Exception as exc:
        ui.notify(f"Не удалось открыть проводник ОС: {exc}", type="negative")


def _select_in_os_explorer(path: str) -> None:
    """Open Windows Explorer with the file selected (explorer /select,<path>)."""
    value = str(path or "").strip()
    if not value:
        return
    try:
        subprocess.Popen(["explorer", f"/select,{value}"])
    except Exception as exc:
        ui.notify(f"Не удалось открыть проводник ОС: {exc}", type="negative")


def _within_catalog(root: Path, candidate: Path) -> bool:
    try:
        candidate.resolve().relative_to(root.resolve())
        return True
    except Exception:
        return False


def _safe_explorer_path(state: PageState) -> Path:
    root = Path(str(state.cfg.get("catalog_path") or ""))
    if state.explorer_path:
        candidate = Path(state.explorer_path)
        if candidate.exists() and _within_catalog(root, candidate):
            return candidate
    state.explorer_path = str(root)
    return root


# ─────────────────────────── Cloud Drive helpers ─────────────────────────────

_CD_SERVICE_CACHE: Dict[tuple[str, ...], CloudDriveService] = {}
_CD_SERVICE_CACHE_LOCK = threading.Lock()


def _cd_service_cache_key(cfg: Dict[str, Any]) -> tuple[str, ...]:
    return tuple(
        str(cfg.get(key) or "").strip()
        for key in (
            "cloud_drive_db_path",
            "cloud_drive_storage",
            "cloud_drive_storage_root",
            "cloud_drive_bucket",
            "cloud_drive_s3_endpoint",
            "cloud_drive_s3_region",
            "cloud_drive_s3_access_key",
            "cloud_drive_s3_secret_key",
        )
    )


def _cd_cached_service(cfg: Dict[str, Any]) -> Optional["CloudDriveService"]:
    try:
        if not str(cfg.get("cloud_drive_db_path") or "").strip():
            return None
        key = _cd_service_cache_key(cfg)
        with _CD_SERVICE_CACHE_LOCK:
            service = _CD_SERVICE_CACHE.get(key)
            if service is None:
                service = CloudDriveService.from_config(cfg)
                _CD_SERVICE_CACHE[key] = service
            return service
    except Exception:
        return None


def _cd_get_service(cfg: Dict[str, Any]) -> Optional["CloudDriveService"]:
    if not cfg.get("cloud_drive_enabled"):
        return None
    return _cd_cached_service(cfg)


def _cd_list_children(
    service: "CloudDriveService",
    cd_path: str,
    *,
    cfg: Dict[str, Any] | None = None,
    user: Dict[str, Any] | None = None,
) -> "tuple[list, list]":
    try:
        if cd_path:
            folder = service.registry.get_folder_by_path(cd_path)
        else:
            folder = service.registry.get_root_folder()
        if folder is None:
            return [], []
        folders = service.registry.list_child_folders(folder.id)
        files = service.registry.list_files_in_folder(folder.id)
        if cfg is not None and user is not None:
            folders = [
                item for item in folders
                if _cd_acl_allows(cfg, user, item.path)
            ]
            files = [
                item for item in files
                if _cd_acl_allows(cfg, user, item.path)
            ]
            nodes = [
                *((str(item.path or "").strip().strip("/"), "") for item in folders),
                *((str(item.path or "").strip().strip("/"), str(item.id or "")) for item in files),
            ]
            try:
                access = service.user_access_map(
                    username=str((user or {}).get("username") or ""),
                    role=str((user or {}).get("role") or ""),
                    groups=[str(group_id) for group_id in ((user or {}).get("group_ids") or [])],
                    nodes=nodes,
                    required_level="viewer",
                )
                folders = [
                    item for item in folders
                    if access.get((str(item.path or "").strip().strip("/"), ""), False)
                ]
                files = [
                    item for item in files
                    if access.get(
                        (str(item.path or "").strip().strip("/"), str(item.id or "")),
                        False,
                    )
                ]
            except Exception:
                folders = [
                    item for item in folders
                    if _cd_registry_acl_allows(cfg, user, item.path, service=service)
                ]
                files = [
                    item for item in files
                    if _cd_registry_acl_allows(cfg, user, item.path, file_id=item.id, service=service)
                ]
        return folders, files
    except Exception:
        return [], []


def _cd_breadcrumb_chain(service: "CloudDriveService", cd_path: str) -> "list":
    if not cd_path:
        root = service.registry.get_root_folder()
        return [root] if root else []
    parts: list = []
    segments = cd_path.replace("\\", "/").split("/")
    built = ""
    for seg in segments:
        if not seg:
            continue
        built = (built + "/" + seg).lstrip("/")
        folder = service.registry.get_folder_by_path(built)
        if folder is not None:
            parts.append(folder)
    return parts


def _cd_file_size(size_bytes: int) -> str:
    if size_bytes < 1024:
        return f"{size_bytes} Б"
    if size_bytes < 1024 ** 2:
        return f"{size_bytes / 1024:.1f} КБ"
    if size_bytes < 1024 ** 3:
        return f"{size_bytes / 1024 ** 2:.1f} МБ"
    return f"{size_bytes / 1024 ** 3:.2f} ГБ"


def _cd_search_by_name(
    registry: "Any",
    query: str,
    *,
    max_folders: int = 5,
    max_files: int = 5,
) -> "tuple[list, list]":
    pattern = f"%{query.lower()}%"
    try:
        with registry._connect() as conn:
            folder_rows = conn.execute(
                "SELECT * FROM cloud_folders WHERE lower(name) LIKE ? AND is_root=0 AND deleted_at='' LIMIT ?",
                (pattern, max_folders),
            ).fetchall()
            file_rows = conn.execute(
                "SELECT * FROM cloud_files WHERE lower(name) LIKE ? AND deleted_at='' LIMIT ?",
                (pattern, max_files),
            ).fetchall()
        folders = [registry._folder_from_row(r) for r in folder_rows]
        files = [registry._file_from_row(r) for r in file_rows]
        return folders, files
    except Exception:
        return [], []


def _filter_cd_name_matches(
    cfg: Dict[str, Any],
    user: Dict[str, Any] | None,
    folders: list[Any],
    files: list[Any],
    *,
    service: CloudDriveService,
) -> tuple[list[Any], list[Any]]:
    allowed_folders = [
        folder
        for folder in folders
        if _cd_registry_acl_allows(
            cfg,
            user,
            str(getattr(folder, "path", "") or ""),
            service=service,
        )
    ]
    allowed_files = [
        file_row
        for file_row in files
        if _cd_registry_acl_allows(
            cfg,
            user,
            str(getattr(file_row, "path", "") or ""),
            file_id=str(getattr(file_row, "id", "") or ""),
            service=service,
        )
    ]
    return allowed_folders, allowed_files


def _cd_file_jobs_map(registry: "Any", file_ids: "list[str]") -> "Dict[str, Dict[str, str]]":
    """Single query: returns {file_id: {status, job_type, last_error}} for latest job per file."""
    if not file_ids:
        return {}
    try:
        jobs = registry.list_latest_jobs_for_files(file_ids, job_types=["reindex", "cleanup", "ocr", "preview"])
        return {
            str(file_id): {
                "status": str(job.status or ""),
                "job_type": str(job.job_type or ""),
                "last_error": str(job.last_error or ""),
            }
            for file_id, job in jobs.items()
        }
    except Exception:
        return {}


def _cd_acl_allows(cfg: Dict[str, Any], user: Dict[str, Any] | None, path: str) -> bool:
    """Opt-in Cloud Drive ACL hook.

    Config shape:
      cloud_drive_acl = {
        "users": {"ivan": ["Public", "Projects/A"], "*": ["Public"]},
        "roles": {"admin": ["*"], "user": ["Public"]},
      }
    Empty/missing ACL means allow, so existing installs keep working.
    """
    acl = cfg.get("cloud_drive_acl")
    if not isinstance(acl, dict) or not acl:
        return True
    username = str((user or {}).get("username") or "").strip().lower()
    role = str((user or {}).get("role") or "").strip().lower()
    group_keys = {
        str(group_id or "").strip().lower()
        for group_id in ((user or {}).get("group_ids") or [])
        if str(group_id or "").strip()
    }
    for group in ((user or {}).get("groups") or []):
        if isinstance(group, dict):
            for key in (group.get("id"), group.get("name")):
                if str(key or "").strip():
                    group_keys.add(str(key).strip().lower())
    normalized_path = str(path or "").strip().strip("/")

    def _prefixes(section: str, key: str) -> list[str]:
        block = acl.get(section)
        if not isinstance(block, dict):
            return []
        value = block.get(key)
        if value is None and key != "*":
            value = block.get("*")
        if isinstance(value, str):
            return [value]
        if isinstance(value, list):
            return [str(item) for item in value]
        return []

    allowed = [
        *_prefixes("users", username),
        *_prefixes("roles", role),
        *(prefix for group_key in group_keys for prefix in _prefixes("groups", group_key)),
    ]
    if not allowed:
        return False
    for prefix in allowed:
        clean = str(prefix or "").strip().strip("/")
        if clean in {"", "*"}:
            return True
        if normalized_path == clean or normalized_path.startswith(f"{clean}/"):
            return True
    return False


def _cd_registry_acl_allows(
    cfg: Dict[str, Any],
    user: Dict[str, Any] | None,
    path: str,
    *,
    file_id: str = "",
    required_level: str = "viewer",
    service: CloudDriveService | None = None,
) -> bool:
    if not _cd_acl_allows(cfg, user, path):
        return False
    try:
        svc = service or _cd_cached_service(cfg)
        if svc is None:
            return False
        return bool(
            svc.user_can_access(
                username=str((user or {}).get("username") or ""),
                role=str((user or {}).get("role") or ""),
                groups=[str(group_id) for group_id in ((user or {}).get("group_ids") or [])],
                path=path,
                file_id=file_id,
                required_level=required_level,
            )
        )
    except Exception:
        return False


def _cd_registry_acl_filter(
    cfg: Dict[str, Any],
    user: Dict[str, Any] | None,
    items: List[Dict[str, Any]],
    *,
    service: CloudDriveService | None = None,
    required_level: str = "viewer",
) -> List[Dict[str, Any]]:
    """Filter Cloud Drive rows with one registry ACL snapshot; fail closed."""
    if not items:
        return []
    candidates = [item for item in items if _cd_acl_allows(cfg, user, str(item.get("path") or ""))]
    if not candidates:
        return []
    try:
        svc = service or _cd_cached_service(cfg)
        if svc is None:
            return []
        nodes = [
            (
                str(item.get("path") or ""),
                str(item.get("id") or "") if str(item.get("node_type") or "") == "file" else "",
            )
            for item in candidates
        ]
        decisions = svc.user_access_map(
            username=str((user or {}).get("username") or ""),
            role=str((user or {}).get("role") or ""),
            groups=[str(group_id) for group_id in ((user or {}).get("group_ids") or [])],
            nodes=nodes,
            required_level=required_level,
        )
        return [item for item, node in zip(candidates, nodes) if decisions.get(node, False)]
    except Exception:
        return []


def _cd_registry_node_for_search_item(
    cfg: Dict[str, Any],
    service: CloudDriveService,
    item: Dict[str, Any],
) -> Any:
    cloud_path = str(item.get("cloud_path") or "").strip()
    if cloud_path:
        node = service.registry.get_node_by_path(cloud_path)
        if node is not None:
            return node
    cloud_file_id = str(item.get("cloud_file_id") or "").strip()
    if cloud_file_id:
        node = service.registry.get_file_by_id(cloud_file_id)
        if node is not None:
            return node

    raw_path = str(
        item.get("source_path")
        or item.get("full_path")
        or item.get("filepath")
        or item.get("path")
        or ""
    ).strip()
    if not raw_path:
        return None

    node = service.registry.get_node_by_source_path(raw_path)
    if node is not None:
        return node

    catalog = str(cfg.get("catalog_path") or "").strip()
    if catalog:
        try:
            rel = Path(raw_path).resolve().relative_to(Path(catalog).resolve())
            node = service.registry.get_node_by_path(str(rel).replace("\\", "/"))
            if node is not None:
                return node
        except Exception:
            pass
    return None


def _filter_cloud_drive_search_results(
    cfg: Dict[str, Any],
    user: Dict[str, Any] | None,
    results: List[Dict[str, Any]],
    *,
    service: CloudDriveService | None = None,
) -> List[Dict[str, Any]]:
    if not results:
        return results
    effective_user = dict(user or {})
    username = str(effective_user.get("username") or "").strip().lower()
    if username:
        try:
            fresh_user = UserAuthDB(str(_users_db_path(cfg))).get_user(username=username)
            if fresh_user is not None:
                effective_user = fresh_user
        except Exception:
            effective_user["group_ids"] = []
            effective_user["groups"] = []
    try:
        effective_service = service or _cd_cached_service(cfg)
    except Exception:
        effective_service = None
    if effective_service is None:
        return [] if str(cfg.get("cloud_drive_db_path") or "").strip() else results

    filtered: List[Dict[str, Any]] = []
    is_admin = str(effective_user.get("role") or "").strip().lower() == "admin"
    resolved: List[tuple[Dict[str, Any], str, str]] = []
    for item in results:
        node = _cd_registry_node_for_search_item(cfg, effective_service, item)
        if node is None:
            if is_admin:
                filtered.append(item)
            continue

        node_path = str(getattr(node, "path", "") or "")
        cloud_file_id = str(getattr(node, "id", "") or "") if hasattr(node, "folder_id") else ""
        resolved.append((item, node_path, cloud_file_id))

    nodes = [(path.strip().strip("/"), file_id) for _item, path, file_id in resolved]
    try:
        decisions = effective_service.user_access_map(
            username=str(effective_user.get("username") or ""),
            role=str(effective_user.get("role") or ""),
            groups=[str(group_id) for group_id in (effective_user.get("group_ids") or [])],
            nodes=nodes,
            required_level="viewer",
        )
    except Exception:
        decisions = {}

    for item, node_path, cloud_file_id in resolved:
        clean_path = node_path.strip().strip("/")
        allowed = _cd_acl_allows(cfg, effective_user, clean_path) and decisions.get(
            (clean_path, cloud_file_id),
            False,
        )
        if allowed:
            item.setdefault("cloud_path", node_path)
            if cloud_file_id:
                item.setdefault("cloud_file_id", cloud_file_id)
            filtered.append(item)
    return filtered


def _run_authorized_quick_name_search(
    searcher: RAGSearcher,
    *,
    cfg: Dict[str, Any],
    user: Dict[str, Any] | None,
    query: str,
    limit: int,
    file_type: Optional[str],
) -> List[Dict[str, Any]]:
    """Run quick retrieval and ACL filtering in the same worker thread."""
    results = _run_quick_name_search(
        searcher,
        query=query,
        limit=limit,
        file_type=file_type,
    )
    return _filter_cloud_drive_search_results(cfg, user, results)


def _run_authorized_catalog_search(
    searcher: RAGSearcher,
    *,
    cfg: Dict[str, Any],
    user: Dict[str, Any] | None,
    query: str,
    query_original: str,
    query_used: str,
    limit: int,
    file_type: Optional[str],
    content_only: bool,
    title_only: bool,
    username: str = "",
) -> List[Dict[str, Any]]:
    """Run full retrieval and ACL filtering without blocking NiceGUI."""
    results = _run_catalog_search(
        searcher,
        query=query,
        query_original=query_original,
        query_used=query_used,
        limit=limit,
        file_type=file_type,
        content_only=content_only,
        title_only=title_only,
        username=username,
    )
    return _filter_cloud_drive_search_results(cfg, user, results)


# ─────────────────────────── search helpers ─────────────────────────────────

_STOP_WORDS = frozenset(["и", "в", "на", "с", "по", "для", "из", "к", "от", "за", "о", "а", "но", "не"])


def _highlight_query_terms(text: str, query: str) -> str:
    """Return HTML-escaped text with query terms wrapped in <mark> tags."""
    if not text or not query:
        return html.escape(text or "")
    terms = [t for t in re.split(r"\s+", query.strip()) if len(t) >= 2 and t.lower() not in _STOP_WORDS]
    if not terms:
        return html.escape(text)
    escaped = html.escape(text)
    for term in terms:
        pattern = re.compile(re.escape(html.escape(term)), re.IGNORECASE)
        escaped = pattern.sub(lambda m: f'<mark class="rag-highlight">{m.group(0)}</mark>', escaped)
    return escaped


# ─────────────────────────── index stats / telemetry ────────────────────────

def _read_index_stats(cfg: Dict[str, Any]) -> Dict[str, Any]:
    state_file = Path(str(cfg.get("qdrant_db_path") or "")) / "index_state.db"
    out: Dict[str, Any] = {
        "found": False,
        "state_file": str(state_file),
        "total": 0,
        "total_size_bytes": 0,
        "by_ext": {},
        "by_ext_size": {},
        "by_stage": {},
    }
    if not state_file.exists():
        return out
    try:
        db = IndexStateDB(str(state_file))
        stats = db.stats()
    except Exception as exc:
        out["error"] = str(exc)
        return out
    out.update({
        "found": True,
        "total": int(stats.get("total") or 0),
        "total_size_bytes": int(stats.get("total_size_bytes") or 0),
        "by_ext": dict(stats.get("by_ext") or {}),
        "by_ext_size": dict(stats.get("by_ext_size") or {}),
        "by_stage": dict(stats.get("by_stage") or {}),
    })
    try:
        out["last_modified"] = time.strftime("%d.%m.%Y %H:%M", time.localtime(state_file.stat().st_mtime))
    except Exception:
        pass
    return out


def _index_stage_from_note(note: str) -> str:
    match = re.search(r"(?:^|\b)stage=([a-z_]+)", str(note or "").lower())
    stage = match.group(1) if match else ""
    return stage if stage in {"metadata", "small", "large", "content", "all", "full"} else ""


def _find_headless_active_stages(db_path: Path) -> List[Dict[str, Any]]:
    """Return stage-progress rows for a live headless indexer (PID alive but DB status='cancelled')."""
    import json as _json

    live_pid = 0
    marker_stage = ""
    marker_path = PROJECT_ROOT / "runtime" / "index_active.json"
    if marker_path.exists():
        try:
            data = _json.loads(marker_path.read_text(encoding="utf-8"))
            pid = int(data.get("pid") or 0)
            if pid > 0 and _process_matches_module(pid, "rag_catalog.core.index_rag"):
                live_pid = pid
                marker_stage = str(data.get("stage") or "").strip().lower()
        except Exception:
            pass

    if not live_pid:
        pids = _find_module_process_pids("rag_catalog.core.index_rag")
        if pids:
            live_pid = pids[0]

    if not live_pid:
        return []

    run_rows = _db_query_dicts(
        db_path,
        """
        SELECT run_id,
               note,
               CAST((julianday(CURRENT_TIMESTAMP) - julianday(ts_started)) * 86400 AS INTEGER) AS duration_sec
        FROM index_runs
        WHERE worker_pid=?
        ORDER BY ts_started DESC
        LIMIT 1
        """,
        (live_pid,),
    )
    if not run_rows:
        return [
            {
                "run_id": "",
                "stage": marker_stage or "all",
                "status": "running",
                "run_status": "running",
                "run_note": "runtime_marker",
                "processed_files": 0,
                "total_files": 0,
                "added_files": 0,
                "updated_files": 0,
                "skipped_files": 0,
                "error_files": 0,
                "points_added": 0,
                "duration_sec": 0,
                "_progress_unknown": True,
            }
        ]
    run_id = str(run_rows[0].get("run_id") or "")
    if not run_id:
        return []
    run_note = str(run_rows[0].get("note") or "").strip()

    rows = _db_query_dicts(
        db_path,
        """
        SELECT isp.*,
               'running' AS run_status,
               'headless' AS run_note,
               CAST((julianday(COALESCE(isp.ts_finished, CURRENT_TIMESTAMP)) - julianday(isp.ts_started)) * 86400 AS INTEGER) AS duration_sec
        FROM index_stage_progress AS isp
        WHERE isp.run_id=?
        ORDER BY CASE isp.stage WHEN 'metadata' THEN 1 WHEN 'small' THEN 2 WHEN 'large' THEN 3 WHEN 'content' THEN 4 ELSE 9 END
        """,
        (run_id,),
    )
    if not rows:
        return [
            {
                "run_id": run_id,
                "stage": marker_stage or _index_stage_from_note(run_note) or "all",
                "status": "running",
                "run_status": "running",
                "run_note": run_note or "headless",
                "processed_files": 0,
                "total_files": 0,
                "added_files": 0,
                "updated_files": 0,
                "skipped_files": 0,
                "error_files": 0,
                "points_added": 0,
                "duration_sec": int(run_rows[0].get("duration_sec") or 0),
                "_progress_unknown": True,
            }
        ]
    # PID is alive → watchdog-set 'cancelled' status is wrong; correct it to 'running'
    for r in rows:
        if r.get("status") == "cancelled":
            r["status"] = "running"
    return rows


def _find_headless_active_ocr(db_path: Path) -> Optional[Dict[str, Any]]:
    live_pid = 0
    marker_path = PROJECT_ROOT / "runtime" / "ocr_active.json"
    if marker_path.exists():
        try:
            data = json.loads(marker_path.read_text(encoding="utf-8"))
            pid = int(data.get("pid") or 0)
            if pid > 0 and _process_matches_module(pid, "rag_catalog.core.ocr_pdfs"):
                live_pid = pid
        except Exception:
            pass

    if not live_pid:
        pids = _find_module_process_pids("rag_catalog.core.ocr_pdfs")
        if pids:
            live_pid = pids[0]

    if not live_pid:
        return None

    rows = _db_query_dicts(
        db_path,
        """
        SELECT *,
               CAST((julianday(CURRENT_TIMESTAMP) - julianday(ts_started)) * 86400 AS INTEGER) AS duration_sec
        FROM ocr_runs
        WHERE worker_pid=?
        ORDER BY ts_started DESC LIMIT 1
        """,
        (live_pid,),
    )
    if rows:
        row = dict(rows[0])
        if row.get("status") == "cancelled":
            row["status"] = "running"
        return row
    return {
        "ocr_run_id": "",
        "status": "running",
        "worker_pid": live_pid,
        "found_scanned": 0,
        "processed_pdfs": 0,
        "duration_sec": 0,
        "note": "process_scan",
        "_progress_unknown": True,
    }


def _read_ocr_inventory(cfg: Dict[str, Any]) -> Dict[str, Any]:
    """Return OCR inventory across index_state and cached OCR results."""
    qdrant_dir = Path(str(cfg.get("qdrant_db_path") or "")).expanduser()
    state_file = qdrant_dir / "index_state.db"
    telemetry_file = _telemetry_db_path(cfg)
    cache_key = f"{state_file}|{telemetry_file}"
    now = time.monotonic()
    cached = _OCR_INVENTORY_CACHE.get(cache_key)
    if cached and now - float(cached.get("ts") or 0) < OCR_INVENTORY_CACHE_TTL_SECONDS:
        data = dict(cached.get("data") or {})
        data["status_counts"] = dict(data.get("status_counts") or {})
        data["engine_counts"] = dict(data.get("engine_counts") or {})
        data["cached"] = True
        return data

    out: Dict[str, Any] = {
        "found": False,
        "ocr_capable_total": 0,
        "eligible_total": 0,
        "pending_candidates": 0,
        "recognized_files": 0,
        "partial_files": 0,
        "empty_files": 0,
        "error_files": 0,
        "recognized_pages": 0,
        "recognized_chars": 0,
        "recognized_lines": 0,
        "recognized_duration_ms": 0,
        "pages_per_minute": 0.0,
        "seconds_per_page": 0.0,
        "fallback_files": 0,
        "engine_counts": {},
        "status_counts": {},
    }
    state_rows: Dict[str, Dict[str, Any]] = {}
    candidate_paths: set[str] = set()
    small_pdf_mb_raw = cfg.get("small_pdf_mb")
    min_pdf_size = max(0, int(float(2.0 if small_pdf_mb_raw in (None, "") else small_pdf_mb_raw) * 1_048_576))
    if state_file.exists():
        try:
            with sqlite3.connect(str(state_file), timeout=30.0) as conn:
                conn.row_factory = sqlite3.Row
                rows = conn.execute(
                    """
                    SELECT full_path, extension, size_bytes, mtime, stage, status, indexed_stage
                    FROM state_entries
                    WHERE lower(extension) IN (?,?,?,?,?,?,?,?,?)
                    """,
                    tuple(sorted(OCR_CAPABLE_EXTENSIONS)),
                ).fetchall()
            for row in rows:
                path = str(row["full_path"] or "").strip()
                if not path:
                    continue
                ext = str(row["extension"] or "").lower()
                stage = str(row["stage"] or "")
                status = str(row["status"] or "")
                indexed_stage = str(row["indexed_stage"] or "")
                state_rows[path] = {
                    "extension": ext,
                    "size_bytes": int(row["size_bytes"] or 0),
                    "mtime": float(row["mtime"] or 0.0),
                    "stage": stage,
                    "status": status,
                    "indexed_stage": indexed_stage,
                }
                needs_content = stage != "content" or status in {"empty", "error", "deferred_ocr"} or indexed_stage in {"", "metadata", "small"}
                if ext == ".pdf":
                    if status == "deferred_ocr" or (int(row["size_bytes"] or 0) >= min_pdf_size and needs_content):
                        candidate_paths.add(path)
                elif ext in OCR_IMAGE_EXTENSIONS and needs_content:
                    candidate_paths.add(path)
        except Exception as exc:
            out["state_error"] = str(exc)

    ocr_rows: List[Dict[str, Any]] = []
    try:
        result_cols = _db_query_dicts(telemetry_file, "PRAGMA table_info(ocr_file_results)")
        result_col_names = {str(row.get("name") or "") for row in result_cols}
        has_line_count = "line_count" in result_col_names
        line_expr = "line_count" if has_line_count else "0 AS line_count"
        duration_expr = "duration_ms" if "duration_ms" in result_col_names else "0 AS duration_ms"
        engine_expr = "engine" if "engine" in result_col_names else "'' AS engine"
        fallback_expr = "fallback_used" if "fallback_used" in result_col_names else "0 AS fallback_used"
        ocr_rows = _db_query_dicts(
            telemetry_file,
            f"""
            SELECT file_path, file_mtime, status, pages, char_count, {line_expr},
                   {duration_expr}, {engine_expr}, {fallback_expr}
            FROM ocr_file_results
            """,
        )
        out["found"] = True
    except Exception as exc:
        out["telemetry_error"] = str(exc)

    result_paths: set[str] = set()
    recognized_paths: set[str] = set()
    empty_paths: set[str] = set()
    error_paths: set[str] = set()
    partial_paths: set[str] = set()
    status_counts: Dict[str, int] = {}
    recognized_pages = 0
    recognized_chars = 0
    recognized_lines = 0
    recognized_duration_ms = 0
    fallback_files = 0
    engine_counts: Dict[str, int] = {}

    for row in ocr_rows:
        path = str(row.get("file_path") or "").strip()
        if not path:
            continue
        state = state_rows.get(path)
        if state is None or abs(float(row.get("file_mtime") or 0.0) - float(state.get("mtime") or 0.0)) >= 1.0:
            continue
        result_paths.add(path)
        status = str(row.get("status") or "unknown").strip().lower() or "unknown"
        status_counts[status] = status_counts.get(status, 0) + 1
        engine = str(row.get("engine") or "").strip().lower()
        if engine:
            engine_counts[engine] = engine_counts.get(engine, 0) + 1
        if bool(row.get("fallback_used")):
            fallback_files += 1
        chars = int(row.get("char_count") or 0)
        if status == "error":
            error_paths.add(path)
            continue
        if status == "empty" or chars <= 0:
            empty_paths.add(path)
            continue
        recognized_paths.add(path)
        recognized_pages += int(row.get("pages") or 0)
        recognized_chars += chars
        recognized_lines += int(row.get("line_count") or 0)
        recognized_duration_ms += max(0, int(row.get("duration_ms") or 0))
        if (
            str(state.get("stage") or "") != "content"
            or str(state.get("indexed_stage") or "") not in {"content", "large"}
            or str(state.get("status") or "") in {"empty", "error", "deferred_ocr"}
        ):
            partial_paths.add(path)

    relevant_paths = candidate_paths | result_paths
    pending_paths = candidate_paths - recognized_paths - empty_paths - error_paths
    pages_per_minute = (
        recognized_pages * 60_000.0 / recognized_duration_ms
        if recognized_pages > 0 and recognized_duration_ms > 0
        else 0.0
    )
    out.update(
        {
            "found": bool(state_rows or ocr_rows or out.get("found")),
            "ocr_capable_total": len(state_rows),
            "eligible_total": len(relevant_paths),
            "pending_candidates": len(pending_paths),
            "recognized_files": len(recognized_paths),
            "partial_files": len(partial_paths),
            "empty_files": len(empty_paths),
            "error_files": len(error_paths),
            "recognized_pages": recognized_pages,
            "recognized_chars": recognized_chars,
            "recognized_lines": recognized_lines,
            "recognized_duration_ms": recognized_duration_ms,
            "pages_per_minute": pages_per_minute,
            "seconds_per_page": 60.0 / pages_per_minute if pages_per_minute > 0 else 0.0,
            "fallback_files": fallback_files,
            "engine_counts": engine_counts,
            "status_counts": status_counts,
        }
    )
    cached_data = dict(out)
    cached_data["status_counts"] = dict(status_counts)
    cached_data["engine_counts"] = dict(engine_counts)
    _OCR_INVENTORY_CACHE[cache_key] = {"ts": now, "data": cached_data}
    return out


def _read_index_activity(cfg: Dict[str, Any]) -> Dict[str, Any]:
    """Return index/OCR activity without scanning the full OCR file inventory."""
    db_path = _telemetry_db_path(cfg)
    active_runs = _db_query_dicts(
        db_path,
        """
        SELECT *,
               CAST((julianday(CURRENT_TIMESTAMP) - julianday(ts_started)) * 86400 AS INTEGER) AS duration_sec
        FROM index_runs
        WHERE status='running'
        ORDER BY ts_started DESC
        LIMIT 3
        """,
    )
    active_run_ids = [str(row.get("run_id") or "") for row in active_runs if str(row.get("run_id") or "")]
    active_stages: List[Dict[str, Any]] = []
    if active_run_ids:
        placeholders = ",".join("?" for _ in active_run_ids)
        active_stages = _db_query_dicts(
            db_path,
            f"""
            SELECT isp.*,
                   ir.status AS run_status,
                   ir.note AS run_note,
                   CAST((julianday(COALESCE(isp.ts_finished, CURRENT_TIMESTAMP)) - julianday(isp.ts_started)) * 86400 AS INTEGER) AS duration_sec
            FROM index_stage_progress AS isp
            LEFT JOIN index_runs AS ir ON ir.run_id = isp.run_id
            WHERE isp.run_id IN ({placeholders})
              AND isp.status='running'
            ORDER BY isp.ts_started, isp.stage
            """,
            tuple(active_run_ids),
        )
    if not active_stages:
        active_stages = _find_headless_active_stages(db_path)

    latest_stages = _db_query_dicts(
        db_path,
        """
        SELECT isp.*,
               ir.status AS run_status,
               ir.note AS run_note,
               CAST((julianday(COALESCE(isp.ts_finished, CURRENT_TIMESTAMP)) - julianday(isp.ts_started)) * 86400 AS INTEGER) AS duration_sec
        FROM index_stage_progress AS isp
        LEFT JOIN index_runs AS ir ON ir.run_id = isp.run_id
        WHERE isp.run_id=(SELECT run_id FROM index_runs ORDER BY ts_started DESC LIMIT 1)
        ORDER BY CASE isp.stage WHEN 'metadata' THEN 1 WHEN 'small' THEN 2 WHEN 'large' THEN 3 WHEN 'content' THEN 4 ELSE 9 END
        """,
    )
    stage_summary = _db_query_dicts(
        db_path,
        """
        WITH finished AS (
            SELECT isp.run_id,
                   isp.stage,
                   isp.status,
                   isp.total_files,
                   isp.processed_files,
                   isp.added_files,
                   isp.updated_files,
                   isp.skipped_files,
                   isp.error_files,
                   isp.points_added,
                   isp.ts_started,
                   isp.ts_finished,
                   CASE
                       WHEN isp.status='cancelled'
                            AND (
                                COALESCE(ir.note, '') LIKE '%recovery%'
                                OR COALESCE(ir.note, '') LIKE '%switch_to_ocr%'
                                OR COALESCE(ir.note, '') LIKE '%active_ocr_running%'
                            )
                       THEN 1
                       ELSE 0
                   END AS is_recovery_cancelled,
                   CAST((julianday(isp.ts_finished) - julianday(isp.ts_started)) * 86400 AS INTEGER) AS duration_sec
            FROM index_stage_progress AS isp
            LEFT JOIN index_runs AS ir ON ir.run_id=isp.run_id
            WHERE isp.ts_finished IS NOT NULL
        ),
        latest AS (
            SELECT *
            FROM (
                SELECT f.*,
                       ROW_NUMBER() OVER (
                           PARTITION BY f.stage
                           ORDER BY f.is_recovery_cancelled ASC, f.ts_started DESC
                       ) AS rn
                FROM finished f
            )
            WHERE rn=1
        ),
        avg_by_stage AS (
            SELECT stage,
                   COUNT(*) AS runs_count,
                   AVG(duration_sec) AS avg_duration_sec,
                   AVG(processed_files) AS avg_processed_files
            FROM finished
            GROUP BY stage
        )
        SELECT latest.stage,
               latest.run_id,
               latest.status,
               latest.total_files,
               latest.processed_files,
               latest.added_files,
               latest.updated_files,
               latest.skipped_files,
               latest.error_files,
               latest.points_added,
               latest.ts_started,
               latest.ts_finished,
               ir.status AS run_status,
               ir.note AS run_note,
               latest.duration_sec AS last_duration_sec,
               avg_by_stage.runs_count,
               CAST(avg_by_stage.avg_duration_sec AS INTEGER) AS avg_duration_sec,
               CAST(avg_by_stage.avg_processed_files AS INTEGER) AS avg_processed_files
        FROM latest
        JOIN avg_by_stage ON avg_by_stage.stage=latest.stage
        LEFT JOIN index_runs AS ir ON ir.run_id=latest.run_id
        ORDER BY CASE latest.stage WHEN 'metadata' THEN 1 WHEN 'small' THEN 2 WHEN 'large' THEN 3 WHEN 'content' THEN 4 ELSE 9 END
        """,
    )
    overall = _db_query_dicts(
        db_path,
        """
        SELECT COUNT(*) AS runs_count,
               CAST(AVG((julianday(ts_finished) - julianday(ts_started)) * 86400) AS INTEGER) AS avg_duration_sec,
               MAX(total_files) AS max_total_files,
               AVG(total_files) AS avg_total_files
        FROM index_runs
        WHERE ts_finished IS NOT NULL
        """,
    )
    active_ocr = _db_query_dicts(
        db_path,
        """
        SELECT *,
               CAST((julianday(CURRENT_TIMESTAMP) - julianday(ts_started)) * 86400 AS INTEGER) AS duration_sec
        FROM ocr_runs
        WHERE status='running'
        ORDER BY ts_started DESC
        LIMIT 1
        """,
    )
    active_ocr_row = active_ocr[0] if active_ocr else _find_headless_active_ocr(db_path)
    return {
        "active_runs": active_runs,
        "active_stages": active_stages,
        "latest_stages": latest_stages,
        "stage_summary": stage_summary,
        "overall": overall[0] if overall else {},
        "active_ocr": active_ocr_row,
    }


def _read_index_telemetry(cfg: Dict[str, Any]) -> Dict[str, Any]:
    db_path = _telemetry_db_path(cfg)
    activity = _read_index_activity(cfg)
    last_run = _db_query_dicts(
        db_path,
        """
        SELECT *,
               CAST((julianday(COALESCE(ts_finished, CURRENT_TIMESTAMP)) - julianday(ts_started)) * 86400 AS INTEGER) AS duration_sec
        FROM index_runs
        ORDER BY ts_started DESC
        LIMIT 1
        """,
    )
    last_ocr = _db_query_dicts(
        db_path,
        """
        SELECT *,
               CAST((julianday(COALESCE(ts_finished, CURRENT_TIMESTAMP)) - julianday(ts_started)) * 86400 AS INTEGER) AS duration_sec
        FROM ocr_runs
        ORDER BY ts_started DESC
        LIMIT 1
        """,
    )
    ocr_summary = _db_query_dicts(
        db_path,
        """
        SELECT COUNT(*) AS runs_count,
               CAST(AVG((julianday(ts_finished) - julianday(ts_started)) * 86400) AS INTEGER) AS avg_duration_sec,
               AVG(found_scanned) AS avg_found_scanned,
               AVG(processed_pdfs) AS avg_processed_pdfs
        FROM ocr_runs
        WHERE ts_finished IS NOT NULL
        """,
    )
    ocr_inventory = _read_ocr_inventory(cfg)
    return {
        "last_run": last_run[0] if last_run else None,
        "active_runs": activity["active_runs"],
        "active_stages": activity["active_stages"],
        "latest_stages": activity["latest_stages"],
        "stage_summary": activity["stage_summary"],
        "overall": activity["overall"],
        "active_ocr": activity["active_ocr"],
        "last_ocr": last_ocr[0] if last_ocr else None,
        "ocr_summary": ocr_summary[0] if ocr_summary else {},
        "ocr_inventory": ocr_inventory,
    }


# ─────────────────────────── searcher / admin helpers ───────────────────────

_SEARCHER_CACHE: Dict[tuple[str, ...], RAGSearcher] = {}


def _searcher_cache_key(cfg: Dict[str, Any]) -> tuple[str, ...]:
    return (
        str(cfg.get("qdrant_url") or ""),
        str(cfg.get("qdrant_db_path") or ""),
        str(cfg.get("collection_name") or ""),
        str(cfg.get("embedding_model") or ""),
        str(cfg.get("embedding_collection_versioning") or ""),
        str(cfg.get("embedding_collection_suffix") or ""),
        str(cfg.get("retrieval_preset") or ""),
        str(cfg.get("retrieval_pipeline") or ""),
    )


def _qdrant_http_ready(cfg: Dict[str, Any], *, timeout: float = 0.75) -> bool:
    qdrant_url = str(cfg.get("qdrant_url") or "").strip().rstrip("/")
    if not qdrant_url:
        return True
    try:
        req = urllib.request.Request(f"{qdrant_url}/collections", method="GET")
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            return 200 <= int(resp.status) < 500
    except (OSError, urllib.error.URLError, urllib.error.HTTPError):
        return False


def _warm_searcher_cache(cfg: Dict[str, Any]) -> None:
    """Preload shared searcher/embedder after startup so first user search is not cold."""
    try:
        if not _qdrant_http_ready(cfg):
            return
        key = _searcher_cache_key(cfg)
        searcher = _SEARCHER_CACHE.get(key)
        if searcher is None:
            searcher = RAGSearcher(cfg)
            _SEARCHER_CACHE[key] = searcher
        if searcher.connected:
            searcher.warm_retrieval_cache()
            searcher.embedder.encode(["warmup"])
    except Exception:
        # Warmup is an optimization; search path will report real errors to the user.
        pass


def _cached_searcher_if_ready(cfg: Dict[str, Any]) -> Optional[RAGSearcher]:
    if not _qdrant_http_ready(cfg):
        return None
    searcher = _SEARCHER_CACHE.get(_searcher_cache_key(cfg))
    if searcher is not None and searcher.connected:
        return searcher
    return None


def _ensure_searcher(state: PageState) -> Optional[RAGSearcher]:
    if state.searcher is not None:
        return state.searcher
    try:
        key = _searcher_cache_key(state.cfg)
        cached = _SEARCHER_CACHE.get(key)
        if cached is None or not cached.connected:
            cached = RAGSearcher(state.cfg)
            _SEARCHER_CACHE[key] = cached
        state.searcher = cached
    except Exception as exc:
        state.searcher_error = str(exc)
        return None
    if not state.searcher.connected:
        state.searcher_error = "Нет подключения к Qdrant."
    return state.searcher


def _is_admin(state: PageState) -> bool:
    return str((state.current_user or {}).get("role") or "") == "admin"


def _show_system_files(state: PageState) -> bool:
    if not _is_admin(state):
        return False
    try:
        auth = state.auth_db if state.auth_db is not None else _get_auth_db(state)
        return bool(auth.get_show_system_files_for_admin())
    except Exception:
        return False


def _is_system_file(path_or_ext: str | Path) -> bool:
    value = str(path_or_ext or "")
    ext = Path(value).suffix.lower() or value.lower()
    name = Path(value).name.lower()
    if not ext:
        return False
    if ext in SYSTEM_FILE_EXTENSIONS:
        return True
    if ext.endswith("_"):
        return True
    if ext in {".sample", ".backup", ".sav", ".asd"}:
        return True
    return name.startswith("~$") or name.endswith(".tmp")


# ─────────────────────────── result grouping / icons ────────────────────────

def _result_kind(result: Dict[str, Any]) -> str:
    if result.get("type") == "folder_metadata":
        return "Каталог"
    ext = str(result.get("extension") or "").lower()
    if ext in {".xlsx", ".xls", ".csv"}:
        return "Таблица"
    if ext == ".pdf":
        return "PDF"
    if ext in {".docx", ".doc"}:
        return "Документ"
    return "Файл"


def _result_group(result: Dict[str, Any]) -> str:
    text = " ".join(
        str(result.get(key, "") or "").lower()
        for key in ("filename", "path", "type", "text", "extension")
    )
    if str(result.get("type") or "") == "folder_metadata":
        return "Каталоги"
    if any(word in text for word in ("птс", "псм", "стс", "техпаспорт", "техническ", "паспорт транспорт", "экскаватор")):
        return "Техпаспорта ТС"
    if any(word in text for word in ("паспорт", "удостоверен")):
        return "Паспорта и удостоверения"
    if any(word in text for word in ("договор", "соглашен", "контракт")):
        return "Договоры"
    if any(word in text for word in ("счет", "счёт", "оплат", "платеж")):
        return "Счета и платежи"
    if str(result.get("extension") or "").lower() in {".xlsx", ".xls", ".csv"}:
        return "Таблицы"
    if str(result.get("extension") or "").lower() == ".pdf":
        return "PDF"
    return "Другие файлы"


def _grouped_results(results: List[Dict[str, Any]]) -> List[tuple[str, List[Dict[str, Any]]]]:
    order = [
        "Каталоги",
        "Техпаспорта ТС",
        "Паспорта и удостоверения",
        "Договоры",
        "Счета и платежи",
        "Таблицы",
        "PDF",
        "Другие файлы",
    ]
    grouped: Dict[str, List[Dict[str, Any]]] = {}
    for result in results:
        grouped.setdefault(_result_group(result), []).append(result)
    return [
        (group, sorted(grouped[group], key=lambda item: float(item.get("score") or 0), reverse=True))
        for group in order
        if group in grouped
    ]


def _file_icon_svg(path_or_ext: str, kind: str = "Файл") -> str:
    ext = Path(str(path_or_ext or "")).suffix.lower() or str(path_or_ext or "").lower()
    if kind == "Каталог":
        color, label = "#f2b84b", ""
    elif ext in {".doc", ".docx"}:
        color, label = "#2b579a", "W"
    elif ext in {".xls", ".xlsx", ".csv"}:
        color, label = "#217346", "X"
    elif ext in {".ppt", ".pptx"}:
        color, label = "#d24726", "P"
    elif ext == ".pdf":
        color, label = "#d32f2f", "PDF"
    elif ext in {".zip", ".rar", ".7z"}:
        color, label = "#d39a18", "ZIP"
    elif ext in {".jpg", ".jpeg", ".png", ".gif", ".webp"}:
        color, label = "#7b4ab8", "IMG"
    elif ext in {".txt", ".log", ".md"}:
        color, label = "#64748b", "TXT"
    elif ext in {".mp3", ".wav", ".flac"}:
        color, label = "#0e7490", "AUD"
    elif ext in {".mp4", ".avi", ".mov", ".mkv"}:
        color, label = "#7c3aed", "VID"
    elif ext in {".json", ".xml", ".html", ".htm", ".css", ".js", ".py", ".ps1", ".cmd", ".bat"}:
        color, label = "#475569", "DEV"
    elif ext in {".dwg", ".svg", ".psd", ".ico"}:
        color, label = "#0891b2", "ART"
    elif _is_system_file(path_or_ext):
        color, label = "#94a3b8", (ext.replace(".", "").upper()[:3] or "SYS")
    else:
        color, label = "#0f766e", (ext.replace(".", "").upper()[:3] or "FILE")

    if kind == "Каталог":
        svg = f"""
        <svg viewBox="0 0 56 56" aria-hidden="true">
          <path fill="#d99d2b" d="M5 15a6 6 0 0 1 6-6h12l5 6h17a6 6 0 0 1 6 6v4H5z"/>
          <path fill="{color}" d="M5 20h46v21a6 6 0 0 1-6 6H11a6 6 0 0 1-6-6z"/>
          <path fill="#ffd56d" opacity=".75" d="M8 23h40v4H8z"/>
        </svg>
        """
    else:
        font_size = "12" if len(label) <= 1 else "9"
        svg = f"""
        <svg viewBox="0 0 56 56" aria-hidden="true">
          <path fill="#ffffff" d="M12 4h22l10 10v36a4 4 0 0 1-4 4H12a4 4 0 0 1-4-4V8a4 4 0 0 1 4-4z"/>
          <path fill="#dbe4ef" d="M34 4v10h10z"/>
          <path fill="{color}" d="M6 29h37a4 4 0 0 1 4 4v12a4 4 0 0 1-4 4H6z"/>
          <text x="25" y="42" fill="#fff" text-anchor="middle" font-family="Arial, sans-serif" font-size="{font_size}" font-weight="700">{html.escape(label)}</text>
          <path fill="none" stroke="#cbd5e1" d="M12.5 4.5h21.3L43.5 14v35.5a4 4 0 0 1-4 4h-27a4 4 0 0 1-4-4v-41a4 4 0 0 1 4-4z"/>
        </svg>
        """
    icon_class = "rag-file-icon system" if kind != "Каталог" and _is_system_file(path_or_ext) else "rag-file-icon"
    return f'<span class="{icon_class}">{svg}</span>'


def _file_badge_html(path_or_ext: str, kind: str = "Файл") -> str:
    """CSS badge <span> for ui.html() in list/table rows — uses rag-file-badge CSS."""
    ext = Path(str(path_or_ext or "")).suffix.lower() or str(path_or_ext or "").lower()
    if kind == "Каталог":
        return (
            '<span class="rag-file-badge fld" aria-label="Папка">'
            '<span class="material-icons rag-folder-badge-icon" aria-hidden="true">folder</span>'
            "</span>"
        )
    elif ext in {".doc", ".docx"}:
        css_cls, label = "doc", "DOC"
    elif ext in {".xls", ".xlsx", ".csv"}:
        css_cls, label = "xls", "XLS"
    elif ext in {".ppt", ".pptx"}:
        css_cls, label = "ppt", "PPT"
    elif ext == ".pdf":
        css_cls, label = "pdf", "PDF"
    elif ext in {".zip", ".rar", ".7z"}:
        css_cls, label = "txt", "ZIP"
    elif ext in {".jpg", ".jpeg", ".png", ".gif", ".webp", ".svg", ".ico"}:
        css_cls, label = "img", "IMG"
    elif ext in {".txt", ".log", ".md"}:
        css_cls, label = "txt", "TXT"
    elif ext in {".mp3", ".wav", ".flac"}:
        css_cls, label = "doc", "AUD"
    elif ext in {".mp4", ".avi", ".mov", ".mkv"}:
        css_cls, label = "ppt", "VID"
    else:
        css_cls = "txt"
        label = ext.replace(".", "").upper()[:3] or "FIL"
    return f'<span class="rag-file-badge {css_cls}">{html.escape(label)}</span>'


# ─────────────────────────── explorer helpers ────────────────────────────────

def _path_sort_key(path: Path, sort_by: str) -> Any:
    if sort_by == "По размеру":
        name = path.name.lower()
        if path.is_file():
            try:
                return (path.stat().st_size, name)
            except Exception:
                return (0, name)
        return (-1, name)
    if sort_by == "По дате":
        name = path.name.lower()
        try:
            return (path.stat().st_mtime, name)
        except Exception:
            return (0, name)
    return path.name.lower()


def _file_rows(path: Path, state: PageState) -> tuple[List[Path], List[Path], int]:
    try:
        entries = [x for x in path.iterdir() if not x.name.startswith(".") and not x.name.startswith("~$")]
    except Exception:
        return [], [], 0

    dirs = [x for x in entries if x.is_dir()]
    files = [x for x in entries if x.is_file()]
    if not _show_system_files(state):
        files = [x for x in files if not _is_system_file(x)]
    needle = state.explorer_filter.strip().lower()
    if needle:
        dirs = [x for x in dirs if needle in x.name.lower()]
        files = [x for x in files if needle in x.name.lower()]
    if state.explorer_ext != "Все":
        files = [x for x in files if x.suffix.lower() == state.explorer_ext.lower()]

    reverse = bool(state.explorer_desc)
    dirs.sort(key=lambda x: _path_sort_key(x, state.explorer_sort), reverse=reverse)
    files.sort(key=lambda x: _path_sort_key(x, state.explorer_sort), reverse=reverse)
    return dirs, files, len(files)


def _event_input_value(event: Any, fallback: Any = "") -> str:
    value = getattr(event, "value", None)
    if value is not None:
        return str(value)
    args = getattr(event, "args", None)
    if isinstance(args, dict):
        for key in ("value", "modelValue"):
            if args.get(key) is not None:
                return str(args[key])
        target = args.get("target")
        if isinstance(target, dict) and target.get("value") is not None:
            return str(target["value"])
    return str(fallback or "")


def _apply_explorer_filter_input(state: PageState, event: Any, fallback: Any = "") -> None:
    state.explorer_filter = _event_input_value(event, fallback)
    state.explorer_page = 0


# ─────────────────────────── user state persistence ─────────────────────────

def _load_user_state(state: PageState) -> None:
    username = _username(state)
    if not username:
        return
    auth_db = _get_auth_db(state)
    settings = auth_db.get_user_settings(username=username)
    explorer = settings.get("explorer") if isinstance(settings.get("explorer"), dict) else {}
    if explorer:
        state.explorer_view = str(explorer.get("view") or state.explorer_view)
        state.explorer_sort = str(explorer.get("sort") or state.explorer_sort)
        state.explorer_desc = bool(explorer.get("desc", state.explorer_desc))
        state.explorer_ext = str(explorer.get("ext") or state.explorer_ext)
        state.explorer_hidden_paths = [
            str(path)
            for path in (explorer.get("hidden_paths") or [])
            if str(path or "").strip()
        ]
        state.explorer_show_hidden = bool(explorer.get("show_hidden", state.explorer_show_hidden))
    ui_settings = settings.get("ui") if isinstance(settings.get("ui"), dict) else {}
    if ui_settings:
        theme = str(ui_settings.get("theme") or "").strip().lower()
        if theme in {"light", "dark"}:
            state.theme = theme
        if "ai_search_expand" in ui_settings:
            state.ai_search_expand = bool(ui_settings.get("ai_search_expand"))
    state.favorites = auth_db.list_favorites(username=username)
    state.saved_searches = auth_db.list_saved_searches(username=username)
    if not state.history:
        try:
            _get_telemetry(state)
        except Exception:
            pass
        state.history = _my_recent_queries(state.cfg, username, limit=24)


def _save_ui_settings(state: PageState) -> None:
    username = _username(state)
    if not username:
        return
    auth_db = _get_auth_db(state)
    settings = auth_db.get_user_settings(username=username)
    ui_settings = settings.get("ui") if isinstance(settings.get("ui"), dict) else {}
    ui_settings["theme"] = state.theme if state.theme in {"light", "dark"} else "light"
    ui_settings["ai_search_expand"] = bool(state.ai_search_expand)
    settings["ui"] = ui_settings
    auth_db.save_user_settings(username=username, settings=settings)


def _save_explorer_settings(state: PageState) -> None:
    username = _username(state)
    if not username:
        return
    auth_db = _get_auth_db(state)
    settings = auth_db.get_user_settings(username=username)
    settings["explorer"] = {
        "view": state.explorer_view,
        "sort": state.explorer_sort,
        "desc": state.explorer_desc,
        "ext": state.explorer_ext,
        "hidden_paths": list(dict.fromkeys(str(path) for path in state.explorer_hidden_paths if str(path or "").strip())),
        "show_hidden": bool(state.explorer_show_hidden),
    }
    auth_db.save_user_settings(username=username, settings=settings)
    _log_app_event(state, "explorer", "save_settings", details=settings["explorer"])
