"""NiceGUI web frontend for RAG Catalog."""

from __future__ import annotations

import argparse
import html
import json
import mimetypes
import os
import re
import sqlite3
import subprocess
import sys
import time
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path, PureWindowsPath
from typing import Any, Dict, List, Optional
from urllib.parse import quote

from fastapi import HTTPException
from fastapi.responses import FileResponse
from nicegui import app, events, run, ui

from rag_catalog.core.indexer_control import (
    get_current_command as _get_indexer_command,
    write_indexer_control as _write_indexer_control,
)
from rag_catalog.core.rag_core import RAGSearcher, load_config, save_config
from rag_catalog.core.telemetry_db import TelemetryDB
from rag_catalog.core.user_auth_db import UserAuthDB


PROJECT_ROOT = Path(__file__).resolve().parents[3]
APP_ICON_PATH = PROJECT_ROOT / "assets" / "brand" / "ico" / "favicon.ico"
LOGO_PATH = PROJECT_ROOT / "assets" / "brand" / "png" / "app-badge-128.png"

_STAGE_LABELS: Dict[str, str] = {
    "all": "Все этапы",
    "metadata": "metadata",
    "small": "small chunks",
    "large": "large chunks",
    "content": "content",
}
_CADENCE_LABELS: Dict[str, str] = {
    "hourly": "Каждый час",
    "daily": "Ежедневно",
    "weekly": "Еженедельно",
}
_DAY_LABELS = ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]
_DAY_RU = {"Mon": "Пн", "Tue": "Вт", "Wed": "Ср", "Thu": "Чт", "Fri": "Пт", "Sat": "Сб", "Sun": "Вс"}


def _open_log(log_path: "Path", label: str) -> "Any":
    """Открыть лог-файл на дозапись, записать заголовок с временем."""
    log_path.parent.mkdir(exist_ok=True)
    fh = open(log_path, "a", encoding="utf-8", errors="replace")  # noqa: WPS515
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    fh.write(f"\n{'='*60}\n{label}  {ts}\n{'='*60}\n")
    fh.flush()
    return fh


def _windows_detached_creationflags() -> int:
    flags = 0
    for name in ("CREATE_NO_WINDOW", "DETACHED_PROCESS", "CREATE_NEW_PROCESS_GROUP", "CREATE_BREAKAWAY_FROM_JOB"):
        flags |= int(getattr(subprocess, name, 0) or 0)
    return flags


def _is_process_alive(pid: int) -> bool:
    if int(pid or 0) <= 0:
        return False
    try:
        os.kill(int(pid), 0)
    except PermissionError:
        return True
    except ProcessLookupError:
        return False
    except OSError:
        return False
    return True


def _terminate_process(pid: int) -> bool:
    """
    Жесткая остановка процесса. На Windows — taskkill /F, на POSIX — SIGTERM.
    Используется как fallback после кооперативного cancel.
    """
    pid_int = int(pid or 0)
    if pid_int <= 0:
        return False
    try:
        if sys.platform.startswith("win"):
            subprocess.run(
                ["taskkill", "/PID", str(pid_int), "/F", "/T"],
                capture_output=True,
                check=False,
            )
        else:
            import signal as _signal
            os.kill(pid_int, _signal.SIGTERM)
        return True
    except OSError:
        return False


def _find_live_running_index_run(telemetry: TelemetryDB) -> Optional[Dict[str, Any]]:
    rows = telemetry.fetch_dicts(
        "SELECT * FROM index_runs WHERE status='running' ORDER BY ts_started DESC LIMIT 20"
    )
    for row in rows:
        pid = _safe_int(row.get("worker_pid"), 0)
        if _is_process_alive(pid):
            return row
    return None


def _find_live_running_ocr_run(telemetry: TelemetryDB) -> Optional[Dict[str, Any]]:
    rows = telemetry.fetch_dicts(
        "SELECT * FROM ocr_runs WHERE status='running' ORDER BY ts_started DESC LIMIT 20"
    )
    for row in rows:
        pid = _safe_int(row.get("worker_pid"), 0)
        if _is_process_alive(pid):
            return row
    return None


def _launch_indexer(
    cfg: Dict[str, Any],
    *,
    stage: str = "all",
    recreate: bool = False,
    workers: Optional[int] = None,
    max_chunks: Optional[int] = None,
    skip_inline_ocr: bool = False,
) -> int:
    """Запустить index_rag как фоновый процесс. Возвращает PID."""
    telemetry = TelemetryDB(str(_telemetry_db_path(cfg)))
    active_run = _find_live_running_index_run(telemetry)
    if active_run:
        active_pid = _safe_int(active_run.get("worker_pid"), 0)
        raise RuntimeError(
            f"Индексация уже запущена (PID {active_pid}). Дождитесь завершения текущего процесса."
        )
    env = os.environ.copy()
    env["PYTHONPATH"] = str(PROJECT_ROOT / "src")
    env["PYTHONIOENCODING"] = "utf-8"
    args = [
        sys.executable, "-m", "rag_catalog.core.index_rag",
        "--catalog", str(cfg.get("catalog_path") or ""),
        "--collection", str(cfg.get("collection_name") or ""),
        "--stage", stage,
        "--workers", str(int(workers or cfg.get("index_read_workers") or 4)),
        "--max-chunks", str(int(max_chunks or cfg.get("index_max_chunks") or 2000)),
    ]
    qdrant_url = str(cfg.get("qdrant_url") or "")
    if qdrant_url:
        args += ["--url", qdrant_url]
    else:
        args += ["--db", str(cfg.get("qdrant_db_path") or "")]
    if recreate:
        args.append("--recreate")
    if skip_inline_ocr:
        args.append("--no-ocr")
    # Сбрасываем контрольный файл — на случай если в нём осталась команда
    # cancel/pause от предыдущего прогона.
    try:
        _write_indexer_control("running")
    except (OSError, ValueError):
        pass
    log_fh = _open_log(PROJECT_ROOT / "logs" / "indexer.log", f"INDEXER  stage={stage}")
    try:
        proc = subprocess.Popen(
            args,
            cwd=str(PROJECT_ROOT),
            env=env,
            stdout=log_fh,
            stderr=log_fh,
            creationflags=_windows_detached_creationflags(),
        )
    finally:
        log_fh.close()
    return proc.pid


def _launch_ocr(cfg: Dict[str, Any], *, min_text_len: int = 50) -> int:
    """Запустить ocr_pdfs как фоновый процесс. Возвращает PID."""
    telemetry = TelemetryDB(str(_telemetry_db_path(cfg)))
    active_run = _find_live_running_ocr_run(telemetry)
    if active_run:
        active_pid = _safe_int(active_run.get("worker_pid"), 0)
        raise RuntimeError(
            f"OCR уже запущен (PID {active_pid}). Дождитесь завершения текущего процесса."
        )
    env = os.environ.copy()
    env["PYTHONPATH"] = str(PROJECT_ROOT / "src")
    env["PYTHONIOENCODING"] = "utf-8"
    args = [sys.executable, "-m", "rag_catalog.core.ocr_pdfs",
            "--min-text-len", str(int(min_text_len))]
    log_fh = _open_log(PROJECT_ROOT / "logs" / "ocr.log", "OCR")
    try:
        proc = subprocess.Popen(
            args,
            cwd=str(PROJECT_ROOT),
            env=env,
            stdout=log_fh,
            stderr=log_fh,
            creationflags=_windows_detached_creationflags(),
        )
    finally:
        log_fh.close()
    return proc.pid


def _schedules_due(schedules: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Вернуть расписания, которые должны запуститься прямо сейчас (±1 мин)."""
    now = datetime.now(timezone.utc)
    due = []
    for sched in schedules:
        if not int(sched.get("enabled") or 0):
            continue
        sched_time = str(sched.get("time") or "03:00")
        try:
            hh, mm = int(sched_time[:2]), int(sched_time[3:5])
        except (ValueError, IndexError):
            continue
        cadence = str(sched.get("cadence") or "daily")
        days = sched.get("days") or []
        day_name = now.strftime("%a")  # "Mon", "Tue", ...
        if cadence == "weekly" and day_name not in days:
            continue
        if cadence == "daily" and days and day_name not in days:
            continue
        # Проверяем совпадение часа и минуты (±1 мин)
        if now.hour != hh or abs(now.minute - mm) > 1:
            continue
        # Не запускать дважды в одну минуту
        last_run = str(sched.get("last_run_at") or "")
        if last_run:
            try:
                lr = datetime.fromisoformat(last_run.replace("Z", "+00:00"))
                if (now - lr).total_seconds() < 90:
                    continue
            except ValueError:
                pass
        due.append(sched)
    return due

SEARCH_PRESETS = [
    ("Договоры", "договор поставки"),
    ("Счета", "счет на оплату"),
    ("Паспорта", "паспорт техника"),
    ("PDF", "pdf скан"),
    ("Таблицы", "реестр xlsx"),
]

FILE_PREVIEW_EXTENSIONS = {".txt", ".log", ".csv", ".json", ".md", ".py", ".ps1", ".xml", ".html", ".css"}
INLINE_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".svg"}
OFFICE_PREVIEW_EXTENSIONS = {".docx", ".xlsx", ".xls", ".rtf"}
PAGE_SIZE = 80
SYSTEM_FILE_EXTENSIONS = {
    ".dll",
    ".exe",
    ".msi",
    ".sys",
    ".inf",
    ".cat",
    ".pnf",
    ".bak",
    ".tmp",
    ".temp",
    ".dat",
    ".db",
    ".db3",
    ".sqlite",
    ".sqlite-journal",
    ".db-shm",
    ".db-wal",
    ".lock",
    ".cfg",
    ".ini",
    ".config",
    ".lnk",
    ".chm",
    ".toc",
    ".cldbin",
}

if LOGO_PATH.exists():
    app.add_static_file(local_file=LOGO_PATH, url_path="/rag-logo.png")

_MARK_SVG_PATH = PROJECT_ROOT / "assets" / "brand" / "svg" / "mark.svg"
if _MARK_SVG_PATH.exists():
    app.add_static_file(local_file=_MARK_SVG_PATH, url_path="/rag-mark.svg")


@dataclass
class PageState:
    cfg: Dict[str, Any]
    searcher: Optional[RAGSearcher] = None
    searcher_error: str = ""
    screen: str = "search"
    query: str = ""
    file_type: Optional[str] = None
    limit: int = 50
    content_only: bool = False
    title_only: bool = False
    history: List[str] = field(default_factory=list)
    results: List[Dict[str, Any]] = field(default_factory=list)
    search_error: str = ""
    searched_query: str = ""
    expanded_query: str = ""       # расширенный запрос от LLM (пусто если не менялся)
    rag_answer_text: str = ""      # RAG Q&A ответ (пусто если LLM отключён)
    rag_answer_loading: bool = False
    settings_section: str = "profile"  # активная секция в настройках (IDE-навигация)
    displayed_count: int = 10
    active_type_filter: Optional[str] = None
    explorer_path: Optional[str] = None
    explorer_filter: str = ""
    explorer_ext: str = "Все"
    explorer_sort: str = "По имени"
    explorer_desc: bool = False
    explorer_view: str = "Таблица"
    explorer_page: int = 0
    auth_db: Optional[UserAuthDB] = None
    current_user: Optional[Dict[str, Any]] = None
    auth_token: str = ""
    theme: str = "light"
    favorites: List[Dict[str, Any]] = field(default_factory=list)
    header_explorer_actions: Optional[ui.row] = None
    header_breadcrumbs: Optional[ui.row] = None
    telemetry: Optional[TelemetryDB] = None
    search_refine_terms: List[str] = field(default_factory=list)  # "уточнить в найденном"
    selected_result_paths: List[str] = field(default_factory=list)  # выбранные результаты
    analytics_tab: str = "overview"  # активная вкладка аналитики


def _file_url(full_path: str) -> str:
    try:
        p = PureWindowsPath(full_path)
        if not p.parts:
            return ""
        drive = p.drive
        encoded = "/".join(quote(part, safe="") for part in p.parts[1:])
        return "file:///" + drive + "/" + encoded
    except Exception:
        return ""


def _folder_url(full_path: str) -> str:
    try:
        p = PureWindowsPath(full_path).parent
        if not p.parts:
            return ""
        drive = p.drive
        encoded = "/".join(quote(part, safe="") for part in p.parts[1:])
        return "file:///" + drive + "/" + encoded
    except Exception:
        return ""


def _resolve_catalog_file(cfg: Dict[str, Any], raw_path: str) -> Optional[Path]:
    try:
        catalog = Path(str(cfg.get("catalog_path") or "")).resolve()
        if not catalog.exists() or not catalog.is_dir():
            return None
        candidate = Path(str(raw_path or "")).resolve()
        candidate.relative_to(catalog)
    except Exception:
        return None
    if not candidate.exists() or not candidate.is_file():
        return None
    return candidate


def _viewer_file_url(full_path: str) -> str:
    value = str(full_path or "").strip()
    if not value:
        return ""
    return f"/api/view-file?path={quote(value, safe='')}"


def _preview_office_file(path: Path, limit: int = 12000) -> str:
    ext = path.suffix.lower()
    if ext == ".docx":
        try:
            from docx import Document  # noqa: PLC0415

            doc = Document(path)
            parts: List[str] = [p.text for p in doc.paragraphs if p.text]
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [str(cell.text or "").strip() for cell in row.cells]
                    line = " | ".join(item for item in row_cells if item)
                    if line:
                        parts.append(line)
            return "\n".join(parts)[:limit] or "Документ пуст."
        except Exception as exc:
            return f"Не удалось прочитать DOCX: {exc}"
    if ext == ".xlsx":
        try:
            from openpyxl import load_workbook  # noqa: PLC0415

            wb = load_workbook(path, read_only=True, data_only=True)
            parts: List[str] = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                parts.append(f"Лист: {sheet_name}")
                for row in ws.iter_rows(values_only=True):
                    row_text = " | ".join(str(c) if c is not None else "" for c in row)
                    if row_text.strip():
                        parts.append(row_text)
                    if len("\n".join(parts)) >= limit:
                        return "\n".join(parts)[:limit]
            return "\n".join(parts)[:limit] or "Таблица пуста."
        except Exception as exc:
            return f"Не удалось прочитать XLSX: {exc}"
    if ext == ".xls":
        try:
            import xlrd  # type: ignore  # noqa: PLC0415

            book = xlrd.open_workbook(str(path))
            parts = []
            for sheet in book.sheets():
                parts.append(f"Лист: {sheet.name}")
                for row_idx in range(sheet.nrows):
                    row = sheet.row_values(row_idx)
                    row_text = " | ".join(str(v) if v not in ("", None) else "" for v in row)
                    if row_text.strip():
                        parts.append(row_text)
                    if len("\n".join(parts)) >= limit:
                        return "\n".join(parts)[:limit]
            return "\n".join(parts)[:limit] or "Таблица пуста."
        except Exception as exc:
            return f"Не удалось прочитать XLS: {exc}"
    if ext == ".rtf":
        try:
            raw = path.read_text(encoding="utf-8", errors="replace")
            no_ctrl = re.sub(r"\\[a-z]+\d* ?", " ", raw)
            no_braces = re.sub(r"[{}]", "", no_ctrl)
            clean = re.sub(r"\s+", " ", no_braces).strip()
            return clean[:limit] or "Документ пуст."
        except Exception as exc:
            return f"Не удалось прочитать RTF: {exc}"
    return "Для этого типа файла доступно открытие или скачивание."


@app.get("/api/view-file")
def api_view_file(path: str) -> FileResponse:
    resolved = _resolve_catalog_file(load_config(), path)
    if resolved is None:
        raise HTTPException(status_code=404, detail="Файл не найден или недоступен")
    media_type = mimetypes.guess_type(str(resolved))[0] or "application/octet-stream"
    return FileResponse(str(resolved), media_type=media_type, filename=resolved.name)


def _telemetry_db_path(cfg: Dict[str, Any]) -> Path:
    explicit = str(cfg.get("telemetry_db_path") or "").strip()
    if explicit:
        return Path(explicit)
    return Path(str(cfg.get("qdrant_db_path") or "")) / "rag_telemetry.db"


def _safe_int(value: Any, default: int = 0) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return int(default)


def _resolve_index_recovery_stage(telemetry: TelemetryDB, active_run: Dict[str, Any]) -> str:
    run_id = str(active_run.get("run_id") or "")
    if run_id:
        stage_rows = telemetry.fetch_dicts(
            """
            SELECT stage
            FROM index_stage_progress
            WHERE run_id=? AND status='running'
            ORDER BY ts_updated DESC, ts_started DESC
            LIMIT 1
            """,
            [run_id],
        )
        if stage_rows:
            candidate = str(stage_rows[0].get("stage") or "").strip().lower()
            if candidate in _STAGE_LABELS:
                return candidate
    note = str(active_run.get("note") or "")
    match = re.search(r"stage=(all|metadata|small|large|content)", note.lower())
    if match:
        return match.group(1)
    return "all"


def _recover_background_tasks(cfg: Dict[str, Any]) -> None:
    telemetry = TelemetryDB(str(_telemetry_db_path(cfg)))
    settings = telemetry.get_index_settings() if hasattr(telemetry, "get_index_settings") else {}
    workers = _safe_int(settings.get("workers") or cfg.get("index_read_workers") or 4, 4)
    max_chunks = _safe_int(settings.get("max_chunks") or cfg.get("index_max_chunks") or 2000, 2000)
    skip_inline_ocr = bool(settings.get("skip_inline_ocr"))
    ocr_min_text_len = _safe_int(settings.get("ocr_min_text_len") or 50, 50)

    live_index = _find_live_running_index_run(telemetry)
    active_index = telemetry.get_active_index_run() if hasattr(telemetry, "get_active_index_run") else None
    if not live_index and active_index:
        recovery_stage = _resolve_index_recovery_stage(telemetry, active_index)
        telemetry.finalize_running_index_runs(
            status="cancelled",
            note="server_restart_recovery",
        )
        _launch_indexer(
            cfg,
            stage=recovery_stage,
            workers=workers,
            max_chunks=max_chunks,
            skip_inline_ocr=skip_inline_ocr,
        )

    live_ocr = _find_live_running_ocr_run(telemetry)
    active_ocr = telemetry.get_active_ocr_run() if hasattr(telemetry, "get_active_ocr_run") else None
    if not live_ocr and active_ocr:
        if hasattr(telemetry, "finalize_running_ocr_runs"):
            telemetry.finalize_running_ocr_runs(
                status="cancelled",
                note="server_restart_recovery",
            )
        _launch_ocr(
            cfg,
            min_text_len=ocr_min_text_len,
        )


CONFIG_PATH_KEYS = {
    "catalog_path",
    "qdrant_db_path",
    "qdrant_url",
    "log_file",
    "collection_name",
    "telemetry_db_path",
}


def _save_config_patch(values: Dict[str, Any]) -> Dict[str, Any]:
    clean = {key: str(values.get(key) or "").strip() for key in CONFIG_PATH_KEYS if key in values}
    cfg = load_config()
    cfg.update(clean)
    save_config(cfg)
    return cfg


def _users_db_path(cfg: Dict[str, Any]) -> Path:
    explicit = str(cfg.get("users_db_path") or "").strip()
    if explicit:
        return Path(explicit)
    return Path(str(cfg.get("qdrant_db_path") or ".")) / "rag_users.db"


def _get_auth_db(state: PageState) -> UserAuthDB:
    path = _users_db_path(state.cfg)
    if state.auth_db is None or Path(getattr(state.auth_db, "db_path", "")) != path:
        state.auth_db = UserAuthDB(str(path))
    return state.auth_db


def _refresh_current_user(state: PageState) -> None:
    if not state.current_user:
        return
    user = _get_auth_db(state).get_user(username=str(state.current_user.get("username") or ""))
    if user:
        state.current_user = user


def _username(state: PageState) -> str:
    return str((state.current_user or {}).get("username") or "").strip().lower()


def _get_telemetry(state: PageState) -> TelemetryDB:
    path = _telemetry_db_path(state.cfg)
    if state.telemetry is None or Path(getattr(state.telemetry, "db_path", "")) != path:
        state.telemetry = TelemetryDB(str(path))
    return state.telemetry


def _log_app_event(
    state: PageState,
    feature: str,
    action: str,
    *,
    ok: bool = True,
    details: Optional[Dict[str, Any]] = None,
) -> None:
    try:
        _get_telemetry(state).log_app_event(
            username=_username(state),
            screen=state.screen,
            feature=feature,
            action=action,
            ok=ok,
            details=details or {},
        )
    except Exception:
        pass


def _load_user_state(state: PageState) -> None:
    username = _username(state)
    if not username:
        return
    auth_db = _get_auth_db(state)
    settings = auth_db.get_user_settings(username=username)
    explorer = settings.get("explorer") if isinstance(settings.get("explorer"), dict) else {}
    if explorer:
        state.explorer_view = str(explorer.get("view") or state.explorer_view)
        state.explorer_sort = str(explorer.get("sort") or state.explorer_sort)
        state.explorer_desc = bool(explorer.get("desc", state.explorer_desc))
        state.explorer_ext = str(explorer.get("ext") or state.explorer_ext)
    ui_settings = settings.get("ui") if isinstance(settings.get("ui"), dict) else {}
    if ui_settings:
        theme = str(ui_settings.get("theme") or "").strip().lower()
        if theme in {"light", "dark"}:
            state.theme = theme
    state.favorites = auth_db.list_favorites(username=username)
    # Восстановить личную историю поиска из БД (персистентна между сессиями)
    if not state.history:
        try:
            _get_telemetry(state)
        except Exception:
            pass
        state.history = _my_recent_queries(state.cfg, username, limit=24)


def _save_ui_settings(state: PageState) -> None:
    username = _username(state)
    if not username:
        return
    auth_db = _get_auth_db(state)
    settings = auth_db.get_user_settings(username=username)
    ui_settings = settings.get("ui") if isinstance(settings.get("ui"), dict) else {}
    ui_settings["theme"] = state.theme if state.theme in {"light", "dark"} else "light"
    settings["ui"] = ui_settings
    auth_db.save_user_settings(username=username, settings=settings)


def _save_explorer_settings(state: PageState) -> None:
    username = _username(state)
    if not username:
        return
    auth_db = _get_auth_db(state)
    settings = auth_db.get_user_settings(username=username)
    settings["explorer"] = {
        "view": state.explorer_view,
        "sort": state.explorer_sort,
        "desc": state.explorer_desc,
        "ext": state.explorer_ext,
    }
    auth_db.save_user_settings(username=username, settings=settings)
    _log_app_event(state, "explorer", "save_settings", details=settings["explorer"])


def _favorite_key(path: str) -> str:
    return str(path or "").strip().casefold()


def _is_favorite(state: PageState, path: str) -> bool:
    key = _favorite_key(path)
    return any(_favorite_key(str(item.get("path") or "")) == key for item in state.favorites)


def _favorite_type(path: Path) -> str:
    return "folder" if path.is_dir() else "file"


def _toggle_favorite(state: PageState, path: Path, *, item_type: Optional[str] = None, title: str = "") -> bool:
    username = _username(state)
    if not username:
        ui.notify("Войдите, чтобы сохранять избранное.", type="warning")
        return False
    auth_db = _get_auth_db(state)
    path_value = str(path)
    active = _is_favorite(state, path_value)
    if active:
        auth_db.remove_favorite(username=username, path=path_value)
        _log_app_event(state, "favorites", "remove", details={"path": path_value})
    else:
        auth_db.add_favorite(
            username=username,
            item_type=item_type or _favorite_type(path),
            path=path_value,
            title=title or path.name or path_value,
        )
        _log_app_event(state, "favorites", "add", details={"path": path_value, "item_type": item_type or _favorite_type(path)})
    state.favorites = auth_db.list_favorites(username=username)
    return not active


def _db_query_dicts(db_path: Path, query: str, params: Optional[tuple] = None) -> List[Dict[str, Any]]:
    if not db_path.exists():
        return []
    conn = sqlite3.connect(str(db_path))
    conn.row_factory = sqlite3.Row
    try:
        cur = conn.execute(query, params or ())
        return [dict(row) for row in cur.fetchall()]
    except Exception:
        return []
    finally:
        conn.close()


def _dedupe_queries(values: List[str], limit: int = 12) -> List[str]:
    out: List[str] = []
    seen = set()
    for value in values:
        query = re.sub(r"\s+", " ", str(value or "")).strip()
        key = query.lower()
        if not query or key in seen:
            continue
        seen.add(key)
        out.append(query)
        if len(out) >= limit:
            break
    return out


def _my_recent_queries(cfg: Dict[str, Any], username: str = "", limit: int = 12) -> List[str]:
    """Последние запросы текущего пользователя из БД."""
    rows = _db_query_dicts(
        _telemetry_db_path(cfg),
        """
        SELECT COALESCE(NULLIF(query_original, ''), query) AS query
        FROM search_logs
        WHERE query <> '' AND lower(username) = ?
        ORDER BY id DESC
        LIMIT 80
        """,
        (username.strip().lower(),),
    )
    return _dedupe_queries([str(row.get("query") or "") for row in rows], limit=limit)


def _popular_queries(cfg: Dict[str, Any], exclude_username: str = "", limit: int = 10) -> List[str]:
    """Топ запросов по частоте среди всех пользователей (кроме текущего)."""
    rows = _db_query_dicts(
        _telemetry_db_path(cfg),
        """
        SELECT COALESCE(NULLIF(query_original, ''), query) AS query, COUNT(*) AS cnt
        FROM search_logs
        WHERE query <> '' AND lower(username) != ?
        GROUP BY lower(COALESCE(NULLIF(query_original, ''), query))
        ORDER BY cnt DESC
        LIMIT 40
        """,
        (exclude_username.strip().lower(),),
    )
    return _dedupe_queries([str(row.get("query") or "") for row in rows], limit=limit)


def _search_suggestions(state: PageState, typed: str = "") -> List[str]:
    username = _username(state)
    personal = _dedupe_queries([*state.history, *_my_recent_queries(state.cfg, username, limit=12)], limit=24)
    needle = typed.strip().lower()
    if not needle:
        return personal[:12]
    starts = [item for item in personal if item.lower().startswith(needle)]
    contains = [item for item in personal if needle in item.lower() and item not in starts]
    return [*starts, *contains][:12]


def _telegram_deeplink(bot_link: str, purpose: str, token: str) -> str:
    base = str(bot_link or "").strip()
    value = str(token or "").strip()
    if not base or not value:
        return ""
    joiner = "&" if "?" in base else "?"
    return f"{base}{joiner}start={purpose}_{quote(value, safe='')}"


def _remember_query(state: PageState, query: str) -> None:
    clean = re.sub(r"\s+", " ", str(query or "")).strip()
    if clean:
        state.history = _dedupe_queries([clean, *state.history], limit=24)


def _normalize_search_results(results: Any) -> List[Dict[str, Any]]:
    if results is None:
        return []
    if not isinstance(results, list):
        return []
    return [item for item in results if isinstance(item, dict)]


def _run_catalog_search(
    searcher: RAGSearcher,
    *,
    query: str,
    query_original: str,
    query_used: str,
    limit: int,
    file_type: Optional[str],
    content_only: bool,
    title_only: bool,
    username: str = "",
) -> List[Dict[str, Any]]:
    results = _normalize_search_results(
        searcher.search(
            query,
            limit=limit,
            file_type=file_type,
            content_only=content_only,
            title_only=title_only,
            source="nicegui",
            username=username,
            query_original=query_original,
        )
    )
    if content_only or title_only:
        return results

    # If the vector path returns an empty/invalid value, keep name/path search usable.
    # Merge with fallback lexical results, keeping the best rank_score for each unique item.
    try:
        fallback = _normalize_search_results(
            searcher._lexical_catalog_search(  # noqa: SLF001 - UI fallback for catalog metadata search
                query=query_used,
                limit=max(limit, 10),
                file_type=file_type,
                content_only=False,
                title_only=title_only,
            )
        )
    except Exception:
        return results

    if not fallback:
        return results

    def _rank_key(item: Dict[str, Any]) -> float:
        return float(item.get("rank_score") or item.get("score") or 0)

    seen: Dict[str, Dict[str, Any]] = {}
    for item in [*results, *fallback]:
        key = f"{item.get('full_path')}::{item.get('chunk_index')}::{item.get('type')}"
        if key not in seen or _rank_key(item) > _rank_key(seen[key]):
            seen[key] = item
    merged = sorted(seen.values(), key=_rank_key, reverse=True)
    return merged[:limit]


def _format_file_size(size_b: int) -> str:
    if size_b >= 1_048_576:
        return f"{size_b / 1_048_576:.1f} МБ"
    if size_b >= 1024:
        return f"{size_b / 1024:.1f} КБ"
    return f"{size_b} Б"


def _clean_text(value: Any) -> str:
    raw = str(value or "")
    raw = re.sub(r"<[^>]{1,200}>", " ", raw)
    raw = html.unescape(raw)
    raw = re.sub(r"[ \t]+", " ", raw)
    raw = re.sub(r"\s*\n\s*", "\n", raw)
    return raw.strip()


def _directory_children(path: str, limit: int = 60) -> Dict[str, Any]:
    p = Path(path)
    out: Dict[str, Any] = {"exists": p.exists(), "is_dir": p.is_dir(), "dirs": [], "files": []}
    if not p.exists() or not p.is_dir():
        return out
    try:
        entries = sorted(
            [x for x in p.iterdir() if not x.name.startswith(".") and not x.name.startswith("~$")],
            key=lambda x: (not x.is_dir(), x.name.lower()),
        )
    except Exception as exc:
        out["error"] = str(exc)
        return out
    for child in entries[:limit]:
        item = {"name": child.name, "path": str(child)}
        if child.is_dir():
            out["dirs"].append(item)
        elif child.is_file():
            try:
                item["size"] = _format_file_size(child.stat().st_size)
            except Exception:
                item["size"] = ""
            out["files"].append(item)
    out["truncated"] = len(entries) > limit
    return out


def _preview_file(path: Path, limit: int = 6000) -> str:
    if not path.exists() or not path.is_file():
        return "Файл недоступен."
    if path.suffix.lower() not in FILE_PREVIEW_EXTENSIONS:
        return "Для этого типа файла доступно открытие или скачивание."
    try:
        return path.read_text(encoding="utf-8", errors="replace")[:limit]
    except Exception as exc:
        return f"Не удалось прочитать файл: {exc}"


def _open_os_path(path: str) -> None:
    value = str(path or "").strip()
    if not value:
        return
    try:
        subprocess.Popen(["explorer", value])
    except Exception as exc:
        ui.notify(f"Не удалось открыть проводник ОС: {exc}", type="negative")


def _within_catalog(root: Path, candidate: Path) -> bool:
    try:
        candidate.resolve().relative_to(root.resolve())
        return True
    except Exception:
        return False


def _safe_explorer_path(state: PageState) -> Path:
    root = Path(str(state.cfg.get("catalog_path") or ""))
    if state.explorer_path:
        candidate = Path(state.explorer_path)
        if candidate.exists() and _within_catalog(root, candidate):
            return candidate
    state.explorer_path = str(root)
    return root


def _read_index_stats(cfg: Dict[str, Any]) -> Dict[str, Any]:
    state_file = Path(str(cfg.get("qdrant_db_path") or "")) / "index_state.json"
    out: Dict[str, Any] = {
        "found": False,
        "state_file": str(state_file),
        "total": 0,
        "total_size_bytes": 0,
        "by_ext": {},
        "by_ext_size": {},
    }
    if not state_file.exists():
        return out
    try:
        data = json.loads(state_file.read_text(encoding="utf-8"))
        files = data.get("files", {})
    except Exception as exc:
        out["error"] = str(exc)
        return out
    by_ext: Dict[str, int] = {}
    by_ext_size: Dict[str, int] = {}
    total_size = 0
    for key, meta in files.items():
        path = Path(str(key))
        ext = path.suffix.lower() or "(без расширения)"
        by_ext[ext] = by_ext.get(ext, 0) + 1
        size = 0
        if isinstance(meta, dict):
            fingerprint = str(meta.get("fingerprint") or "")
            try:
                size = int(float(fingerprint.split("_", 1)[0])) if "_" in fingerprint else 0
            except (TypeError, ValueError):
                size = 0
            if not size and meta.get("size"):
                try:
                    size = int(float(meta.get("size") or 0))
                except (TypeError, ValueError):
                    size = 0
        if not size:
            try:
                size = path.stat().st_size
            except OSError:
                size = 0
        total_size += size
        by_ext_size[ext] = by_ext_size.get(ext, 0) + size
    out.update({
        "found": True,
        "total": len(files),
        "total_size_bytes": total_size,
        "by_ext": dict(sorted(by_ext.items(), key=lambda x: x[1], reverse=True)),
        "by_ext_size": dict(sorted(by_ext_size.items(), key=lambda x: x[1], reverse=True)),
    })
    try:
        out["last_modified"] = time.strftime("%d.%m.%Y %H:%M", time.localtime(state_file.stat().st_mtime))
    except Exception:
        pass
    return out


def _format_bytes(value: Any) -> str:
    try:
        size = float(value or 0)
    except (TypeError, ValueError):
        size = 0.0
    units = ["Б", "КБ", "МБ", "ГБ", "ТБ"]
    idx = 0
    while size >= 1024 and idx < len(units) - 1:
        size /= 1024
        idx += 1
    if idx == 0:
        return f"{int(size)} {units[idx]}"
    return f"{size:.1f} {units[idx]}"


def _format_duration_seconds(value: Any) -> str:
    try:
        seconds = int(float(value or 0))
    except (TypeError, ValueError):
        seconds = 0
    if seconds <= 0:
        return "0 сек"
    hours, rem = divmod(seconds, 3600)
    minutes, secs = divmod(rem, 60)
    if hours:
        return f"{hours} ч {minutes:02d} мин"
    if minutes:
        return f"{minutes} мин {secs:02d} сек"
    return f"{secs} сек"


def _duration_between(start: Any, finish: Any = None) -> int:
    start_s = str(start or "")
    finish_s = str(finish or "")
    if not start_s:
        return 0
    try:
        start_dt = datetime.fromisoformat(start_s)
        finish_dt = datetime.fromisoformat(finish_s) if finish_s else datetime.now(start_dt.tzinfo)
        return max(0, int((finish_dt - start_dt).total_seconds()))
    except ValueError:
        return 0


def _read_index_telemetry(cfg: Dict[str, Any]) -> Dict[str, Any]:
    db_path = _telemetry_db_path(cfg)
    last_run = _db_query_dicts(
        db_path,
        """
        SELECT *,
               CAST((julianday(COALESCE(ts_finished, CURRENT_TIMESTAMP)) - julianday(ts_started)) * 86400 AS INTEGER) AS duration_sec
        FROM index_runs
        ORDER BY ts_started DESC
        LIMIT 1
        """,
    )
    active_runs = _db_query_dicts(
        db_path,
        """
        SELECT *,
               CAST((julianday(CURRENT_TIMESTAMP) - julianday(ts_started)) * 86400 AS INTEGER) AS duration_sec
        FROM index_runs
        WHERE status='running'
        ORDER BY ts_started DESC
        LIMIT 3
        """,
    )
    active_run_ids = [str(row.get("run_id") or "") for row in active_runs if str(row.get("run_id") or "")]
    active_stages: List[Dict[str, Any]] = []
    if active_run_ids:
        placeholders = ",".join("?" for _ in active_run_ids)
        active_stages = _db_query_dicts(
            db_path,
            f"""
            SELECT *,
                   CAST((julianday(COALESCE(ts_finished, CURRENT_TIMESTAMP)) - julianday(ts_started)) * 86400 AS INTEGER) AS duration_sec
            FROM index_stage_progress
            WHERE run_id IN ({placeholders})
            ORDER BY ts_started, stage
            """,
            tuple(active_run_ids),
        )
    latest_stages = _db_query_dicts(
        db_path,
        """
        SELECT *,
               CAST((julianday(COALESCE(ts_finished, CURRENT_TIMESTAMP)) - julianday(ts_started)) * 86400 AS INTEGER) AS duration_sec
        FROM index_stage_progress
        WHERE run_id=(SELECT run_id FROM index_runs ORDER BY ts_started DESC LIMIT 1)
        ORDER BY CASE stage WHEN 'metadata' THEN 1 WHEN 'small' THEN 2 WHEN 'large' THEN 3 WHEN 'content' THEN 4 ELSE 9 END
        """,
    )
    stage_summary = _db_query_dicts(
        db_path,
        """
        WITH finished AS (
            SELECT stage,
                   status,
                   total_files,
                   processed_files,
                   added_files,
                   updated_files,
                   skipped_files,
                   error_files,
                   points_added,
                   ts_started,
                   ts_finished,
                   CAST((julianday(ts_finished) - julianday(ts_started)) * 86400 AS INTEGER) AS duration_sec
            FROM index_stage_progress
            WHERE ts_finished IS NOT NULL
        ),
        latest AS (
            SELECT f.*
            FROM finished f
            JOIN (
                SELECT stage, MAX(ts_started) AS ts_started
                FROM finished
                GROUP BY stage
            ) m ON m.stage=f.stage AND m.ts_started=f.ts_started
        ),
        avg_by_stage AS (
            SELECT stage,
                   COUNT(*) AS runs_count,
                   AVG(duration_sec) AS avg_duration_sec,
                   AVG(processed_files) AS avg_processed_files
            FROM finished
            GROUP BY stage
        )
        SELECT latest.stage,
               latest.status,
               latest.total_files,
               latest.processed_files,
               latest.added_files,
               latest.updated_files,
               latest.skipped_files,
               latest.error_files,
               latest.points_added,
               latest.ts_started,
               latest.ts_finished,
               latest.duration_sec AS last_duration_sec,
               avg_by_stage.runs_count,
               CAST(avg_by_stage.avg_duration_sec AS INTEGER) AS avg_duration_sec,
               CAST(avg_by_stage.avg_processed_files AS INTEGER) AS avg_processed_files
        FROM latest
        JOIN avg_by_stage ON avg_by_stage.stage=latest.stage
        ORDER BY CASE latest.stage WHEN 'metadata' THEN 1 WHEN 'small' THEN 2 WHEN 'large' THEN 3 WHEN 'content' THEN 4 ELSE 9 END
        """,
    )
    overall = _db_query_dicts(
        db_path,
        """
        SELECT COUNT(*) AS runs_count,
               CAST(AVG((julianday(ts_finished) - julianday(ts_started)) * 86400) AS INTEGER) AS avg_duration_sec,
               MAX(total_files) AS max_total_files,
               AVG(total_files) AS avg_total_files
        FROM index_runs
        WHERE ts_finished IS NOT NULL
        """,
    )
    active_ocr = _db_query_dicts(
        db_path,
        """
        SELECT *,
               CAST((julianday(CURRENT_TIMESTAMP) - julianday(ts_started)) * 86400 AS INTEGER) AS duration_sec
        FROM ocr_runs
        WHERE status='running'
        ORDER BY ts_started DESC
        LIMIT 1
        """,
    )
    last_ocr = _db_query_dicts(
        db_path,
        """
        SELECT *,
               CAST((julianday(COALESCE(ts_finished, CURRENT_TIMESTAMP)) - julianday(ts_started)) * 86400 AS INTEGER) AS duration_sec
        FROM ocr_runs
        ORDER BY ts_started DESC
        LIMIT 1
        """,
    )
    ocr_summary = _db_query_dicts(
        db_path,
        """
        SELECT COUNT(*) AS runs_count,
               CAST(AVG((julianday(ts_finished) - julianday(ts_started)) * 86400) AS INTEGER) AS avg_duration_sec,
               AVG(found_scanned) AS avg_found_scanned,
               AVG(processed_pdfs) AS avg_processed_pdfs
        FROM ocr_runs
        WHERE ts_finished IS NOT NULL
        """,
    )
    return {
        "last_run": last_run[0] if last_run else None,
        "active_runs": active_runs,
        "active_stages": active_stages,
        "latest_stages": latest_stages,
        "stage_summary": stage_summary,
        "overall": overall[0] if overall else {},
        "active_ocr": active_ocr[0] if active_ocr else None,
        "last_ocr": last_ocr[0] if last_ocr else None,
        "ocr_summary": ocr_summary[0] if ocr_summary else {},
    }


def _ensure_searcher(state: PageState) -> Optional[RAGSearcher]:
    if state.searcher is not None:
        return state.searcher
    try:
        state.searcher = RAGSearcher(state.cfg)
    except Exception as exc:
        state.searcher_error = str(exc)
        return None
    if not state.searcher.connected:
        state.searcher_error = "Нет подключения к Qdrant."
    return state.searcher


def _is_admin(state: PageState) -> bool:
    return str((state.current_user or {}).get("role") or "") == "admin"


def _show_system_files(state: PageState) -> bool:
    if not _is_admin(state):
        return False
    try:
        auth = state.auth_db if state.auth_db is not None else _get_auth_db(state)
        return bool(auth.get_show_system_files_for_admin())
    except Exception:
        return False


def _is_system_file(path_or_ext: str | Path) -> bool:
    value = str(path_or_ext or "")
    ext = Path(value).suffix.lower() or value.lower()
    name = Path(value).name.lower()
    if not ext:
        return False
    if ext in SYSTEM_FILE_EXTENSIONS:
        return True
    if ext.endswith("_"):
        return True
    if ext in {".sample", ".backup", ".sav", ".asd"}:
        return True
    return name.startswith("~$") or name.endswith(".tmp")


def _result_kind(result: Dict[str, Any]) -> str:
    if result.get("type") == "folder_metadata":
        return "Каталог"
    ext = str(result.get("extension") or "").lower()
    if ext in {".xlsx", ".xls", ".csv"}:
        return "Таблица"
    if ext == ".pdf":
        return "PDF"
    if ext in {".docx", ".doc"}:
        return "Документ"
    return "Файл"


def _result_group(result: Dict[str, Any]) -> str:
    text = " ".join(
        str(result.get(key, "") or "").lower()
        for key in ("filename", "path", "type", "text", "extension")
    )
    if str(result.get("type") or "") == "folder_metadata":
        return "Каталоги"
    if any(word in text for word in ("птс", "псм", "стс", "техпаспорт", "техническ", "паспорт транспорт", "экскаватор")):
        return "Техпаспорта ТС"
    if any(word in text for word in ("паспорт", "удостоверен")):
        return "Паспорта и удостоверения"
    if any(word in text for word in ("договор", "соглашен", "контракт")):
        return "Договоры"
    if any(word in text for word in ("счет", "счёт", "оплат", "платеж")):
        return "Счета и платежи"
    if str(result.get("extension") or "").lower() in {".xlsx", ".xls", ".csv"}:
        return "Таблицы"
    if str(result.get("extension") or "").lower() == ".pdf":
        return "PDF"
    return "Другие файлы"


def _grouped_results(results: List[Dict[str, Any]]) -> List[tuple[str, List[Dict[str, Any]]]]:
    order = [
        "Каталоги",
        "Техпаспорта ТС",
        "Паспорта и удостоверения",
        "Договоры",
        "Счета и платежи",
        "Таблицы",
        "PDF",
        "Другие файлы",
    ]
    grouped: Dict[str, List[Dict[str, Any]]] = {}
    for result in results:
        grouped.setdefault(_result_group(result), []).append(result)
    return [
        (group, sorted(grouped[group], key=lambda item: float(item.get("rank_score") or item.get("score") or 0), reverse=True))
        for group in order
        if group in grouped
    ]


def _file_icon_svg(path_or_ext: str, kind: str = "Файл") -> str:
    ext = Path(str(path_or_ext or "")).suffix.lower() or str(path_or_ext or "").lower()
    if kind == "Каталог":
        color, label = "#f2b84b", ""
    elif ext in {".doc", ".docx"}:
        color, label = "#2b579a", "W"
    elif ext in {".xls", ".xlsx", ".csv"}:
        color, label = "#217346", "X"
    elif ext in {".ppt", ".pptx"}:
        color, label = "#d24726", "P"
    elif ext == ".pdf":
        color, label = "#d32f2f", "PDF"
    elif ext in {".zip", ".rar", ".7z"}:
        color, label = "#d39a18", "ZIP"
    elif ext in {".jpg", ".jpeg", ".png", ".gif", ".webp"}:
        color, label = "#7b4ab8", "IMG"
    elif ext in {".txt", ".log", ".md"}:
        color, label = "#64748b", "TXT"
    elif ext in {".mp3", ".wav", ".flac"}:
        color, label = "#0e7490", "AUD"
    elif ext in {".mp4", ".avi", ".mov", ".mkv"}:
        color, label = "#7c3aed", "VID"
    elif ext in {".json", ".xml", ".html", ".htm", ".css", ".js", ".py", ".ps1", ".cmd", ".bat"}:
        color, label = "#475569", "DEV"
    elif ext in {".dwg", ".svg", ".psd", ".ico"}:
        color, label = "#0891b2", "ART"
    elif _is_system_file(path_or_ext):
        color, label = "#94a3b8", (ext.replace(".", "").upper()[:3] or "SYS")
    else:
        color, label = "#0f766e", (ext.replace(".", "").upper()[:3] or "FILE")

    if kind == "Каталог":
        svg = f"""
        <svg viewBox="0 0 56 56" aria-hidden="true">
          <path fill="#d99d2b" d="M5 15a6 6 0 0 1 6-6h12l5 6h17a6 6 0 0 1 6 6v4H5z"/>
          <path fill="{color}" d="M5 20h46v21a6 6 0 0 1-6 6H11a6 6 0 0 1-6-6z"/>
          <path fill="#ffd56d" opacity=".75" d="M8 23h40v4H8z"/>
        </svg>
        """
    else:
        font_size = "12" if len(label) <= 1 else "9"
        svg = f"""
        <svg viewBox="0 0 56 56" aria-hidden="true">
          <path fill="#ffffff" d="M12 4h22l10 10v36a4 4 0 0 1-4 4H12a4 4 0 0 1-4-4V8a4 4 0 0 1 4-4z"/>
          <path fill="#dbe4ef" d="M34 4v10h10z"/>
          <path fill="{color}" d="M6 29h37a4 4 0 0 1 4 4v12a4 4 0 0 1-4 4H6z"/>
          <text x="25" y="42" fill="#fff" text-anchor="middle" font-family="Arial, sans-serif" font-size="{font_size}" font-weight="700">{html.escape(label)}</text>
          <path fill="none" stroke="#cbd5e1" d="M12.5 4.5h21.3L43.5 14v35.5a4 4 0 0 1-4 4h-27a4 4 0 0 1-4-4v-41a4 4 0 0 1 4-4z"/>
        </svg>
        """
    icon_class = "rag-file-icon system" if kind != "Каталог" and _is_system_file(path_or_ext) else "rag-file-icon"
    return f'<span class="{icon_class}">{svg}</span>'


def _path_sort_key(path: Path, sort_by: str) -> Any:
    if sort_by == "По размеру":
        name = path.name.lower()
        if path.is_file():
            try:
                return (path.stat().st_size, name)
            except Exception:
                return (0, name)
        return (-1, name)
    if sort_by == "По дате":
        name = path.name.lower()
        try:
            return (path.stat().st_mtime, name)
        except Exception:
            return (0, name)
    return path.name.lower()


def _file_rows(path: Path, state: PageState) -> tuple[List[Path], List[Path], int]:
    try:
        entries = [x for x in path.iterdir() if not x.name.startswith(".") and not x.name.startswith("~$")]
    except Exception:
        return [], [], 0

    dirs = [x for x in entries if x.is_dir()]
    files = [x for x in entries if x.is_file()]
    if not _show_system_files(state):
        files = [x for x in files if not _is_system_file(x)]
    needle = state.explorer_filter.strip().lower()
    if needle:
        dirs = [x for x in dirs if needle in x.name.lower()]
        files = [x for x in files if needle in x.name.lower()]
    if state.explorer_ext != "Все":
        files = [x for x in files if x.suffix.lower() == state.explorer_ext.lower()]

    reverse = bool(state.explorer_desc)
    dirs.sort(key=lambda x: _path_sort_key(x, state.explorer_sort), reverse=reverse)
    files.sort(key=lambda x: _path_sort_key(x, state.explorer_sort), reverse=reverse)
    return dirs, files, len(files)


def _event_input_value(event: Any, fallback: Any = "") -> str:
    value = getattr(event, "value", None)
    if value is not None:
        return str(value)
    args = getattr(event, "args", None)
    if isinstance(args, dict):
        for key in ("value", "modelValue"):
            if args.get(key) is not None:
                return str(args[key])
        target = args.get("target")
        if isinstance(target, dict) and target.get("value") is not None:
            return str(target["value"])
    return str(fallback or "")


def _apply_explorer_filter_input(state: PageState, event: Any, fallback: Any = "") -> None:
    state.explorer_filter = _event_input_value(event, fallback)
    state.explorer_page = 0


def _install_css() -> None:
    ui.add_head_html(
        '<link href="https://fonts.googleapis.com/css2?family=Manrope:wght@400;500;600;700;800&family=Inter:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap" rel="stylesheet">'
    )
    ui.add_css(
        """
        :root {
          /* Brandbook palette */
          --rag-accent: #3d63ff;
          --rag-accent-2: #10b981;
          --rag-amber: #f59e0b;
          --rag-danger: #dc2626;
          --rag-ok: #16a34a;
          --rag-warn: #f59e0b;

          /* Backgrounds */
          --rag-bg: #fafaf7;
          --rag-surface: #ffffff;
          --rag-bg-sunken: #f4f3ee;

          /* Lines / borders */
          --rag-border: #e6e3da;
          --rag-border-strong: #d8d3c4;

          /* Ink / text */
          --rag-text: #14141a;
          --rag-ink-2: #3a3a44;
          --rag-muted: #6c6c78;
          --rag-ink-4: #9a9aa2;

          /* Fonts */
          --rag-display-font: 'Manrope', system-ui, sans-serif;
          --rag-text-font: 'Inter', system-ui, sans-serif;
          --rag-mono-font: 'JetBrains Mono', ui-monospace, monospace;

          /* Radii */
          --rag-radius-sm: 6px;
          --rag-radius-md: 10px;
          --rag-radius-lg: 14px;
          --rag-radius-xl: 22px;

          /* Shadows */
          --rag-shadow: 0 1px 0 rgba(20,20,26,.04), 0 1px 2px rgba(20,20,26,.05);
          --rag-shadow-2: 0 6px 16px -8px rgba(20,20,26,.12), 0 2px 4px rgba(20,20,26,.04);
          --rag-shadow-3: 0 24px 48px -16px rgba(20,20,26,.18), 0 8px 16px -8px rgba(20,20,26,.08);
        }
        .body--dark {
          --rag-bg: #0c0c0f;
          --rag-surface: #15151a;
          --rag-bg-sunken: #08080a;
          --rag-border: #23232b;
          --rag-border-strong: #2e2e38;
          --rag-text: #f4f4f7;
          --rag-ink-2: #c8c8d0;
          --rag-muted: #8a8a96;
          --rag-ink-4: #5a5a64;
          --rag-accent: #6385ff;
        }
        body {
          background: var(--rag-bg) !important;
          color: var(--rag-text) !important;
          font-family: var(--rag-text-font) !important;
          font-size: 87.5% !important;
          letter-spacing: -0.01em !important;
        }
        .q-page { background: transparent; }
        .rag-header {
          background: rgba(250, 250, 247, 0.88) !important;
          color: var(--rag-text) !important;
          border-bottom: 1px solid var(--rag-border) !important;
          backdrop-filter: blur(20px) saturate(160%);
          -webkit-backdrop-filter: blur(20px) saturate(160%);
          padding-left: 16px !important;
          padding-right: 16px !important;
        }
        .body--dark .rag-header {
          background: rgba(21, 21, 26, 0.92) !important;
          border-bottom-color: #23232b !important;
        }
        .rag-header-brand {
          display: flex;
          align-items: center;
          gap: 8px;
          text-decoration: none;
          cursor: default;
        }
        .rag-header-brand-name {
          font-family: var(--rag-display-font);
          font-size: 15px;
          font-weight: 700;
          color: var(--rag-text);
          letter-spacing: -0.02em;
        }
        .rag-header-breadcrumbs .q-btn { min-height: 32px; padding: 0 6px; }
        .rag-header-actions .q-btn { min-width: 34px; min-height: 34px; }

        /* ── Drawer — всегда тёмная боковая панель ── */
        .rag-drawer {
          background: #0d0d11 !important;
          border-right: none !important;
          width: 224px !important;
          padding: 0 !important;
        }
        .rag-drawer-brand {
          display: flex;
          align-items: center;
          gap: 10px;
          padding: 18px 16px 12px;
          border-bottom: 1px solid rgba(255,255,255,.07);
          margin-bottom: 8px;
        }
        .rag-drawer-brand-mark {
          width: 28px;
          height: 28px;
          flex-shrink: 0;
        }
        .rag-drawer-brand-name {
          font-family: var(--rag-display-font);
          font-size: 14px;
          font-weight: 700;
          color: #ffffff;
          letter-spacing: -0.02em;
        }
        .rag-drawer-body {
          display: flex;
          flex-direction: column;
          height: 100%;
          padding: 0 8px;
        }
        .rag-drawer-nav { flex: 1; }
        .rag-drawer-bottom {
          padding: 8px 0 12px;
          border-top: 1px solid rgba(255,255,255,.07);
          margin-top: 8px;
        }
        .rag-nav-button {
          color: rgba(255,255,255,.6) !important;
          border-radius: var(--rag-radius-md) !important;
          transition: background .12s ease, color .12s ease !important;
        }
        .rag-nav-button:hover {
          background: rgba(255,255,255,.07) !important;
          color: rgba(255,255,255,.9) !important;
        }
        .rag-nav-button.text-primary,
        .rag-nav-button[color=primary] {
          background: rgba(99,133,255,.18) !important;
          color: #a0b4ff !important;
        }
        .rag-nav-button .q-btn__content {
          justify-content: flex-start;
          width: 100%;
          text-align: left;
        }
        .rag-nav-button .q-icon { margin-right: 10px; }
        .rag-page {
          width: min(1440px, calc(100vw - 24px));
          margin: 0 auto;
          padding: 10px 0 32px;
        }
        .rag-page.search { padding-top: 4px; }
        .rag-title { font-size: clamp(22px, 3.5vw, 34px); font-weight: 760; line-height: 1.05; letter-spacing: 0; font-family: var(--rag-display-font); }
        .rag-subtitle { color: var(--rag-muted); font-size: 13px; max-width: 820px; }
        .rag-card {
          background: var(--rag-surface);
          border: 1px solid var(--rag-border);
          border-radius: var(--rag-radius-lg);
          box-shadow: var(--rag-shadow-2);
          transition: transform 0.2s ease, box-shadow 0.2s ease;
        }
        .rag-card:hover {
          transform: translateY(-2px);
          box-shadow: var(--rag-shadow-3);
        }
        .rag-search-shell { position: relative; z-index: 5; }
        .rag-search-box {
          background: var(--rag-surface);
          border: 1px solid var(--rag-border);
          border-radius: var(--rag-radius-lg);
          box-shadow: var(--rag-shadow-2);
          transition: box-shadow 0.2s ease;
        }
        .rag-search-box:focus-within {
          box-shadow: 0 0 0 2px var(--rag-accent), var(--rag-shadow-2);
        }
        .rag-suggest {
          position: absolute;
          left: 0;
          right: 0;
          top: calc(100% + 8px);
          background: var(--rag-surface);
          border: 1px solid var(--rag-border);
          border-radius: var(--rag-radius-lg);
          box-shadow: var(--rag-shadow-3);
          overflow: hidden;
          z-index: 30;
        }
        .rag-result {
          background: var(--rag-surface);
          border: 1px solid var(--rag-border);
          border-radius: var(--rag-radius-md);
          padding: 12px;
          box-shadow: var(--rag-shadow);
          width: 100%;
          box-sizing: border-box;
          transition: all 0.2s ease;
        }
        .rag-result:hover {
          background: var(--rag-surface);
          transform: translateY(-2px);
          box-shadow: var(--rag-shadow-2);
          border-color: rgba(61, 99, 255, 0.3);
        }
        .rag-meta { color: var(--rag-muted); font-size: 12px; }
        .rag-chip {
          display: inline-flex;
          align-items: center;
          min-height: 28px;
          padding: 0 12px;
          border: 1px solid var(--rag-border);
          border-radius: 14px;
          color: var(--rag-muted);
          background: var(--rag-bg-sunken);
          font-size: 12px;
          font-weight: 500;
          cursor: pointer;
          user-select: none;
          transition: all 0.2s ease;
        }
        .rag-chip:hover {
          background: var(--rag-surface);
          color: var(--rag-accent);
          border-color: var(--rag-accent);
          transform: translateY(-1px);
          box-shadow: 0 4px 6px -1px rgba(61, 99, 255, 0.1);
        }
        .rag-chip-active {
          background: var(--rag-accent);
          color: #ffffff;
          border-color: transparent;
          font-weight: 600;
          box-shadow: 0 4px 12px rgba(61, 99, 255, 0.3);
        }
        .rag-chip-active:hover {
          transform: translateY(-1px);
          box-shadow: 0 6px 14px rgba(61, 99, 255, 0.4);
        }
        .rag-path {
          word-break: break-word;
          overflow-wrap: anywhere;
          color: var(--rag-muted);
          font-size: 12px;
        }
        .rag-actions { display: flex; flex-wrap: wrap; gap: 6px; }
        .rag-feedback-btn {
          width: 30px;
          height: 30px;
          min-width: 30px;
        }
        .rag-nav-button { justify-content: flex-start; border-radius: 8px; text-align: left; }
        .rag-nav-button .q-btn__content { justify-content: flex-start; width: 100%; text-align: left; }
        .rag-nav-button .q-icon { margin-right: 10px; }
        .rag-suggest-item .block {
          min-width: 0;
          overflow: hidden;
          text-overflow: ellipsis;
          white-space: nowrap;
        }
        .rag-group-panel {
          width: 100%;
          border: 1px solid var(--rag-border);
          border-radius: var(--rag-radius-md);
          background: var(--rag-surface);
          overflow: hidden;
        }
        .rag-file-icon {
          display: inline-flex;
          width: 34px;
          height: 34px;
          flex: 0 0 34px;
        }
        .rag-file-icon svg { width: 34px; height: 34px; display: block; }
        .rag-file-icon.system {
          opacity: .42;
          filter: grayscale(1);
        }
        .rag-file-icon svg { width: 42px; height: 42px; display: block; }
        .rag-explorer-grid {
          display: grid;
          grid-template-columns: repeat(auto-fill, minmax(220px, 1fr));
          gap: 10px;
        }
        .rag-explorer-grid.medium { grid-template-columns: repeat(auto-fill, minmax(150px, 1fr)); }
        .rag-explorer-grid.small {
          grid-template-columns: repeat(auto-fill, minmax(82px, 92px));
          gap: 8px;
        }
        .rag-explorer-item {
          width: 100%;
          min-width: 0;
          background: var(--rag-surface);
          border: 1px solid var(--rag-border);
          border-radius: var(--rag-radius-md);
          color: var(--rag-text);
          transition: all 0.2s ease;
        }
        .rag-explorer-item:hover {
          background: var(--rag-surface);
          border-color: rgba(61, 99, 255, 0.3);
          transform: translateY(-2px);
          box-shadow: 0 10px 20px -10px rgba(61, 99, 255, 0.15);
        }
        .rag-explorer-item.system {
          opacity: .55;
          color: var(--rag-muted);
        }
        .rag-explorer-item.system:hover {
          opacity: .78;
          background: var(--rag-bg-sunken);
          border-color: var(--rag-border-strong);
        }
        .rag-explorer-item { position: relative; }
        .rag-explorer-grid.small .rag-explorer-item {
          min-height: 96px;
          max-height: 106px;
          padding: 6px;
          overflow: hidden;
        }
        .rag-explorer-grid.small .rag-file-icon,
        .rag-explorer-grid.small .rag-file-icon svg {
          width: 34px;
          height: 34px;
          flex-basis: 34px;
        }
        .rag-explorer-grid.small .rag-favorite-star {
          position: absolute;
          top: 2px;
          right: 2px;
          z-index: 2;
          background: rgba(255, 255, 255, 0.72);
        }
        .rag-explorer-grid.small .rag-explorer-opener {
          width: 100%;
          min-width: 0;
          overflow: hidden;
        }
        .rag-favorite-star {
          opacity: 0;
          color: rgba(0, 0, 0, 0.45);
          transition: opacity .12s ease, color .12s ease, transform .12s ease;
        }
        .rag-explorer-item:hover .rag-favorite-star,
        .rag-favorite-star.active {
          opacity: 1;
        }
        .rag-favorite-star:hover {
          color: #d89b00;
          transform: scale(1.08);
        }
        .rag-favorite-star.active {
          color: #f6b700;
        }
        .rag-tile-star-wrap {
          position: absolute;
          top: 4px;
          right: 4px;
          z-index: 2;
        }
        .rag-favorite-star.header {
          opacity: .65;
        }
        .rag-favorite-star.header:hover {
          opacity: 1;
          color: #d89b00;
        }
        .rag-bookmarks {
          display: flex;
          width: 100%;
          gap: 8px;
          overflow-x: auto;
          padding: 8px 0;
          align-items: center;
          flex-wrap: nowrap;
        }
        .rag-bookmark {
          position: relative;
          flex: 0 0 auto;
          width: 220px;
          min-width: 160px;
          height: 42px;
          border: 1px solid var(--rag-border);
          background: var(--rag-surface);
          border-radius: var(--rag-radius-md);
          overflow: hidden;
          transition: background .12s ease, border-color .12s ease, box-shadow .12s ease;
        }
        .rag-bookmark:hover {
          background: var(--rag-bg-sunken);
          border-color: var(--rag-border-strong);
        }
        .rag-bookmark-main {
          position: absolute;
          inset: 0 36px 0 0;
          display: flex;
          align-items: center;
          min-width: 0;
        }
        .rag-bookmark:hover .rag-bookmark-main {
          box-shadow: 18px 0 24px rgba(23, 32, 44, 0.12);
        }
        .rag-bookmark-remove {
          position: absolute;
          right: 0;
          top: 0;
          width: 36px;
          height: 100%;
          display: flex;
          align-items: center;
          justify-content: center;
          opacity: 0;
          background: var(--rag-surface);
          border-left: 1px solid var(--rag-border);
          color: var(--rag-muted);
          transition: opacity .12s ease, color .12s ease, background .12s ease;
        }
        .rag-bookmark:hover .rag-bookmark-remove {
          opacity: 1;
        }
        .rag-bookmark-remove:hover {
          background: #fff1f1;
          color: var(--rag-danger);
        }
        .rag-bookmark .q-btn {
          min-width: 0;
          width: 100%;
          height: 100%;
          padding-right: 4px;
        }
        .rag-bookmark .q-btn__content {
          min-width: 0;
          flex-wrap: nowrap;
          overflow: hidden;
        }
        .rag-bookmark .block {
          min-width: 0;
          overflow: hidden;
          text-overflow: ellipsis;
          white-space: nowrap;
        }
        .rag-bookmark-more {
          flex: 0 0 auto;
          width: 42px;
          height: 42px;
        }
        .rag-context-menu {
          position: fixed;
          z-index: 10000;
          min-width: 220px;
          background: var(--rag-surface);
          border: 1px solid var(--rag-border);
          border-radius: var(--rag-radius-md);
          box-shadow: var(--rag-shadow-3);
          padding: 6px;
          display: none;
        }
        .rag-context-menu button {
          display: block;
          width: 100%;
          padding: 8px 10px;
          border: 0;
          background: transparent;
          text-align: left;
          border-radius: var(--rag-radius-sm);
          color: var(--rag-text);
          cursor: pointer;
        }
        .rag-context-menu button:hover { background: var(--rag-bg-sunken); }
        .rag-favorites-dialog-row {
          display: grid;
          grid-template-columns: auto minmax(0, 1fr) auto;
          gap: 8px;
          align-items: center;
          width: 100%;
        }
        .rag-explorer-name {
          width: 100%;
          min-width: 0;
          overflow-wrap: anywhere;
          word-break: break-word;
          line-height: 1.2;
        }
        .rag-explorer-grid.small .rag-explorer-name {
          display: -webkit-box;
          width: 100%;
          max-width: 100%;
          -webkit-line-clamp: 2;
          -webkit-box-orient: vertical;
          overflow: hidden;
          text-overflow: ellipsis;
          overflow-wrap: anywhere;
          word-break: break-word;
          font-size: 12px;
          line-height: 1.15;
        }
        .rag-explorer-list {
          display: grid;
          grid-template-columns: 1fr;
          gap: 4px;
        }
        .rag-code {
          white-space: pre-wrap;
          word-break: break-word;
          font-family: var(--rag-mono-font);
          font-size: 12px;
          background: var(--rag-bg-sunken);
          border: 1px solid var(--rag-border);
          border-radius: var(--rag-radius-sm);
          padding: 12px;
        }
        @media (max-width: 760px) {
          .rag-page { width: calc(100vw - 20px); padding-top: 18px; }
          .rag-title { font-size: 28px; }
          .rag-actions .q-btn { width: auto; }
          .rag-search-box { box-shadow: var(--rag-shadow); }
        }

        /* ═══════════════════════════════════════
           LOGIN — split layout
        ═══════════════════════════════════════ */
        .rag-login-root {
          position: fixed;
          inset: 0;
          z-index: 9000;
          display: flex;
          background: var(--rag-bg);
        }
        .rag-login-brand {
          flex: 0 0 400px;
          background: #0c0c0f;
          display: flex;
          flex-direction: column;
          padding: 48px 40px 32px;
          position: relative;
          overflow: hidden;
        }
        .rag-login-brand::before {
          content: '';
          position: absolute;
          inset: 0;
          background:
            radial-gradient(ellipse at 25% 15%, rgba(61,99,255,.28) 0, transparent 55%),
            radial-gradient(ellipse at 75% 85%, rgba(99,133,255,.14) 0, transparent 55%);
          pointer-events: none;
        }
        .rag-login-brand-inner {
          position: relative;
          z-index: 1;
          display: flex;
          flex-direction: column;
          flex: 1;
        }
        .rag-login-mark-wrap {
          width: 56px;
          height: 56px;
          margin-bottom: 28px;
          flex-shrink: 0;
        }
        .rag-login-mark-wrap svg { width: 56px; height: 56px; }
        .rag-login-brand-name {
          font-family: var(--rag-display-font);
          font-size: 26px;
          font-weight: 800;
          color: #ffffff;
          letter-spacing: -0.02em;
          line-height: 1.1;
        }
        .rag-login-brand-sub {
          font-size: 13px;
          color: rgba(255,255,255,.45);
          margin-top: 6px;
          letter-spacing: 0.02em;
        }
        .rag-login-stats {
          margin-top: auto;
          padding-top: 48px;
        }
        .rag-login-stat-val {
          font-family: var(--rag-display-font);
          font-size: 44px;
          font-weight: 800;
          color: #ffffff;
          letter-spacing: -0.03em;
          line-height: 1;
        }
        .rag-login-stat-lbl {
          font-size: 13px;
          color: rgba(255,255,255,.45);
          margin-top: 6px;
        }
        .rag-login-brand-footer {
          position: relative;
          z-index: 1;
          margin-top: 32px;
          padding-top: 20px;
          border-top: 1px solid rgba(255,255,255,.1);
        }
        .rag-login-badge {
          display: inline-flex;
          align-items: center;
          gap: 6px;
          font-size: 12px;
          color: rgba(255,255,255,.4);
        }
        .rag-login-badge-dot {
          width: 6px;
          height: 6px;
          border-radius: 50%;
          background: #10b981;
          flex-shrink: 0;
        }
        .rag-login-form-side {
          flex: 1;
          display: flex;
          flex-direction: column;
          align-items: center;
          justify-content: center;
          padding: 48px 40px;
          overflow-y: auto;
          background: var(--rag-surface);
        }
        .rag-login-form-inner {
          width: 100%;
          max-width: 380px;
          display: flex;
          flex-direction: column;
          gap: 0;
        }
        .rag-login-greeting {
          font-family: var(--rag-display-font);
          font-size: 28px;
          font-weight: 700;
          color: var(--rag-text);
          letter-spacing: -0.02em;
          line-height: 1.15;
          margin-bottom: 6px;
        }
        .rag-login-greeting-sub {
          font-size: 14px;
          color: var(--rag-muted);
          margin-bottom: 28px;
        }
        .rag-login-divider {
          display: flex;
          align-items: center;
          gap: 12px;
          margin: 16px 0;
        }
        .rag-login-divider-line {
          flex: 1;
          height: 1px;
          background: var(--rag-border);
        }
        .rag-login-divider-text {
          font-size: 12px;
          color: var(--rag-ink-4);
          white-space: nowrap;
        }
        .rag-login-reg-wrap {
          margin-top: 20px;
          padding-top: 20px;
          border-top: 1px solid var(--rag-border);
          display: flex;
          flex-direction: column;
          gap: 12px;
        }
        .rag-login-reg-wrap .rag-login-reg-title {
          font-family: var(--rag-display-font);
          font-size: 16px;
          font-weight: 600;
          color: var(--rag-text);
        }
        .rag-login-input-gap { margin-bottom: 12px; }
        .rag-login-btn-primary {
          background: var(--rag-accent) !important;
          color: #ffffff !important;
          border-radius: var(--rag-radius-md) !important;
          font-weight: 600 !important;
          font-size: 14px !important;
          min-height: 44px !important;
          margin-top: 20px;
          margin-bottom: 12px;
        }
        @media (max-width: 800px) {
          .rag-login-brand { display: none; }
          .rag-login-form-side { padding: 32px 24px; }
        }

        /* ═══════════════════════════════════════
           SEARCH — stats bar, score badge, toolbar
        ═══════════════════════════════════════ */
        .rag-search-stats-bar {
          display: flex;
          align-items: center;
          gap: 10px;
          flex-wrap: wrap;
          padding: 8px 12px;
          background: var(--rag-bg-sunken);
          border: 1px solid var(--rag-border);
          border-radius: var(--rag-radius-md);
          font-size: 12px;
          color: var(--rag-muted);
        }
        .rag-search-stats-total {
          font-weight: 600;
          color: var(--rag-text);
          font-size: 13px;
        }
        .rag-search-stats-sep {
          width: 1px;
          height: 12px;
          background: var(--rag-border-strong);
          flex-shrink: 0;
        }
        .rag-sort-toolbar {
          display: flex;
          align-items: center;
          gap: 8px;
          flex-wrap: wrap;
        }
        .rag-sort-btn {
          display: inline-flex;
          align-items: center;
          gap: 4px;
          padding: 4px 10px;
          border-radius: var(--rag-radius-sm);
          font-size: 12px;
          font-weight: 500;
          color: var(--rag-muted);
          background: transparent;
          border: 1px solid var(--rag-border);
          cursor: pointer;
          transition: all 0.15s ease;
        }
        .rag-sort-btn:hover {
          color: var(--rag-text);
          border-color: var(--rag-border-strong);
        }
        .rag-sort-btn.active {
          color: var(--rag-accent);
          border-color: var(--rag-accent);
          background: rgba(61,99,255,.06);
        }
        .rag-score-badge {
          display: inline-flex;
          align-items: center;
          padding: 2px 7px;
          border-radius: 10px;
          font-size: 11px;
          font-weight: 600;
          font-family: var(--rag-mono-font);
          background: rgba(61,99,255,.08);
          color: rgba(61,99,255,.7);
          border: 1px solid rgba(61,99,255,.15);
          flex-shrink: 0;
        }
        .rag-result-date {
          font-size: 11px;
          color: var(--rag-ink-4);
          white-space: nowrap;
        }
        .rag-result-kind {
          font-size: 11px;
          color: var(--rag-muted);
          background: var(--rag-bg-sunken);
          border: 1px solid var(--rag-border);
          border-radius: 6px;
          padding: 1px 6px;
        }
        .rag-refine-bar {
          display: flex;
          align-items: center;
          gap: 8px;
          flex-wrap: wrap;
          padding: 8px 0 2px;
        }
        .rag-refine-label {
          font-size: 12px;
          color: var(--rag-muted);
          white-space: nowrap;
        }
        """
    )
    ui.add_body_html(
        """
        <div id="rag-global-context-menu" class="rag-context-menu" role="menu"></div>
        <script>
        (() => {
          if (window.__ragContextMenuInstalled) return;
          window.__ragContextMenuInstalled = true;
          const menu = () => document.getElementById('rag-global-context-menu');
          const hide = () => { const m = menu(); if (m) m.style.display = 'none'; };
          const addButton = (m, label, action) => {
            const b = document.createElement('button');
            b.textContent = label;
            b.onclick = () => { hide(); action(); };
            m.appendChild(b);
          };
          const show = (event) => {
            const root = event.target.closest('.q-layout');
            if (!root) return;
            event.preventDefault();
            const m = menu();
            if (!m) return;
            const item = event.target.closest('[data-rag-context="explorer-item"]');
            m.innerHTML = '';
            if (item) {
              const itemType = item.dataset.ragType || 'file';
              const itemPath = decodeURIComponent(item.dataset.ragPath || '');
              const itemUrl = item.dataset.ragUrl || '';
              addButton(m, 'Открыть', () => {
                item.querySelector('[data-rag-open]')?.click();
              });
              if (itemType === 'file' && itemUrl) {
                addButton(m, 'Скачать', () => {
                  const a = document.createElement('a');
                  a.href = itemUrl;
                  a.download = '';
                  document.body.appendChild(a);
                  a.click();
                  a.remove();
                });
              }
              addButton(m, 'Показать в ОС', () => item.querySelector('[data-rag-os]')?.click());
              addButton(m, item.dataset.ragFavorite === 'true' ? 'Убрать из избранного' : 'Добавить в избранное', () => item.querySelector('[data-rag-favorite-button]')?.click());
              addButton(m, 'Поделиться путем', () => navigator.clipboard && navigator.clipboard.writeText(itemPath));
            } else {
              addButton(m, 'Обновить экран', () => location.reload());
              addButton(m, 'Скопировать адрес экрана', () => navigator.clipboard && navigator.clipboard.writeText(location.href));
              addButton(m, 'Настройки', () => { location.href = '/settings'; });
            }
            m.style.left = Math.min(event.clientX, window.innerWidth - 240) + 'px';
            m.style.top = Math.min(event.clientY, window.innerHeight - 160) + 'px';
            m.style.display = 'block';
          };
          document.addEventListener('contextmenu', show);
          document.addEventListener('click', hide);
          document.addEventListener('scroll', hide, true);
          document.addEventListener('keydown', (e) => { if (e.key === 'Escape') hide(); });
        })();
        </script>
        """
    )
    # ── v2 wireframe CSS ────────────────────────────────────────────────
    ui.add_css(
        """
        /* === KPI tiles (Home + Analytics) === */
        .rag-kpi-tile {
          background: var(--rag-surface);
          border: 1px solid var(--rag-border);
          border-radius: var(--rag-radius-lg);
          padding: 20px 20px 14px;
          flex: 1;
          min-width: 160px;
          display: flex;
          flex-direction: column;
          gap: 4px;
          box-shadow: var(--rag-shadow);
        }
        .rag-kpi-label {
          font-family: var(--rag-mono-font);
          font-size: 10px;
          font-weight: 600;
          color: var(--rag-ink-4);
          text-transform: uppercase;
          letter-spacing: .08em;
        }
        .rag-kpi-value {
          font-family: var(--rag-display-font);
          font-size: 28px;
          font-weight: 800;
          letter-spacing: -0.03em;
          color: var(--rag-text);
          line-height: 1;
        }
        .rag-kpi-delta {
          font-size: 11px;
          font-weight: 600;
        }
        .rag-kpi-delta.up   { color: var(--rag-ok); }
        .rag-kpi-delta.down { color: var(--rag-danger); }
        .rag-kpi-delta.neutral { color: var(--rag-muted); }
        .rag-kpi-sparkline {
          display: flex;
          align-items: flex-end;
          gap: 2px;
          height: 28px;
          margin-top: 8px;
        }
        .rag-kpi-sparkline span {
          flex: 1;
          background: var(--rag-accent);
          opacity: .55;
          border-radius: 2px 2px 0 0;
          min-height: 2px;
          transition: opacity .15s;
        }
        .rag-kpi-sparkline span:last-child { opacity: .9; }

        /* === Refine bar (Search) === */
        .rag-refine-bar {
          background: rgba(245,158,11,.1);
          border: 1px solid rgba(245,158,11,.35);
          border-radius: var(--rag-radius-md);
          padding: 8px 14px;
          display: flex;
          align-items: center;
          gap: 10px;
          font-size: 13px;
        }
        .rag-refine-bar-label {
          font-size: 11px;
          font-weight: 700;
          color: var(--rag-muted);
          white-space: nowrap;
          text-transform: uppercase;
          letter-spacing: .05em;
        }
        .rag-refine-chip {
          display: inline-flex;
          align-items: center;
          gap: 4px;
          padding: 2px 8px 2px 10px;
          background: var(--rag-accent);
          color: #fff;
          font-size: 11px;
          font-weight: 600;
          border-radius: 12px;
          cursor: pointer;
        }
        .rag-refine-chip:hover { opacity: .85; }

        /* === Selection bar (Search) === */
        .rag-selection-bar {
          position: sticky;
          bottom: 16px;
          z-index: 20;
          background: var(--rag-text);
          color: #fff;
          border-radius: var(--rag-radius-lg);
          padding: 10px 16px;
          display: flex;
          align-items: center;
          gap: 10px;
          box-shadow: var(--rag-shadow-3);
          font-size: 13px;
          font-weight: 600;
        }
        .rag-selection-bar .q-btn { color: #fff !important; }

        /* === Phase row (Index screen) === */
        .rag-phase-row {
          background: var(--rag-surface);
          border: 1px solid var(--rag-border);
          border-radius: var(--rag-radius-md);
          padding: 14px 16px;
          display: flex;
          align-items: flex-start;
          gap: 14px;
        }
        .rag-phase-row.running {
          border-color: var(--rag-accent);
          box-shadow: 0 0 0 1px var(--rag-accent)22;
        }
        .rag-phase-circle {
          width: 36px; height: 36px;
          border-radius: 50%;
          display: grid;
          place-items: center;
          font-size: 16px;
          font-weight: 700;
          flex-shrink: 0;
          color: #fff;
        }
        .rag-freshness-bar {
          height: 6px;
          border-radius: 3px;
          width: 180px;
          background: var(--rag-border);
          overflow: hidden;
          flex-shrink: 0;
        }
        .rag-freshness-bar-fill {
          height: 100%;
          border-radius: 3px;
          transition: width .3s ease;
        }

        /* === Preview hint (Search) === */
        .rag-preview-hint {
          border: 1px dashed var(--rag-border-strong);
          border-radius: var(--rag-radius-md);
          padding: 10px 16px;
          text-align: center;
          font-size: 12px;
          color: var(--rag-muted);
          cursor: pointer;
        }
        .rag-preview-hint:hover {
          border-color: var(--rag-accent);
          color: var(--rag-accent);
        }

        /* === Home screen === */
        .rag-home-greeting {
          font-family: var(--rag-display-font);
          font-size: 28px;
          font-weight: 700;
          letter-spacing: -0.025em;
          color: var(--rag-text);
        }
        .rag-home-date {
          font-size: 13px;
          color: var(--rag-muted);
          margin-top: 2px;
        }
        .rag-spoiler {
          background: var(--rag-surface);
          border: 1px solid var(--rag-border);
          border-radius: var(--rag-radius-lg);
          overflow: hidden;
        }
        .rag-spoiler-head {
          display: flex;
          align-items: center;
          padding: 10px 16px;
          cursor: pointer;
          user-select: none;
          border-bottom: 1px solid var(--rag-border);
          gap: 8px;
        }
        .rag-spoiler-head:hover { background: var(--rag-bg-sunken); }
        .rag-spoiler-title {
          font-family: var(--rag-display-font);
          font-size: 14px;
          font-weight: 600;
          color: var(--rag-text);
          flex: 1;
        }
        .rag-spoiler-count {
          font-family: var(--rag-mono-font);
          font-size: 11px;
          color: var(--rag-ink-4);
        }
        .rag-spoiler-body { padding: 16px; }

        /* === Analytics sub-tabs === */
        .rag-analytics-tabs {
          display: flex;
          gap: 4px;
          padding: 0;
          flex-wrap: wrap;
        }
        .rag-analytics-tab {
          padding: 5px 12px;
          font-size: 12px;
          font-weight: 500;
          border-radius: 6px;
          cursor: pointer;
          color: var(--rag-muted);
          border: 1px solid transparent;
          transition: all .15s;
        }
        .rag-analytics-tab:hover {
          background: var(--rag-bg-sunken);
          color: var(--rag-text);
        }
        .rag-analytics-tab.active {
          background: var(--rag-accent);
          color: #fff;
          border-color: transparent;
        }

        /* === AI insight panel === */
        .rag-ai-insight {
          background: rgba(245,158,11,.08);
          border: 1px solid rgba(245,158,11,.3);
          border-radius: var(--rag-radius-md);
          padding: 10px 16px;
          font-size: 13px;
          display: flex;
          align-items: flex-start;
          gap: 10px;
        }
        .rag-ai-insight-icon {
          font-size: 16px;
          flex-shrink: 0;
          margin-top: 1px;
        }

        /* === Index pipeline visual === */
        .rag-pipeline {
          display: flex;
          align-items: center;
          gap: 0;
          flex-wrap: wrap;
        }
        .rag-pipeline-step {
          display: flex;
          flex-direction: column;
          align-items: center;
          gap: 4px;
        }
        .rag-pipeline-dot {
          width: 28px; height: 28px;
          border-radius: 50%;
          display: grid;
          place-items: center;
          font-size: 12px;
          font-weight: 700;
          color: #fff;
        }
        .rag-pipeline-connector {
          width: 32px;
          height: 2px;
          background: var(--rag-border);
          flex-shrink: 0;
          align-self: center;
          margin-bottom: 20px;
        }
        .rag-pipeline-label {
          font-size: 10px;
          color: var(--rag-muted);
          font-family: var(--rag-mono-font);
          white-space: nowrap;
        }

        /* === Stats spoiler (collapsible stats in search) === */
        .rag-stats-spoiler-head {
          display: flex;
          align-items: center;
          gap: 10px;
          padding: 8px 14px;
          cursor: pointer;
          border-radius: var(--rag-radius-md);
          border: 1px solid var(--rag-border);
          background: var(--rag-surface);
          user-select: none;
        }
        .rag-stats-spoiler-head:hover { border-color: var(--rag-accent); }
        """
    )


def _build_page(initial_screen: str = "search") -> None:
    state = PageState(cfg=load_config())
    state.screen = initial_screen
    state.explorer_path = str(Path(str(state.cfg.get("catalog_path") or "")))
    _install_css()
    try:
        stored_token = str(app.storage.user.get("auth_token") or "")
        if stored_token:
            state.auth_token = stored_token
            state.current_user = _get_auth_db(state).get_user_by_session(stored_token)
            if state.current_user:
                _load_user_state(state)
                _get_auth_db(state).log_auth_event(username=_username(state), event_type="session_restore", ok=True)
    except Exception:
        pass

    dark_mode = ui.dark_mode(state.theme == "dark")

    with ui.header(fixed=True, elevated=False).classes("rag-header h-12 items-center"):
        # Гамбургер — только на узких экранах
        menu_button = (
            ui.button(icon="menu", on_click=lambda: drawer.toggle(), color=None)
            .props("flat round dense")
            .classes("lg:hidden")
            .style("color: var(--rag-muted);")
        )
        header_title = ui.label("").classes("hidden")
        header_breadcrumbs = ui.row().classes("rag-header-breadcrumbs items-center gap-1")
        header_actions = ui.row().classes("rag-header-actions items-center gap-1")
        state.header_breadcrumbs = header_breadcrumbs
        state.header_explorer_actions = header_actions
        ui.space()
        status_connected = _ensure_searcher(state) and state.searcher and state.searcher.connected
        status_text = "Qdrant" if status_connected else "Qdrant ✕"
        ui.label(status_text).classes("hidden sm:block rag-chip").style(
            f"color: {'var(--rag-ok)' if status_connected else 'var(--rag-danger)'};"
        )
        theme_button = ui.button(
            icon="light_mode" if state.theme == "dark" else "dark_mode",
            on_click=lambda: toggle_theme(),
            color=None,
        ).props("flat round dense").style("color: var(--rag-muted);")

    with (
        ui.left_drawer(value=True, fixed=True, bordered=False)
        .props("behavior=desktop breakpoint=1024")
        .classes("rag-drawer")
    ) as drawer:
        # Бренд-шапка
        ui.html(
            '<div class="rag-drawer-brand">'
            '  <img src="/rag-mark.svg" class="rag-drawer-brand-mark" alt="">'
            '  <span class="rag-drawer-brand-name">ТЕХНОПОИСК</span>'
            '</div>'
        )
        with ui.column().classes("rag-drawer-body"):
            nav_area = ui.column().classes("rag-drawer-nav w-full gap-1")
            settings_area = ui.column().classes("w-full mt-2")
            bottom_nav_area = ui.column().classes("rag-drawer-bottom w-full gap-1")

    page_root = ui.column().classes("rag-page gap-5")
    with page_root:
        content = ui.column().classes("w-full gap-5")

    def touch_activity() -> None:
        if not state.auth_token or not state.current_user:
            return
        try:
            _get_auth_db(state).touch_session(state.auth_token, min_interval_minutes=60)
        except Exception:
            pass

    ui.timer(3600.0, touch_activity)

    # ── Встроенный планировщик ────────────────────────────────────────
    def _tick_scheduler() -> None:
        """Каждую минуту проверяет расписание и запускает индексатор при совпадении."""
        if not _is_admin(state):
            return
        tdb = _get_telemetry(state)
        if not hasattr(tdb, "list_index_schedules"):
            return
        try:
            schedules = tdb.list_index_schedules()
            due = _schedules_due(schedules)
            cfg_settings = tdb.get_index_settings() if hasattr(tdb, "get_index_settings") else {}
            for sched in due:
                try:
                    pid = _launch_indexer(
                        state.cfg,
                        stage=str(sched.get("stage") or "all"),
                        workers=int(cfg_settings.get("workers") or state.cfg.get("index_read_workers") or 4),
                        max_chunks=int(cfg_settings.get("max_chunks") or state.cfg.get("index_max_chunks") or 2000),
                        skip_inline_ocr=bool(cfg_settings.get("skip_inline_ocr")),
                    )
                except RuntimeError as exc:
                    _log_app_event(
                        state,
                        "index",
                        "schedule_skip",
                        ok=False,
                        details={"sched_id": sched.get("id"), "reason": str(exc)},
                    )
                    continue
                tdb.touch_index_schedule(id=str(sched["id"]))
                _log_app_event(state, "index", "schedule_run",
                               details={"sched_id": sched["id"], "stage": sched.get("stage"), "pid": pid})
        except Exception:
            pass

    ui.timer(60.0, _tick_scheduler)

    def do_logout() -> None:
        auth_db = _get_auth_db(state)
        if state.auth_token:
            auth_db.revoke_session(state.auth_token)
        auth_db.log_auth_event(username=_username(state), event_type="logout", ok=True)
        state.current_user = None
        state.auth_token = ""
        state.theme = "light"
        dark_mode.set_value(False)
        try:
            app.storage.user.pop("auth_token", None)
        except Exception:
            pass
        render()

    def toggle_theme() -> None:
        if state.current_user is None:
            return
        state.theme = "dark" if state.theme == "light" else "light"
        dark_mode.set_value(state.theme == "dark")
        theme_button.set_icon("light_mode" if state.theme == "dark" else "dark_mode")
        _save_ui_settings(state)
        _log_app_event(state, "ui", "theme_toggle", details={"theme": state.theme})

    def set_screen(screen: str, *, close_drawer: bool = False) -> None:
        touch_activity()
        if close_drawer:
            try:
                drawer.set_value(False)
            except Exception:
                pass
        state.screen = screen
        ui.run_javascript(f"history.pushState(null, '', '/{screen}')")
        _log_app_event(state, "navigation", "open_screen", details={"screen": screen})
        render()

    def go_explorer(path: str) -> None:
        value = str(path or "").strip()
        if value:
            p = Path(value)
            state.explorer_path = str(p.parent if p.is_file() else p)
            state.explorer_page = 0
        set_screen("explorer")

    def update_nav() -> None:
        nav_area.clear()
        with nav_area:
            # "Главная" visible to all users (first item)
            color = "primary" if state.screen == "home" else None
            ui.button("Главная", icon="home", on_click=lambda: set_screen("home"), color=color).props("flat align=left no-caps").classes("rag-nav-button w-full")
            for screen, label, icon in [
                ("search", "Поиск", "search"),
                ("explorer", "Проводник", "folder"),
                ("telegram", "Telegram", "send"),
            ]:
                color = "primary" if state.screen == screen else None
                ui.button(label, icon=icon, on_click=lambda s=screen: set_screen(s), color=color).props("flat align=left no-caps").classes("rag-nav-button w-full")
            if str((state.current_user or {}).get("role") or "") == "admin":
                color = "primary" if state.screen == "index" else None
                ui.button("Индекс", icon="analytics", on_click=lambda: set_screen("index"), color=color).props("flat align=left no-caps").classes("rag-nav-button w-full")
                color = "primary" if state.screen == "stats" else None
                ui.button("Аналитика", icon="query_stats", on_click=lambda: set_screen("stats"), color=color).props("flat align=left no-caps").classes("rag-nav-button w-full")

        settings_area.clear()
        with settings_area:
            if state.screen == "search":
                ui.separator()

                def set_content_only(enabled: bool) -> None:
                    state.content_only = bool(enabled)
                    if state.content_only:
                        state.title_only = False
                    render()

                def set_title_only(enabled: bool) -> None:
                    state.title_only = bool(enabled)
                    if state.title_only:
                        state.content_only = False
                    render()

                with ui.expansion("Параметры поиска", icon="tune", value=True).classes("w-full"):
                    ui.select(
                        ["Все", ".docx", ".xlsx", ".xls", ".pdf"],
                        label="Тип файла",
                        value=state.file_type or "Все",
                        on_change=lambda e: setattr(state, "file_type", None if e.value == "Все" else e.value),
                    ).classes("w-full")
                    ui.number("Лимит", value=state.limit, min=1, max=50, step=1, on_change=lambda e: setattr(state, "limit", int(e.value or 10))).classes("w-full")
                    ui.checkbox(
                        "Искать только в содержимом",
                        value=state.content_only,
                        on_change=lambda e: set_content_only(bool(e.value)),
                    )
                    ui.checkbox(
                        "Искать только в названиях",
                        value=state.title_only,
                        on_change=lambda e: set_title_only(bool(e.value)),
                    )

        bottom_nav_area.clear()
        with bottom_nav_area:
            color = "primary" if state.screen == "settings" else None
            user_label = "Настройки"
            if state.current_user:
                user_label = f"Настройки · {state.current_user.get('username')}"
            ui.button(user_label, icon="settings", on_click=lambda: set_screen("settings"), color=color).props("flat align=left no-caps").classes("rag-nav-button w-full")
            if state.current_user:
                ui.button("Выйти", icon="logout", on_click=do_logout, color=None).props("flat align=left no-caps").classes("rag-nav-button w-full")

    async def run_search(explicit_query: Optional[str] = None) -> None:
        touch_activity()
        raw = explicit_query if explicit_query is not None else state.query
        query = re.sub(r"\s+", " ", str(raw or "")).strip()
        if not query:
            ui.notify("Введите запрос.", type="warning")
            return
        state.query = query
        state.search_error = ""
        state.results = []
        state.searched_query = query
        state.expanded_query = ""
        state.rag_answer_text = ""
        state.rag_answer_loading = False
        state.displayed_count = 10
        state.active_type_filter = None
        _remember_query(state, query)
        render_results_loading()
        searcher = _ensure_searcher(state)
        if searcher is None or not searcher.connected:
            state.search_error = state.searcher_error or "Нет подключения к Qdrant."
            render()
            return

        llm_enabled = bool(state.cfg.get("llm_enabled"))
        ollama_url = str(state.cfg.get("ollama_url") or "http://localhost:11434")
        expand_model = str(state.cfg.get("llm_expand_model") or "phi3:mini")
        rag_model = str(state.cfg.get("llm_rag_model") or "qwen3:8b")

        # Расширение запроса через LLM (если включено)
        search_query = query
        if llm_enabled:
            try:
                from rag_catalog.core.llm import expand_query  # noqa: PLC0415
                expanded = await run.io_bound(
                    expand_query, query, model=expand_model, ollama_url=ollama_url
                )
                if expanded and expanded.lower() != query.lower():
                    state.expanded_query = expanded
                    search_query = expanded
            except Exception:
                pass  # падаем молча — поиск всё равно идёт с оригинальным запросом

        try:
            state.results = await run.io_bound(
                _run_catalog_search,
                searcher,
                limit=state.limit,
                file_type=state.file_type,
                content_only=state.content_only,
                title_only=state.title_only,
                username=_username(state),
                query=search_query,
                query_original=query,
                query_used=search_query,
            )
            _log_app_event(
                state,
                "search",
                "run",
                details={
                    "query": query,
                    "query_used": search_query,
                    "results": len(state.results),
                    "content_only": bool(state.content_only),
                    "title_only": bool(state.title_only),
                },
            )
        except Exception as exc:
            state.search_error = str(exc)
            _log_app_event(
                state,
                "search",
                "run",
                ok=False,
                details={
                    "query": query,
                    "query_used": search_query,
                    "error": str(exc),
                    "content_only": bool(state.content_only),
                    "title_only": bool(state.title_only),
                },
            )

        render()

        # RAG Q&A — запускаем ПОСЛЕ рендера результатов, чтобы UI не ждал LLM
        if llm_enabled and state.results and not state.search_error:
            state.rag_answer_loading = True
            render()
            try:
                from rag_catalog.core.llm import rag_answer  # noqa: PLC0415
                answer = await run.io_bound(
                    rag_answer, query, state.results, model=rag_model, ollama_url=ollama_url
                )
                state.rag_answer_text = answer or ""
            except Exception as exc:
                state.rag_answer_text = f"Ошибка LLM: {exc}"
            state.rag_answer_loading = False
            render()

    async def choose_query(query: str) -> None:
        # Прямой async-обработчик: пресеты больше не зависят от ui.timer и гонок с перерисовкой.
        await run_search(query)

    def choose_query_handler(query: str) -> Any:
        async def handler() -> None:
            await choose_query(query)

        return handler

    def render_suggestions(area: ui.column, typed: str) -> None:
        area.clear()
        username = _username(state)
        personal = _dedupe_queries([*state.history, *_my_recent_queries(state.cfg, username, limit=12)], limit=12)
        popular = _popular_queries(state.cfg, exclude_username=username, limit=10)

        needle = typed.strip().lower()
        if needle:
            personal = [q for q in personal if needle in q.lower()]
            popular = [q for q in popular if needle in q.lower()]
        else:
            personal = personal[:8]
            popular = popular[:8]

        if not personal and not popular:
            return

        with area:
            with ui.row().classes("rag-suggest p-3 gap-0 w-full"):
                # Левая колонка — личная история
                if personal:
                    col_cls = "flex-1 gap-1 min-w-0" + (" pr-3 border-r border-gray-200" if popular else "")
                    with ui.column().classes(col_cls):
                        ui.label("Моя история").classes("rag-meta px-2 py-1 font-semibold text-xs uppercase tracking-wide")
                        for item in personal:
                            btn = ui.button(item, icon="history", on_click=choose_query_handler(item), color=None).props("flat align=left no-caps").classes("rag-nav-button rag-suggest-item w-full")
                            btn.tooltip(item)
                # Правая колонка — часто ищут
                if popular:
                    col_cls = "flex-1 gap-1 min-w-0" + (" pl-3" if personal else "")
                    with ui.column().classes(col_cls):
                        ui.label("Часто ищут").classes("rag-meta px-2 py-1 font-semibold text-xs uppercase tracking-wide")
                        for item in popular:
                            btn = ui.button(item, icon="trending_up", on_click=choose_query_handler(item), color=None).props("flat align=left no-caps").classes("rag-nav-button rag-suggest-item w-full")
                            btn.tooltip(item)

    def render_search_box() -> None:
        with ui.column().classes("rag-search-shell w-full max-w-5xl"):
            suggest_area = ui.column().classes("w-full")
            with ui.row().classes("rag-search-box w-full items-center gap-2 p-2"):
                search_input = ui.input(
                    placeholder="Введите название, номер, контрагента или фразу из документа",
                    value=state.query,
                    autocomplete=_search_suggestions(state),
                ).props("borderless dense clearable input-class=text-base").classes("flex-1")

                async def submit_click() -> None:
                    await run_search(str(search_input.value or ""))

                ui.button(icon="search", on_click=submit_click, color="primary").props("unelevated round")

            def handle_input(_: events.GenericEventArguments | None = None) -> None:
                state.query = str(search_input.value or "")
                render_suggestions(suggest_area, state.query)

            async def submit_from_input(_: events.GenericEventArguments | None = None) -> None:
                typed = str(search_input.value or "")
                suggest_area.clear()
                await run_search(typed)

            search_input.on("focus", handle_input)
            search_input.on("input", handle_input)
            search_input.on("keyup.enter", submit_from_input)

    def render_results_loading() -> None:
        content.clear()
        with content:
            render_search_header()
            ui.spinner(size="lg").classes("mt-4")
            ui.label("Ищу совпадения...").classes("rag-meta")

    def render_search_header() -> None:
        with ui.column().classes("w-full gap-2"):
            render_search_box()

    def open_file_viewer(path_value: Path | str) -> None:
        candidate = _resolve_catalog_file(state.cfg, str(path_value or ""))
        if candidate is None:
            ui.notify("Файл недоступен для просмотра.", type="warning")
            return
        viewer_url = _viewer_file_url(str(candidate))
        ext = candidate.suffix.lower()

        with ui.dialog() as dialog, ui.card().classes("w-[min(1100px,96vw)] max-h-[90vh] overflow-auto gap-3"):
            with ui.row().classes("w-full items-center gap-2"):
                with ui.column().classes("min-w-0 flex-1 gap-0"):
                    ui.label(candidate.name).classes("text-lg font-semibold truncate")
                    ui.label(str(candidate)).classes("rag-path")
                ui.button("Скачать", icon="download", on_click=lambda p=candidate: ui.download(p, filename=p.name)).props("outline dense")
                ui.button("Открыть в ОС", icon="open_in_new", on_click=lambda p=candidate: _open_os_path(str(p.parent))).props("outline dense")
                ui.button(icon="close", on_click=dialog.close, color=None).props("flat round dense")

            if ext == ".pdf":
                ui.html(
                    f'<iframe src="{html.escape(viewer_url, quote=True)}" '
                    'style="width:100%; height:72vh; border:1px solid rgba(148,163,184,.45); border-radius:10px;"></iframe>',
                    sanitize=False,
                )
            elif ext in INLINE_IMAGE_EXTENSIONS:
                ui.image(viewer_url).classes("max-w-full max-h-[72vh] object-contain mx-auto")
            elif ext in FILE_PREVIEW_EXTENSIONS:
                ui.label(_preview_file(candidate, limit=32000)).classes("rag-code")
            elif ext in OFFICE_PREVIEW_EXTENSIONS:
                ui.label(_preview_office_file(candidate, limit=32000)).classes("rag-code")
                ui.label("Для офисных форматов показывается текстовый извлеченный фрагмент.").classes("rag-meta")
            else:
                ui.label("Встроенный просмотр для этого формата не поддерживается. Используйте скачивание или открытие в ОС.").classes("rag-meta")
        dialog.open()

    def render_result(result: Dict[str, Any], index: int) -> None:
        name = str(result.get("filename") or "Без имени")
        path = str(result.get("path") or "")
        full_path = str(result.get("full_path") or "")
        score = float(result.get("score") or 0)
        kind = _result_kind(result)
        text = _clean_text(result.get("text") or "")
        preview = text[:280] + ("..." if len(text) > 280 else "")
        p = Path(full_path) if full_path else None

        def rate_result(value: int, result: Dict[str, Any] = result, index: int = index) -> None:
            result_path = str(result.get("full_path") or result.get("path") or "")
            _get_telemetry(state).log_search_feedback(
                username=_username(state),
                source="nicegui",
                query=state.searched_query,
                result_path=result_path,
                result_title=str(result.get("filename") or result_path),
                feedback=value,
                result_rank=index,
                result_score=float(result.get("score") or 0),
                details={"screen": "search", "reason": "explicit"},
            )
            _log_app_event(
                state,
                "search",
                "feedback",
                details={"value": value, "path": result_path, "query": state.searched_query},
            )
            ui.notify("Оценка сохранена.", type="positive")

        def track_result_use(reason: str, result: Dict[str, Any] = result, index: int = index) -> None:
            result_path = str(result.get("full_path") or result.get("path") or "")
            try:
                _get_telemetry(state).log_search_feedback(
                    username=_username(state),
                    source="nicegui",
                    query=state.searched_query,
                    result_path=result_path,
                    result_title=str(result.get("filename") or result_path),
                    feedback=2,
                    result_rank=index,
                    result_score=float(result.get("score") or 0),
                    details={"screen": "search", "reason": reason},
                )
            except Exception:
                pass

        def open_primary() -> None:
            if kind == "Каталог":
                track_result_use("open_folder")
                go_explorer(full_path)
                return
            if p and p.exists() and p.is_file():
                track_result_use("open_viewer")
                open_file_viewer(p)

        modified_raw = result.get("modified") or ""
        modified_str = ""
        if modified_raw:
            try:
                _dt = modified_raw[:10]  # "YYYY-MM-DD"
                modified_str = _dt
            except Exception:
                pass

        with ui.column().classes("rag-result gap-2"):
            with ui.row().classes("w-full items-start gap-2"):
                opener = ui.row().classes("flex-1 min-w-0 items-start gap-2 cursor-pointer").on("click", open_primary)
                with opener:
                    ui.html(_file_icon_svg(full_path or path, kind), sanitize=False)
                    with ui.column().classes("flex-1 min-w-0 gap-0"):
                        title = ui.label(f"{index}. {name}").classes("text-base font-semibold truncate")
                        title.tooltip(name)
                        path_label = ui.label(path or full_path).classes("rag-path truncate")
                        path_label.tooltip(path or full_path)
                with ui.column().classes("items-end gap-1 flex-shrink-0"):
                    ui.html(f'<span class="rag-score-badge">{score:.3f}</span>', sanitize=False)
                    if modified_str:
                        ui.html(f'<span class="rag-result-date">{modified_str}</span>', sanitize=False)
                    ui.html(f'<span class="rag-result-kind">{kind}</span>', sanitize=False)

            with ui.row().classes("w-full items-center justify-between gap-2"):
                with ui.row().classes("rag-actions items-center"):
                    if full_path:
                        if kind == "Каталог":
                            ui.button("В проводник приложения", icon="folder_open", on_click=lambda p=full_path: go_explorer(p)).props("outline dense")
                        else:
                            def open_in_app_explorer(pth: str = full_path) -> None:
                                track_result_use("open_in_app_explorer")
                                go_explorer(pth)

                            ui.button("В проводник приложения", icon="folder", on_click=open_in_app_explorer).props("outline dense")
                            if p and p.exists() and p.is_file():
                                ui.button("Скачать", icon="download", on_click=lambda pth=p: ui.download(pth, filename=pth.name)).props("outline dense")
                        ui.button("Открыть в ОС", icon="open_in_new", on_click=lambda pth=full_path: _open_os_path(str(Path(pth).parent if kind != "Каталог" else pth))).props("outline dense")
                with ui.row().classes("items-center justify-end gap-1"):
                    bad = ui.button(icon="thumb_down", on_click=lambda: rate_result(-3), color=None).props("flat round dense")
                    bad.classes("rag-feedback-btn")
                    bad.tooltip("Не то")
                    good = ui.button(icon="thumb_up", on_click=lambda: rate_result(3), color=None).props("flat round dense")
                    good.classes("rag-feedback-btn")
                    good.tooltip("Полезно")

            if kind == "Каталог":
                with ui.expansion("Раскрыть каталог", icon="account_tree").classes("w-full"):
                    children = _directory_children(full_path)
                    if not children["exists"]:
                        ui.label("Каталог недоступен на диске.").classes("rag-meta")
                    elif children.get("error"):
                        ui.label(f"Не удалось прочитать каталог: {children['error']}").classes("text-red-700")
                    else:
                        if children["dirs"]:
                            ui.label("Папки").classes("font-semibold")
                            with ui.column().classes("w-full gap-1"):
                                for item in children["dirs"]:
                                    btn = ui.button(item["name"], icon="folder", on_click=lambda pth=item["path"]: go_explorer(pth), color=None).props("flat align=left no-caps").classes("rag-nav-button w-full")
                                    btn.tooltip(str(item["path"]))
                        if children["files"]:
                            ui.label("Файлы").classes("font-semibold mt-2")
                            with ui.column().classes("w-full gap-1"):
                                for item in children["files"]:
                                    item_path = Path(str(item["path"]))
                                    with ui.row().classes("w-full items-center gap-2"):
                                        ui.html(_file_icon_svg(str(item_path), "Файл"), sanitize=False)
                                        file_btn = ui.button(
                                            f"{item['name']} · {item.get('size', '')}",
                                            on_click=lambda pth=item_path: open_file_viewer(pth),
                                            color=None,
                                        ).props("flat align=left no-caps dense").classes("rag-nav-button flex-1")
                                        file_btn.tooltip(str(item_path))
                        if children.get("truncated"):
                            ui.label("Показаны первые элементы. Полный список доступен в проводнике приложения.").classes("rag-meta")
            elif preview:
                ui.label(preview).classes("rag-meta")

    def render_search_screen() -> None:
        render_search_header()
        if state.search_error:
            ui.label(state.search_error).classes("text-red-700 rag-card p-4")
        if not state.searched_query:
            # ── Главный экран (нет активного запроса) ───────────────────────
            display_name = str(
                (state.current_user or {}).get("display_name")
                or (state.current_user or {}).get("username")
                or ""
            )
            greeting = f"Добрый день, {display_name}!" if display_name else "Добрый день!"

            _idx_stats = _read_index_stats(state.cfg)
            _doc_total = _idx_stats.get("total", 0)
            _doc_total_str = f"{_doc_total:,}".replace(",", " ")
            _last_mod = _idx_stats.get("last_modified") or "—"
            _by_ext = _idx_stats.get("by_ext") or {}
            _top_ext = sorted(_by_ext.items(), key=lambda x: x[1], reverse=True)[:6]

            # Приветствие
            with ui.column().classes("w-full gap-0 mb-2"):
                ui.html(
                    f'<div style="font-family: var(--rag-display-font); font-size: 24px;'
                    f' font-weight: 700; color: var(--rag-text); letter-spacing: -0.02em;">'
                    f'{greeting}</div>'
                    f'<div style="font-size: 13px; color: var(--rag-muted); margin-top: 4px;">'
                    f'Что ищем сегодня?</div>',
                    sanitize=False,
                )

            # Блок статистики (раскрыт по умолчанию)
            with ui.expansion("Статистика индекса", icon="query_stats", value=True).classes(
                "rag-card w-full"
            ):
                with ui.row().classes("w-full gap-3 flex-wrap p-1"):
                    with ui.column().classes("gap-0"):
                        ui.html(
                            f'<div style="font-family: var(--rag-display-font); font-size: 28px;'
                            f' font-weight: 800; color: var(--rag-text); letter-spacing: -0.03em;'
                            f' line-height: 1;">{_doc_total_str}</div>'
                            f'<div style="font-size: 12px; color: var(--rag-muted); margin-top: 2px;">документов</div>',
                            sanitize=False,
                        )
                    if _top_ext:
                        ui.element("div").style(
                            "width: 1px; background: var(--rag-border); align-self: stretch;"
                        )
                        with ui.column().classes("gap-1"):
                            for ext, cnt in _top_ext:
                                ui.html(
                                    f'<div style="font-size: 12px; color: var(--rag-muted);">'
                                    f'<span style="font-weight: 600; color: var(--rag-text);">'
                                    f'{ext.upper()}</span> — {cnt:,}'.replace(",", " ") + "</div>",
                                    sanitize=False,
                                )
                    ui.element("div").style("flex: 1;")
                    with ui.column().classes("gap-0 items-end"):
                        ui.html(
                            f'<div style="font-size: 12px; color: var(--rag-muted);">Обновлён</div>'
                            f'<div style="font-size: 13px; color: var(--rag-ink-2); font-weight: 500;">'
                            f'{_last_mod}</div>',
                            sanitize=False,
                        )

            # Быстрые запросы (коллапс)
            with ui.expansion("Быстрые запросы", icon="bolt", value=True).classes(
                "rag-card w-full"
            ):
                with ui.row().classes("w-full gap-2 flex-wrap p-1"):
                    for label, query in SEARCH_PRESETS:
                        ui.button(label, on_click=choose_query_handler(query)).props(
                            "outline no-caps"
                        ).style("font-size: 13px;")

            # История запросов (коллапс)
            _recent = list(state.history[:12])
            if _recent:
                with ui.expansion(
                    f"История запросов ({len(_recent)})", icon="history", value=True
                ).classes("rag-card w-full"):
                    with ui.column().classes("w-full gap-1 p-1"):
                        for q in _recent:
                            with ui.row().classes("w-full items-center gap-2"):
                                ui.icon("history", size="16px").style(
                                    "color: var(--rag-ink-4); flex-shrink: 0;"
                                )
                                btn = ui.button(q, on_click=choose_query_handler(q)).props(
                                    "flat dense no-caps align=left"
                                ).classes("flex-1 text-left")
                                btn.style(
                                    "font-size: 13px; color: var(--rag-ink-2);"
                                    " justify-content: flex-start;"
                                )
            return

        # RAG Q&A карточка
        if state.rag_answer_loading:
            with ui.row().classes("rag-card w-full p-3 gap-2 items-center"):
                ui.spinner(size="sm")
                ui.label("Анализирую документы…").classes("rag-meta")
        elif state.rag_answer_text:
            with ui.column().classes("rag-card w-full p-3 gap-1"):
                with ui.row().classes("items-center gap-1"):
                    ui.icon("smart_toy", size="18px").style("color: var(--rag-accent);")
                    ui.label("Ответ ИИ").classes("font-semibold text-sm").style("color: var(--rag-accent);")
                ui.label(state.rag_answer_text).classes("text-sm whitespace-pre-wrap")

        if not state.results:
            ui.label("Совпадений не найдено.").classes("rag-card p-4 rag-meta")
            return

        # Сортировка
        _sort_key = getattr(state, "search_sort", "score")

        def _sorted(results: list, key: str) -> list:
            if key == "name":
                return sorted(results, key=lambda r: (r.get("filename") or "").lower())
            if key == "date":
                return sorted(results, key=lambda r: r.get("modified") or "", reverse=True)
            return sorted(results, key=lambda r: float(r.get("score") or 0), reverse=True)

        sorted_results = _sorted(list(state.results), _sort_key)

        # Считаем кол-во по группам для чипов
        group_counts: Dict[str, int] = {}
        for r in sorted_results:
            group_counts.setdefault(_result_group(r), 0)
            group_counts[_result_group(r)] += 1

        group_order = [
            "Каталоги", "Техпаспорта ТС", "Паспорта и удостоверения",
            "Договоры", "Счета и платежи", "Таблицы", "PDF", "Другие файлы",
        ]

        # ── Refine bar (уточнить в найденном) ────────────────────────────────
        def remove_refine(term: str) -> None:
            if term in state.search_refine_terms:
                state.search_refine_terms.remove(term)
            render()

        refine_input_ref: Dict[str, Any] = {}

        def apply_refine() -> None:
            inp = refine_input_ref.get("input")
            if inp is None:
                return
            term = str(inp.value or "").strip()
            if term and term not in state.search_refine_terms:
                state.search_refine_terms.append(term)
                inp.set_value("")
            render()

        with ui.element("div").classes("rag-refine-bar w-full"):
            ui.html('<span class="rag-refine-bar-label">↳ уточнить:</span>', sanitize=False)
            refine_inp = ui.input(placeholder="введите слово для уточнения…").props(
                "dense borderless"
            ).style("flex: 1; font-size: 13px;").on("keydown.enter", apply_refine)
            refine_input_ref["input"] = refine_inp
            for term in state.search_refine_terms:
                ui.html(
                    f'<span class="rag-refine-chip">{html.escape(term)} ×</span>',
                    sanitize=False,
                ).on("click", lambda t=term: remove_refine(t))
            if state.search_refine_terms:
                ui.html(
                    f'<span style="font-size:11px;color:var(--rag-muted);">'
                    f'{len(state.search_refine_terms)} уточнений</span>',
                    sanitize=False,
                )

        # ── Stats spoiler (коллапс) ────────────────────────────────────────
        # Подсчёт по категориям/типам
        ext_counts: Dict[str, int] = {}
        for r in sorted_results:
            ext = (r.get("extension") or "").upper().lstrip(".") or "—"
            ext_counts.setdefault(ext, 0)
            ext_counts[ext] += 1
        top_exts = sorted(ext_counts.items(), key=lambda x: x[1], reverse=True)[:5]

        # title_match / semantic counts (аппроксимация)
        _title_match = sum(1 for r in sorted_results if float(r.get("score") or 0) >= 0.95)
        _content_match = sum(1 for r in sorted_results if 0.75 <= float(r.get("score") or 0) < 0.95)
        _semantic = len(sorted_results) - _title_match - _content_match

        stats_count_str = f"{len(sorted_results)}"
        if state.search_refine_terms:
            # apply refine filter
            _refine_filtered = [
                r for r in sorted_results
                if all(
                    t.lower() in (r.get("filename") or "").lower()
                    or t.lower() in (r.get("path") or "").lower()
                    for t in state.search_refine_terms
                )
            ]
            stats_count_str = f"{len(sorted_results)} → {len(_refine_filtered)} после уточнения"
            sorted_results = _refine_filtered

        if state.expanded_query:
            exp_note = f" · расширен: {state.expanded_query}"
        else:
            exp_note = ""

        with ui.expansion(
            f"Статистика выдачи  ·  {stats_count_str}{exp_note}",
            icon="bar_chart",
            value=False,
        ).classes("rag-card w-full"):
            with ui.row().classes("w-full gap-4 flex-wrap p-2"):
                for label, value, color in [
                    ("Точные в названиях", _title_match, "var(--rag-ok)"),
                    ("Точные в тексте", _content_match, "var(--rag-accent)"),
                    ("Семантически близкие", _semantic, "var(--rag-muted)"),
                    ("Всего", len(state.results), "var(--rag-text)"),
                ]:
                    with ui.column().classes("gap-0"):
                        ui.html(
                            f'<div style="font-size:10px;font-family:var(--rag-mono-font);'
                            f'color:var(--rag-ink-4);text-transform:uppercase;letter-spacing:.06em;">'
                            f'{label}</div>'
                            f'<div style="font-size:26px;font-weight:800;font-family:var(--rag-display-font);'
                            f'color:{color};letter-spacing:-0.03em;">{value}</div>',
                            sanitize=False,
                        )
                if top_exts:
                    ui.element("div").style("width:1px;background:var(--rag-border);align-self:stretch;")
                    with ui.column().classes("gap-1"):
                        ui.html(
                            '<div style="font-size:10px;font-family:var(--rag-mono-font);'
                            'color:var(--rag-ink-4);text-transform:uppercase;letter-spacing:.06em;">По типам</div>',
                            sanitize=False,
                        )
                        with ui.row().classes("gap-2 flex-wrap"):
                            for ext, cnt in top_exts:
                                ui.label(f"{ext} · {cnt}").classes("rag-chip")

        # ── Sort + filter toolbar ────────────────────────────────────────────
        def set_sort(key: str) -> None:
            state.search_sort = key  # type: ignore[attr-defined]
            render()

        def set_filter(gname: Optional[str]) -> None:
            state.active_type_filter = gname
            state.displayed_count = 10
            render()

        with ui.row().classes("w-full items-center gap-2 flex-wrap"):
            ui.html(
                f'<span style="font-size:13px;font-weight:600;color:var(--rag-ink-2);">'
                f'{len(sorted_results)} результатов</span>',
                sanitize=False,
            )
            ui.element("div").style("width:1px;height:16px;background:var(--rag-border);")
            # Группировка-чипы
            all_active = state.active_type_filter is None
            all_chip = ui.label(f"Все").classes(
                "rag-chip" + (" rag-chip-active" if all_active else "")
            )
            all_chip.on("click", lambda: set_filter(None))
            for gname in group_order:
                cnt = group_counts.get(gname, 0)
                if cnt == 0:
                    continue
                is_active = state.active_type_filter == gname
                chip = ui.label(f"{gname} ({cnt})").classes(
                    "rag-chip" + (" rag-chip-active" if is_active else "")
                )
                chip.on("click", lambda g=gname: set_filter(g))

            ui.element("div").style("flex: 1;")

            # Сортировка
            for sk, slabel in [("score", "Релевантность ▾"), ("name", "Имя ▾"), ("date", "Дата ▾")]:
                is_cur = _sort_key == sk
                btn = ui.html(
                    f'<button class="rag-sort-btn{"  active" if is_cur else ""}">{slabel}</button>',
                    sanitize=False,
                )
                btn.on("click", lambda s=sk: set_sort(s))

        # Применяем фильтр по группе
        if state.active_type_filter:
            visible = [r for r in sorted_results if _result_group(r) == state.active_type_filter]
        else:
            visible = sorted_results

        # Показываем первые displayed_count штук
        to_show = visible[: state.displayed_count]
        with ui.column().classes("w-full gap-3"):
            for idx, result in enumerate(to_show, 1):
                render_result(result, idx)

        # Кнопка «Загрузить ещё»
        remaining = len(visible) - state.displayed_count
        if remaining > 0:
            def load_more() -> None:
                state.displayed_count += 10
                render()

            ui.button(
                f"Загрузить ещё ({remaining})",
                on_click=load_more,
                icon="expand_more",
            ).props("outline no-caps").classes("w-full mt-1")

        # ── Preview hint ─────────────────────────────────────────────────────
        ui.html(
            '<div class="rag-preview-hint">▲ ПРЕВЬЮ — дважды кликните по файлу для просмотра</div>',
            sanitize=False,
        )

        # ── Selection bar ─────────────────────────────────────────────────────
        _sel = state.selected_result_paths
        if _sel:
            with ui.element("div").classes("rag-selection-bar"):
                ui.html(f'<span>выбрано: {len(_sel)}</span>', sanitize=False)
                ui.button("В проводник", icon="folder", on_click=lambda: go_explorer(_sel[0])).props("flat dense no-caps")
                if len(_sel) == 1:
                    _sel_path = Path(_sel[0])
                    if _sel_path.exists():
                        ui.button("Скачать", icon="download", on_click=lambda: ui.download(_sel_path, filename=_sel_path.name)).props("flat dense no-caps")
                ui.button("Открыть в ОС", icon="open_in_new", on_click=lambda: _open_os_path(str(Path(_sel[0]).parent))).props("flat dense no-caps")
                ui.element("div").style("flex:1;")
                ui.button(icon="close", on_click=lambda: (state.selected_result_paths.clear(), render()), color=None).props("flat round dense")

    def render_explorer_screen() -> None:
        root = Path(str(state.cfg.get("catalog_path") or ""))
        if not root.exists():
            ui.label(f"Каталог не найден: {root}").classes("text-red-700 rag-card p-4")
            return

        toolbar = ui.column().classes("w-full gap-3")
        entries_area = ui.column().classes("w-full gap-3")

        def open_folder(path: Path) -> None:
            state.explorer_path = str(path)
            state.explorer_page = 0
            _get_auth_db(state).touch_favorite(username=_username(state), path=str(path))
            _log_app_event(state, "explorer", "open_folder", details={"path": str(path)})
            render()

        def open_file(path: Path) -> None:
            if path.exists() and path.is_file():
                _get_auth_db(state).touch_favorite(username=_username(state), path=str(path))
                _log_app_event(state, "explorer", "open_file", details={"path": str(path)})
                open_file_viewer(path)

        def copy_path(path: Path) -> None:
            ui.run_javascript(f"navigator.clipboard.writeText({json.dumps(str(path))})")
            ui.notify("Путь скопирован.", type="positive")

        def explorer_context_props(path: Path, *, is_dir: bool) -> str:
            item_type = "folder" if is_dir else "file"
            item_url = "" if is_dir else _viewer_file_url(str(path))
            favorite = "true" if _is_favorite(state, str(path)) else "false"
            attrs = {
                "data-rag-context": "explorer-item",
                "data-rag-type": item_type,
                "data-rag-path": quote(str(path), safe=""),
                "data-rag-url": item_url,
                "data-rag-favorite": favorite,
            }
            return " ".join(f'{key}="{html.escape(value, quote=True)}"' for key, value in attrs.items())

        def render_star(path: Path, *, item_type: Optional[str] = None) -> None:
            active = _is_favorite(state, str(path))
            icon = "star" if active else "star_border"
            star = ui.button(icon=icon, color=None).props("flat round dense data-rag-favorite-button")
            star.classes("rag-favorite-star active" if active else "rag-favorite-star")
            star.tooltip("Убрать из избранного" if active else "Добавить в избранное")

            def toggle() -> None:
                _toggle_favorite(state, path, item_type=item_type)
                render()

            star.on("click.stop", toggle)

        def open_favorites_dialog() -> None:
            with ui.dialog() as dialog, ui.card().classes("w-[min(900px,92vw)] max-h-[80vh] overflow-auto gap-3"):
                ui.label("Избранное").classes("text-xl font-semibold")
                if not state.favorites:
                    ui.label("Закладок пока нет.").classes("rag-meta")
                for fav in state.favorites:
                    fav_path = Path(str(fav.get("path") or ""))
                    item_type = str(fav.get("item_type") or "file")
                    label = str(fav.get("title") or fav_path.name or fav_path)
                    with ui.element("div").classes("rag-favorites-dialog-row"):
                        ui.icon("folder" if item_type == "folder" else "description")
                        with ui.column().classes("min-w-0 gap-0"):
                            ui.label(label).classes("font-medium truncate")
                            ui.label(str(fav_path)).classes("rag-path truncate")
                        action = (lambda p=fav_path: (dialog.close(), open_folder(p))) if item_type == "folder" else (lambda p=fav_path: (dialog.close(), open_file(p)))
                        ui.button("Открыть", on_click=action).props("outline dense")
                        ui.button(icon="close", on_click=lambda p=fav_path: (_toggle_favorite(state, p), dialog.close(), render())).props("flat round dense").tooltip("Убрать из избранного")
                ui.button("Закрыть", on_click=dialog.close).props("flat")
            dialog.open()

        def render_tile(path: Path, is_dir: bool, size_class: str) -> None:
            icon = _file_icon_svg(str(path), "Каталог" if is_dir else "Файл")
            click = (lambda p=path: open_folder(p)) if is_dir else (lambda p=path: open_file(p))
            system_class = " system" if not is_dir and _is_system_file(path) else ""
            tile = ui.column().classes(f"rag-explorer-item items-center gap-1 p-2 {size_class}{system_class}")
            tile.props(explorer_context_props(path, is_dir=is_dir))
            with tile:
                with ui.element("div").classes("rag-tile-star-wrap"):
                    render_star(path, item_type="folder" if is_dir else "file")
                opener = ui.column().classes("rag-explorer-opener items-center gap-1 cursor-pointer").on("click", click)
                opener.props("data-rag-open")
                with opener:
                    ui.html(icon, sanitize=False)
                    name_label = ui.label(path.name).classes("rag-explorer-name text-center text-sm")
                    name_label.tooltip(str(path.name))
                os_button = ui.button(on_click=lambda p=path: _open_os_path(str(p.parent if p.is_file() else p))).props("data-rag-os")
                os_button.classes("hidden")

        def render_row(path: Path, is_dir: bool, compact: bool = False) -> None:
            try:
                stat = path.stat()
                size = "" if is_dir else _format_file_size(stat.st_size)
                modified = time.strftime("%d.%m.%Y %H:%M", time.localtime(stat.st_mtime))
            except Exception:
                size, modified = "", ""
            system_class = " system" if not is_dir and _is_system_file(path) else ""
            row = ui.row().classes(f"rag-explorer-item w-full p-2 items-center gap-3{system_class}")
            row.props(explorer_context_props(path, is_dir=is_dir))
            with row:
                ui.html(_file_icon_svg(str(path), "Каталог" if is_dir else "Файл"), sanitize=False)
                action = (lambda p=path: open_folder(p)) if is_dir else (lambda p=path: open_file(p))
                with ui.column().classes("flex-1 gap-0"):
                    open_btn = ui.button(path.name, on_click=action, color=None).props("flat align=left no-caps dense data-rag-open").classes("rag-nav-button w-full")
                    open_btn.tooltip(str(path.name))
                    if not compact:
                        ui.label(f"{'Папка' if is_dir else path.suffix or 'без расширения'} · {size} · {modified}").classes("rag-meta")
                if not compact:
                    if not is_dir:
                        ui.button("Скачать", icon="download", on_click=lambda p=path: (_log_app_event(state, "explorer", "download", details={"path": str(p)}), ui.download(p, filename=p.name))).props("outline dense")
                    ui.button("ОС", icon="open_in_new", on_click=lambda p=path: _open_os_path(str(p.parent if p.is_file() else p))).props("flat dense data-rag-os")
                else:
                    os_button = ui.button(on_click=lambda p=path: _open_os_path(str(p.parent if p.is_file() else p))).props("data-rag-os")
                    os_button.classes("hidden")
                render_star(path, item_type="folder" if is_dir else "file")

        def render_entries() -> None:
            entries_area.clear()
            current = _safe_explorer_path(state)
            if not current.exists():
                state.explorer_path = str(root)
                current = root

            parts: List[Path] = []
            p = current
            while True:
                parts.append(p)
                if p == root or p == p.parent:
                    break
                p = p.parent
            parts.reverse()

            if state.header_breadcrumbs is not None:
                state.header_breadcrumbs.clear()
                with state.header_breadcrumbs:
                    for idx, part in enumerate(parts):
                        label = "Корень" if part == root else part.name
                        ui.button(label, on_click=lambda p=part: (_log_app_event(state, "explorer", "breadcrumb", details={"path": str(p)}), open_folder(p)), color=None).props("flat dense no-caps")
                        if idx < len(parts) - 1:
                            ui.icon("chevron_right").classes("text-slate-400")

            if state.header_explorer_actions is not None:
                state.header_explorer_actions.clear()
                with state.header_explorer_actions:
                    active = _is_favorite(state, str(current))
                    fav = ui.button(icon="star" if active else "star_border", color=None).props("flat round dense")
                    fav.classes("rag-favorite-star header active" if active else "rag-favorite-star header")
                    fav.tooltip("Убрать текущую папку из избранного" if active else "Добавить текущую папку в избранное")
                    fav.on("click", lambda p=current: (_toggle_favorite(state, p, item_type="folder"), render()))

            dirs, files, total_files = _file_rows(current, state)
            state.explorer_page = max(0, min(state.explorer_page, max(0, (len(files) - 1) // PAGE_SIZE)))
            page_files = files[state.explorer_page * PAGE_SIZE : (state.explorer_page + 1) * PAGE_SIZE]

            with entries_area:
                with ui.row().classes("w-full items-center gap-2"):
                    up_button = ui.button(icon="arrow_upward", on_click=lambda: (_log_app_event(state, "explorer", "up", details={"path": str(current.parent)}), open_folder(current.parent)), color=None).props("outline round dense")
                    up_button.tooltip("На уровень выше")
                    if current == root:
                        up_button.disable()
                    ui.label(f"{current} · папок {len(dirs)} · файлов {total_files}").classes("rag-path")

                if state.favorites:
                    with ui.row().classes("rag-bookmarks"):
                        for fav in state.favorites:
                            fav_path = Path(str(fav.get("path") or ""))
                            item_type = str(fav.get("item_type") or "file")
                            label = str(fav.get("title") or fav_path.name or fav_path)
                            icon = "folder" if item_type == "folder" else "description"
                            action = (lambda p=fav_path: open_folder(p)) if item_type == "folder" else (lambda p=fav_path: open_file(p))
                            with ui.element("div").classes("rag-bookmark"):
                                with ui.element("div").classes("rag-bookmark-main"):
                                    button = ui.button(label, icon=icon, on_click=action, color=None).props("flat dense no-caps").classes("rag-nav-button")
                                    button.tooltip(label)
                                with ui.element("div").classes("rag-bookmark-remove"):
                                    remove_button = ui.button(icon="close", color=None).props("flat round dense")
                                    remove_button.tooltip("Убрать из избранного")
                                    remove_button.on("click.stop", lambda p=fav_path: (_toggle_favorite(state, p), render()))
                        ui.button(icon="more_horiz", on_click=open_favorites_dialog, color=None).props("outline round dense").classes("rag-bookmark-more").tooltip("Показать все избранное")

                if not dirs and not files:
                    ui.label("Нет элементов, соответствующих фильтру.").classes("rag-card p-4 rag-meta")
                    return

                if state.explorer_view in {"Крупные значки", "Средние значки", "Мелкие значки"}:
                    grid_class = {
                        "Крупные значки": "",
                        "Средние значки": "medium",
                        "Мелкие значки": "small",
                    }[state.explorer_view]
                    with ui.element("div").classes(f"rag-explorer-grid {grid_class} w-full"):
                        for path in [*dirs, *page_files]:
                            render_tile(path, path.is_dir(), grid_class)
                elif state.explorer_view == "Список":
                    with ui.column().classes("rag-explorer-list w-full"):
                        for path in [*dirs, *page_files]:
                            render_row(path, path.is_dir(), compact=True)
                else:
                    with ui.column().classes("w-full gap-2"):
                        for path in [*dirs, *page_files]:
                            render_row(path, path.is_dir(), compact=False)

                if total_files > PAGE_SIZE:
                    with ui.row().classes("items-center gap-2"):
                        ui.button("Назад", on_click=lambda: (setattr(state, "explorer_page", max(0, state.explorer_page - 1)), render_entries())).props("outline")
                        ui.label(f"Страница {state.explorer_page + 1} из {(total_files + PAGE_SIZE - 1) // PAGE_SIZE}").classes("rag-meta")
                        ui.button("Вперед", on_click=lambda: (setattr(state, "explorer_page", state.explorer_page + 1), render_entries())).props("outline")

        with toolbar:
            with ui.row().classes("rag-card w-full p-3 gap-3 items-center"):
                filter_input = ui.input(placeholder="Фильтр по имени", value=state.explorer_filter).props("dense outlined clearable debounce=0").classes("min-w-64 flex-1")

                def update_explorer_setting(attr: str, value: Any) -> None:
                    setattr(state, attr, value)
                    state.explorer_page = 0
                    _save_explorer_settings(state)
                    _log_app_event(state, "explorer", "change_setting", details={attr: value})
                    render_entries()

                ui.select(["Все", ".docx", ".xlsx", ".xls", ".pdf"], value=state.explorer_ext, on_change=lambda e: update_explorer_setting("explorer_ext", e.value)).props("dense outlined").classes("w-36")
                ui.select(["Крупные значки", "Средние значки", "Мелкие значки", "Список", "Таблица"], value=state.explorer_view, on_change=lambda e: update_explorer_setting("explorer_view", e.value)).props("dense outlined").classes("w-44")
                ui.select(["По имени", "По размеру", "По дате"], value=state.explorer_sort, on_change=lambda e: update_explorer_setting("explorer_sort", e.value)).props("dense outlined").classes("w-40")
                ui.select(["По возрастанию", "По убыванию"], value="По убыванию" if state.explorer_desc else "По возрастанию", on_change=lambda e: update_explorer_setting("explorer_desc", e.value == "По убыванию")).props("dense outlined").classes("w-44")

                def apply_filter(event: events.ValueChangeEventArguments | events.GenericEventArguments | None = None) -> None:
                    _apply_explorer_filter_input(state, event, filter_input.value)
                    render_entries()

                filter_input.on_value_change(apply_filter)

        render_entries()

    def render_index_screen() -> None:  # noqa: C901
        if not _is_admin(state):
            ui.label("Раздел индексирования доступен только администратору.").classes("rag-card p-4 text-red-700")
            return
        stats = _read_index_stats(state.cfg)
        telemetry = _read_index_telemetry(state.cfg)
        settings_db = _get_telemetry(state)
        settings = settings_db.get_index_settings() if hasattr(settings_db, "get_index_settings") else {}

        ui.html(
            '<div style="font-family: var(--rag-display-font); font-size: 26px;'
            ' font-weight: 700; color: var(--rag-text); letter-spacing: -0.02em;">Индексация</div>'
            '<div style="font-size: 13px; color: var(--rag-muted); margin-top: 2px;">'
            "Запуск, расписание, прогресс этапов и OCR.</div>",
            sanitize=False,
        )

        # ── Метрики ──────────────────────────────────────────────────────
        def render_metric(label: str, value: str, icon: str = "analytics") -> None:
            with ui.column().classes("rag-card p-4 gap-1 min-w-52 flex-1"):
                with ui.row().classes("items-center gap-2"):
                    ui.icon(icon).classes("text-xl")
                    ui.label(label).classes("rag-meta")
                ui.label(value).classes("text-xl font-semibold")

        with ui.row().classes("w-full gap-3"):
            render_metric("Файлов в state", f"{stats['total']:,}".replace(",", " "), "description")
            render_metric("Размер файлов", _format_bytes(stats.get("total_size_bytes")), "storage")
            render_metric("State обновлен", str(stats.get("last_modified") or "не найден"), "schedule")
            overall = telemetry.get("overall") or {}
            render_metric("Средняя длительность", _format_duration_seconds(overall.get("avg_duration_sec")), "timer")

        # ── Запустить сейчас ─────────────────────────────────────────────
        with ui.column().classes("rag-card w-full p-4 gap-3"):
            with ui.row().classes("w-full items-center gap-2"):
                ui.icon("play_circle").classes("text-2xl text-indigo-500")
                ui.label("Запустить сейчас").classes("text-xl font-semibold")
            ui.label("Запускается фоновый процесс; прогресс появится в разделе «Прогресс этапов» через несколько секунд.").classes("rag-meta")
            workers_now = int(settings.get("workers") or state.cfg.get("index_read_workers") or 4)
            chunks_now = int(settings.get("max_chunks") or state.cfg.get("index_max_chunks") or 2000)
            skip_ocr_now = bool(settings.get("skip_inline_ocr"))
            ocr_min_len_now = int(settings.get("ocr_min_text_len") or 50)

            def make_run_handler(stage_key: str) -> Any:
                def handler() -> None:
                    try:
                        pid = _launch_indexer(
                            state.cfg,
                            stage=stage_key,
                            workers=workers_now,
                            max_chunks=chunks_now,
                            skip_inline_ocr=skip_ocr_now,
                        )
                    except RuntimeError as exc:
                        ui.notify(str(exc), type="warning")
                        return
                    _log_app_event(state, "index", "run_now", details={"stage": stage_key, "pid": pid})
                    ui.notify(f"Индексация «{_STAGE_LABELS.get(stage_key, stage_key)}» запущена (PID {pid}).", type="positive")
                return handler

            with ui.row().classes("w-full gap-2 flex-wrap"):
                for stage_key, stage_label in _STAGE_LABELS.items():
                    color = "primary" if stage_key == "all" else None
                    ui.button(stage_label, icon="play_arrow", on_click=make_run_handler(stage_key), color=color).props("unelevated" if stage_key == "all" else "outline")

            ui.separator()
            with ui.row().classes("w-full items-center gap-3"):
                ui.icon("scanner").classes("text-xl text-orange-500")
                ui.label("OCR распознавание").classes("font-semibold")

            def run_ocr_now() -> None:
                try:
                    pid = _launch_ocr(state.cfg, min_text_len=ocr_min_len_now)
                except RuntimeError as exc:
                    ui.notify(str(exc), type="warning")
                    return
                _log_app_event(state, "index", "run_ocr_now", details={"pid": pid})
                ui.notify(f"OCR-проход запущен (PID {pid}).", type="positive")

            ui.button("Запустить OCR", icon="document_scanner", on_click=run_ocr_now).props("outline color=orange")

        # ── Прогресс этапов (V2 phase rows) ─────────────────────────────
        with ui.column().classes("rag-card w-full p-4 gap-3"):
            with ui.row().classes("w-full items-center gap-2"):
                ui.icon("account_tree").classes("text-2xl text-indigo-500")
                ui.label("Этапы индексации").classes("text-xl font-semibold")
                ui.space()
                refresh_btn = ui.button(icon="refresh", on_click=lambda: _refresh_progress()).props("flat dense round").tooltip("Обновить")
            phases_area = ui.column().classes("w-full gap-3")

            def _freshness_color(finished_at_str: str) -> str:
                """Return CSS color based on how old the last run was."""
                if not finished_at_str or finished_at_str.startswith("0"):
                    return "var(--rag-border)"
                try:
                    import datetime as _dtt
                    ts = _dtt.datetime.fromisoformat(finished_at_str.replace(" ", "T"))
                    age_h = (_dtt.datetime.now() - ts).total_seconds() / 3600
                    if age_h < 2:
                        return "var(--rag-ok)"
                    elif age_h < 24:
                        return "#f59e0b"
                    elif age_h < 72:
                        return "#f97316"
                    else:
                        return "var(--rag-danger)"
                except Exception:
                    return "var(--rag-border)"

            def _freshness_pct(finished_at_str: str) -> float:
                """Return 0..1 fill for freshness bar (1 = very fresh, 0 = stale)."""
                if not finished_at_str or finished_at_str.startswith("0"):
                    return 0.0
                try:
                    import datetime as _dtt
                    ts = _dtt.datetime.fromisoformat(finished_at_str.replace(" ", "T"))
                    age_h = (_dtt.datetime.now() - ts).total_seconds() / 3600
                    return max(0.0, min(1.0, 1.0 - age_h / 72.0))
                except Exception:
                    return 0.0

            def _normalize_stage_row(stage_key: str, source: Dict[str, Any]) -> Dict[str, Any]:
                """Унифицирует поля: ts_started/ts_finished → started_at/finished_at."""
                row: Dict[str, Any] = dict(source)
                row.setdefault("stage", stage_key)
                if not row.get("started_at"):
                    row["started_at"] = source.get("ts_started") or ""
                if not row.get("finished_at"):
                    row["finished_at"] = source.get("ts_finished") or ""
                # stage_summary возвращает last_duration_sec; active_stages — duration_sec
                if not row.get("duration_sec"):
                    row["duration_sec"] = (
                        source.get("duration_sec")
                        or source.get("last_duration_sec")
                        or 0
                    )
                return row

            def _refresh_progress() -> None:
                fresh = _read_index_telemetry(state.cfg)
                # Используем stage_summary (последний завершённый прогон каждого
                # этапа независимо), а active_stages накладываем сверху для
                # текущих running строк. latest_stages нам НЕ нужен — он
                # ограничен одним последним run_id и врёт когда пользователь
                # запускал отдельные этапы.
                summary = fresh.get("stage_summary") or []
                active = fresh.get("active_stages") or []
                active_runs = fresh.get("active_runs") or []
                active_ocr = fresh.get("active_ocr")
                last_ocr = fresh.get("last_ocr")

                all_rows: Dict[str, Dict[str, Any]] = {}
                for _row in summary:
                    sk = str(_row.get("stage") or "")
                    if sk:
                        all_rows[sk] = _normalize_stage_row(sk, _row)
                for _row in active:
                    sk = str(_row.get("stage") or "")
                    if sk:
                        all_rows[sk] = _normalize_stage_row(sk, _row)

                # OCR живёт в отдельной таблице ocr_runs — добавляем как
                # псевдо-этап с собственными счётчиками.
                ocr_src = active_ocr or last_ocr
                if ocr_src:
                    ocr_row = _normalize_stage_row("ocr", ocr_src)
                    ocr_row.setdefault("processed_files", int(ocr_src.get("processed_pdfs") or 0))
                    ocr_row.setdefault("total_files", int(ocr_src.get("found_scanned") or 0))
                    if active_ocr:
                        ocr_row["status"] = "running"
                    all_rows["ocr"] = ocr_row

                # Текущая команда управления (для отображения paused и выбора кнопок).
                control_cmd = _get_indexer_command()
                indexer_alive = bool(active_runs)
                live_pid = 0
                for r in active_runs:
                    try:
                        pid_i = int(r.get("worker_pid") or 0)
                    except (TypeError, ValueError):
                        pid_i = 0
                    if pid_i > 0:
                        live_pid = pid_i
                        break

                # ── фабрики обработчиков (замыкаем state/settings) ────
                def _make_pause_handler():
                    def _h() -> None:
                        try:
                            _write_indexer_control("pause")
                            _log_app_event(state, "index", "control_pause")
                            ui.notify("Команда «Пауза» отправлена.", type="info")
                        except (OSError, ValueError) as exc:
                            ui.notify(f"Не удалось поставить на паузу: {exc}", type="negative")
                        _refresh_progress()
                    return _h

                def _make_resume_handler():
                    def _h() -> None:
                        try:
                            _write_indexer_control("running")
                            _log_app_event(state, "index", "control_resume")
                            ui.notify("Индексация возобновлена.", type="positive")
                        except (OSError, ValueError) as exc:
                            ui.notify(f"Не удалось возобновить: {exc}", type="negative")
                        _refresh_progress()
                    return _h

                def _make_cancel_handler():
                    def _h() -> None:
                        try:
                            _write_indexer_control("cancel")
                        except (OSError, ValueError) as exc:
                            ui.notify(f"Не удалось отправить отмену: {exc}", type="negative")
                            return
                        _log_app_event(state, "index", "control_cancel", details={"pid": live_pid})
                        ui.notify(
                            "Команда «Отмена» отправлена. "
                            "Индексатор завершится после текущего файла.",
                            type="warning",
                        )
                        _refresh_progress()
                    return _h

                def _make_kill_handler(pid: int):
                    def _h() -> None:
                        if pid <= 0:
                            ui.notify("Не найден живой процесс индексатора.", type="warning")
                            return
                        if _terminate_process(pid):
                            _log_app_event(state, "index", "control_kill", details={"pid": pid})
                            ui.notify(
                                f"Процесс индексатора (PID {pid}) принудительно остановлен.",
                                type="warning",
                            )
                        else:
                            ui.notify(f"Не удалось остановить PID {pid}.", type="negative")
                        _refresh_progress()
                    return _h

                def _make_stage_runner(stage_key: str):
                    def _h() -> None:
                        try:
                            pid = _launch_indexer(
                                state.cfg,
                                stage=stage_key,
                                workers=int(settings.get("workers") or state.cfg.get("index_read_workers") or 4),
                                max_chunks=int(settings.get("max_chunks") or state.cfg.get("index_max_chunks") or 2000),
                                skip_inline_ocr=bool(settings.get("skip_inline_ocr")),
                            )
                        except RuntimeError as exc:
                            ui.notify(str(exc), type="warning")
                            return
                        _log_app_event(state, "index", "run_stage", details={"stage": stage_key, "pid": pid})
                        ui.notify(
                            f"Этап «{_STAGE_LABELS.get(stage_key, stage_key)}» запущен (PID {pid}).",
                            type="positive",
                        )
                        _refresh_progress()
                    return _h

                def _make_ocr_runner():
                    def _h() -> None:
                        try:
                            pid = _launch_ocr(
                                state.cfg,
                                min_text_len=int(settings.get("ocr_min_text_len") or 50),
                            )
                        except RuntimeError as exc:
                            ui.notify(str(exc), type="warning")
                            return
                        _log_app_event(state, "index", "run_ocr_stage", details={"pid": pid})
                        ui.notify(f"OCR-проход запущен (PID {pid}).", type="positive")
                        _refresh_progress()
                    return _h

                # V2 phase order (including OCR as a separate phase)
                phase_order = ["metadata", "small", "large", "content", "ocr"]

                phases_area.clear()
                with phases_area:
                    if not all_rows:
                        ui.html(
                            '<div class="rag-meta" style="padding:16px;text-align:center;">' +
                            'Индексация ещё не запускалась. Нажмите «Запустить сейчас» выше.</div>',
                            sanitize=False,
                        )
                    for stage_key in phase_order:
                        row = all_rows.get(stage_key)
                        if row is None:
                            stage_label = _STAGE_LABELS.get(stage_key, stage_key)
                            with ui.element("div").classes("rag-phase-row"):
                                ui.html(
                                    '<div class="rag-phase-circle" style="background:var(--rag-border);">' +
                                    '<span style="color:var(--rag-muted);font-size:16px;">—</span></div>',
                                    sanitize=False,
                                )
                                with ui.column().classes("flex-1 gap-1"):
                                    ui.html(
                                        f'<span class="font-semibold" style="color:var(--rag-ink-3);">{stage_label}</span>',
                                        sanitize=False,
                                    )
                                    ui.html('<span class="rag-meta">не запускался</span>', sanitize=False)
                                    if not indexer_alive:
                                        with ui.row().classes("gap-2 mt-1"):
                                            if stage_key == "ocr":
                                                ui.button(
                                                    "Запустить", icon="play_arrow",
                                                    on_click=_make_ocr_runner(),
                                                ).props("outline dense color=primary size=sm")
                                            elif stage_key != "content":
                                                ui.button(
                                                    "Запустить", icon="play_arrow",
                                                    on_click=_make_stage_runner(stage_key),
                                                ).props("outline dense color=primary size=sm")
                            continue

                        processed = int(row.get("processed_files") or 0)
                        total_f = int(row.get("total_files") or 0)
                        pct = min(1.0, processed / total_f) if total_f > 0 else (1.0 if str(row.get("status") or "") not in ("running", "") else 0.0)
                        raw_status = str(row.get("status") or "-").lower()
                        is_running_db = raw_status == "running"
                        is_paused_db = raw_status == "paused"
                        # Если в БД running, но UI попросил pause — отображаем как paused.
                        # OCR не использует контрольный файл, на него не распространяем.
                        if stage_key == "ocr":
                            is_paused = is_paused_db
                            is_running = is_running_db
                        else:
                            is_paused = is_paused_db or (is_running_db and control_cmd == "pause")
                            is_running = is_running_db and not is_paused
                        is_done = raw_status in ("done", "completed", "finished")
                        is_error = raw_status in ("error", "failed")
                        is_cancelled = raw_status in ("cancelled", "canceled", "aborted")
                        # Display label учитывает paused-наложение
                        status_str = "paused" if is_paused else raw_status

                        stage_label = _STAGE_LABELS.get(stage_key, stage_key)
                        finished_at = str(row.get("finished_at") or "")
                        started_at = str(row.get("started_at") or "")[:16].replace("T", " ")
                        duration_str = _format_duration_seconds(row.get("duration_sec"))

                        if is_paused:
                            circle_bg = "#f59e0b"
                            circle_icon = "⏸"
                        elif is_running:
                            circle_bg = "var(--rag-accent)"
                            circle_icon = "◷"
                        elif is_done:
                            circle_bg = "var(--rag-ok)"
                            circle_icon = "✓"
                        elif is_error:
                            circle_bg = "var(--rag-danger)"
                            circle_icon = "✗"
                        elif is_cancelled:
                            circle_bg = "var(--rag-muted)"
                            circle_icon = "■"
                        else:
                            circle_bg = "var(--rag-muted)"
                            circle_icon = "—"

                        row_extra = "running" if (is_running or is_paused) else ""
                        fresh_color = _freshness_color(finished_at)
                        fresh_pct = _freshness_pct(finished_at) * 100

                        with ui.element("div").classes(f"rag-phase-row {row_extra}"):
                            ui.html(
                                f'<div class="rag-phase-circle" style="background:{circle_bg};color:#fff;font-size:16px;">' +
                                f'{circle_icon}</div>',
                                sanitize=False,
                            )
                            with ui.column().classes("flex-1 gap-1"):
                                with ui.row().classes("w-full items-center gap-2 flex-wrap"):
                                    ui.html(f'<span class="font-semibold">{stage_label}</span>', sanitize=False)
                                    status_bg = (
                                        "#f59e0b" if is_paused
                                        else "var(--rag-accent)" if is_running
                                        else "var(--rag-ok)" if is_done
                                        else "var(--rag-danger)" if is_error
                                        else "var(--rag-muted)"
                                    )
                                    ui.html(
                                        f'<span class="rag-chip" style="background:{status_bg}18;' +
                                        f'color:{status_bg};border-color:{status_bg};">{status_str}</span>',
                                        sanitize=False,
                                    )
                                    if started_at:
                                        ui.html(f'<span class="rag-meta">▶ {started_at}</span>', sanitize=False)
                                    ui.element("div").style("flex:1;")
                                    if duration_str:
                                        ui.html(f'<span class="rag-meta font-mono">{duration_str}</span>', sanitize=False)

                                if is_running or is_paused or total_f > 0:
                                    progress_color = (
                                        "warning" if is_paused
                                        else "primary" if is_running
                                        else "positive" if is_done
                                        else "negative" if is_error
                                        else "grey"
                                    )
                                    ui.linear_progress(value=pct).props(f"color={progress_color}").classes("w-full")
                                    ui.html(
                                        f'<span class="rag-meta">{processed:,} / {total_f:,} файлов</span>'.replace(",", " "),
                                        sanitize=False,
                                    )

                                if is_done and finished_at:
                                    fin_label = finished_at[:16].replace("T", " ")
                                    ui.html(
                                        f'<div style="display:flex;align-items:center;gap:8px;margin-top:4px;">' +
                                        f'<span class="rag-meta">свежесть:</span>' +
                                        f'<div class="rag-freshness-bar">' +
                                        f'<div class="rag-freshness-bar-fill" style="width:{fresh_pct:.0f}%;background:{fresh_color};"></div>' +
                                        f'</div>' +
                                        f'<span class="rag-meta">■ {fin_label}</span>' +
                                        f'</div>',
                                        sanitize=False,
                                    )

                                stats_parts = []
                                for lbl, key in [
                                    ("добавлено", "added_files"),
                                    ("обновлено", "updated_files"),
                                    ("пропущено", "skipped_files"),
                                    ("ошибок", "error_files"),
                                    ("точек", "points_added"),
                                ]:
                                    val = int(row.get(key) or 0)
                                    if val:
                                        col = "var(--rag-danger)" if key == "error_files" else "var(--rag-muted)"
                                        stats_parts.append(
                                            f'<span style="font-size:12px;color:{col};">' +
                                            f'{lbl} <b style="color:var(--rag-ink-2);">{val:,}</b></span>'.replace(",", " ")
                                        )
                                if stats_parts:
                                    ui.html(
                                        '<div style="display:flex;gap:12px;flex-wrap:wrap;margin-top:2px;">' +
                                        "".join(stats_parts) + "</div>",
                                        sanitize=False,
                                    )

                                with ui.row().classes("gap-2 mt-1"):
                                    if stage_key == "ocr":
                                        # OCR — отдельный процесс, контрольный
                                        # файл его не охватывает. На running —
                                        # только жёсткая остановка.
                                        if is_running:
                                            ocr_pid = int((active_ocr or {}).get("worker_pid") or 0)
                                            ui.button(
                                                "Остановить", icon="stop",
                                                on_click=lambda p=ocr_pid: (
                                                    _terminate_process(p),
                                                    _refresh_progress(),
                                                ),
                                            ).props("outline dense color=negative size=sm")
                                        else:
                                            lbl_btn = "Перезапустить" if is_error else "Запустить"
                                            ico_btn = "replay" if is_error else "play_arrow"
                                            ui.button(
                                                lbl_btn, icon=ico_btn,
                                                on_click=_make_ocr_runner(),
                                            ).props("outline dense color=primary size=sm")
                                    elif is_running or is_paused:
                                        # Активный индекс-этап: Пауза / Возобновить + Отмена + жёсткий kill
                                        if is_paused:
                                            ui.button(
                                                "Возобновить", icon="play_arrow",
                                                on_click=_make_resume_handler(),
                                            ).props("unelevated dense color=primary size=sm")
                                        else:
                                            ui.button(
                                                "Пауза", icon="pause",
                                                on_click=_make_pause_handler(),
                                            ).props("outline dense color=warning size=sm")
                                        ui.button(
                                            "Отмена", icon="stop",
                                            on_click=_make_cancel_handler(),
                                        ).props("outline dense color=negative size=sm")
                                        if live_pid > 0:
                                            ui.button(
                                                icon="power_settings_new",
                                                on_click=_make_kill_handler(live_pid),
                                            ).props("flat dense round color=negative size=sm").tooltip(
                                                f"Принудительно завершить PID {live_pid}"
                                            )
                                    else:
                                        if indexer_alive:
                                            ui.html(
                                                '<span class="rag-meta" style="font-size:12px;">'
                                                'индексатор занят другим этапом</span>',
                                                sanitize=False,
                                            )
                                        elif stage_key != "content":
                                            lbl_btn = "Перезапустить" if is_error else "Запустить"
                                            ico_btn = "replay" if is_error else "play_arrow"
                                            ui.button(
                                                lbl_btn, icon=ico_btn,
                                                on_click=_make_stage_runner(stage_key),
                                            ).props("outline dense color=primary size=sm")

            # Initial render
            _refresh_progress()
            # Auto-refresh every 5 seconds while indexing may be running
            ui.timer(5.0, _refresh_progress)

        # ── Расписание (список) ──────────────────────────────────────────
        with ui.column().classes("rag-card w-full p-4 gap-3"):
            with ui.row().classes("w-full items-center gap-2"):
                ui.icon("event_repeat").classes("text-2xl text-indigo-500")
                ui.label("Расписание индексации").classes("text-xl font-semibold")
            ui.label(
                "Встроенный планировщик проверяет расписание каждую минуту и запускает индексатор в фоне. "
                "Можно создать несколько записей — разные этапы в разное время."
            ).classes("rag-meta")

            schedules_list = settings_db.list_index_schedules() if hasattr(settings_db, "list_index_schedules") else []
            sched_area = ui.column().classes("w-full gap-2")

            def render_schedules() -> None:
                sched_area.clear()
                current = settings_db.list_index_schedules() if hasattr(settings_db, "list_index_schedules") else []
                with sched_area:
                    if not current:
                        ui.label("Расписаний пока нет. Нажмите «+ Добавить» чтобы создать.").classes("rag-meta")
                    for sched in current:
                        days_str = " ".join(_DAY_RU.get(d, d) for d in (sched.get("days") or []))
                        cadence_str = _CADENCE_LABELS.get(str(sched.get("cadence") or "daily"), "")
                        stage_str = _STAGE_LABELS.get(str(sched.get("stage") or "all"), str(sched.get("stage") or ""))
                        last_run = str(sched.get("last_run_at") or "—")
                        enabled_val = bool(int(sched.get("enabled") or 0))
                        color_cls = "" if enabled_val else "opacity-50"
                        with ui.row().classes(f"w-full items-center gap-2 p-2 border border-gray-200 rounded {color_cls}"):
                            ui.icon("check_circle" if enabled_val else "radio_button_unchecked").classes(
                                "text-xl " + ("text-green-500" if enabled_val else "text-gray-400"))
                            ui.label(str(sched.get("label") or "Без названия")).classes("font-semibold min-w-32")
                            ui.label(stage_str).classes("rag-chip")
                            ui.label(f"{cadence_str} в {sched.get('time') or '?'}").classes("rag-meta")
                            ui.label(days_str).classes("rag-meta min-w-28")
                            ui.space()
                            ui.label(f"Последний: {last_run[:16] if last_run != '—' else '—'}").classes("rag-meta text-xs")
                            ui.button(icon="edit", on_click=lambda s=sched: open_sched_dialog(s), color=None).props("flat dense round")
                            ui.button(icon="delete", on_click=lambda s=sched: delete_sched(str(s.get("id") or "")), color=None).props("flat dense round color=red-5")

            def delete_sched(sched_id: str) -> None:
                settings_db.delete_index_schedule(id=sched_id)
                render_schedules()
                ui.notify("Расписание удалено.", type="warning")

            def open_sched_dialog(existing: Optional[Dict[str, Any]] = None) -> None:
                with ui.dialog() as dlg, ui.card().classes("w-full max-w-lg p-4 gap-3"):
                    ui.label("Изменить расписание" if existing else "Новое расписание").classes("text-lg font-semibold")
                    dlg_label = ui.input("Название", value=str((existing or {}).get("label") or "")).props("dense outlined").classes("w-full")
                    with ui.row().classes("w-full gap-3"):
                        dlg_enabled = ui.checkbox("Включено", value=bool(int((existing or {}).get("enabled", 1))))
                        dlg_stage = ui.select(
                            _STAGE_LABELS, value=str((existing or {}).get("stage") or "all"), label="Этап"
                        ).props("dense outlined").classes("flex-1")
                    with ui.row().classes("w-full gap-3"):
                        dlg_cadence = ui.select(
                            _CADENCE_LABELS, value=str((existing or {}).get("cadence") or "daily"), label="Период"
                        ).props("dense outlined").classes("flex-1")
                        dlg_time = ui.input(
                            "Время (ЧЧ:ММ)", value=str((existing or {}).get("time") or "03:00")
                        ).props("dense outlined mask='##:##'").classes("w-32")
                    ui.label("Дни недели (для ежедневного/еженедельного):").classes("rag-meta")
                    existing_days = (existing or {}).get("days") or _DAY_LABELS[:5]
                    day_checks = {d: ui.checkbox(_DAY_RU.get(d, d), value=(d in existing_days)) for d in _DAY_LABELS}
                    with ui.row().classes("w-full gap-2 justify-end"):
                        ui.button("Отмена", on_click=dlg.close).props("flat")
                        def save_sched() -> None:
                            days_sel = [d for d, cb in day_checks.items() if cb.value]
                            settings_db.save_index_schedule(
                                id=(existing or {}).get("id"),
                                label=str(dlg_label.value or ""),
                                enabled=bool(dlg_enabled.value),
                                cadence=str(dlg_cadence.value or "daily"),
                                time=str(dlg_time.value or "03:00"),
                                days=days_sel,
                                stage=str(dlg_stage.value or "all"),
                            )
                            dlg.close()
                            render_schedules()
                            ui.notify("Расписание сохранено.", type="positive")
                        ui.button("Сохранить", icon="save", on_click=save_sched).props("unelevated")
                dlg.open()

            render_schedules()
            ui.button("+ Добавить расписание", icon="add", on_click=lambda: open_sched_dialog()).props("outline")

        # ── Настройки индекса ────────────────────────────────────────────
        with ui.column().classes("rag-card w-full p-4 gap-3"):
            with ui.row().classes("w-full items-center gap-2"):
                ui.icon("tune").classes("text-2xl text-indigo-500")
                ui.label("Параметры индексации").classes("text-xl font-semibold")
            with ui.row().classes("w-full gap-3 flex-wrap"):
                workers_input = ui.number("Потоки чтения", value=int(settings.get("workers") or state.cfg.get("index_read_workers") or 4), min=1, max=32, step=1).props("dense outlined").classes("w-40")
                max_chunks_input = ui.number("Макс. чанков на файл", value=int(settings.get("max_chunks") or state.cfg.get("index_max_chunks") or 2000), min=0, max=100000, step=100).props("dense outlined").classes("w-52")
                recreate_input = ui.checkbox("Пересоздавать коллекцию", value=bool(settings.get("recreate")))
                skip_inline_ocr_input = ui.checkbox("Пропускать OCR внутри индекса", value=bool(settings.get("skip_inline_ocr")))

            ui.separator()
            with ui.row().classes("w-full items-center gap-2"):
                ui.icon("scanner").classes("text-xl text-orange-500")
                ui.label("Настройки OCR").classes("font-semibold")
            with ui.row().classes("w-full gap-3 items-end flex-wrap"):
                ocr_enabled_input = ui.checkbox("Запускать OCR после индексации", value=bool(settings.get("ocr_enabled")))
                with ui.column().classes("gap-0"):
                    ocr_min_text_input = ui.number(
                        "Порог текста для скана (символов)",
                        value=int(settings.get("ocr_min_text_len") or 50),
                        min=1, max=100000, step=10,
                    ).props("dense outlined").classes("w-64")
                    ui.label("Если в PDF меньше указанного числа символов — файл считается сканом.").classes("rag-meta text-xs")

            def save_index_settings() -> None:
                saved = settings_db.save_index_settings({
                    "workers": int(workers_input.value or 4),
                    "max_chunks": int(max_chunks_input.value or 0),
                    "recreate": bool(recreate_input.value),
                    "skip_inline_ocr": bool(skip_inline_ocr_input.value),
                    "ocr_enabled": bool(ocr_enabled_input.value),
                    "ocr_min_text_len": int(ocr_min_text_input.value or 50),
                })
                _log_app_event(state, "index", "save_settings", details=saved)
                ui.notify("Настройки индексирования сохранены.", type="positive")
                render()

            with ui.row().classes("w-full justify-end"):
                ui.button("Сохранить настройки", icon="save", on_click=save_index_settings).props("unelevated")

        # ── OCR статус ───────────────────────────────────────────────────
        active_ocr = telemetry.get("active_ocr")
        last_ocr = telemetry.get("last_ocr")
        ocr_summary = telemetry.get("ocr_summary") or {}
        with ui.column().classes("rag-card w-full p-4 gap-2"):
            ui.label("Статус OCR").classes("text-xl font-semibold")
            if active_ocr:
                found = int(active_ocr.get("found_scanned") or 0)
                processed = int(active_ocr.get("processed_pdfs") or 0)
                pct = min(1.0, processed / found) if found > 0 else 0.0
                ui.label(f"OCR выполняется: {processed:,} / {found:,} PDF".replace(",", " ")).classes("font-semibold")
                ui.linear_progress(value=pct).classes("w-full")
                ui.label(f"Длительность: {_format_duration_seconds(active_ocr.get('duration_sec'))}").classes("rag-meta")
            elif last_ocr:
                ui.label(
                    f"Последний OCR: {last_ocr.get('status') or '-'} · найдено {int(last_ocr.get('found_scanned') or 0):,} · "
                    f"обработано {int(last_ocr.get('processed_pdfs') or 0):,} · {_format_duration_seconds(last_ocr.get('duration_sec'))}".replace(",", " ")
                ).classes("rag-meta")
            else:
                ui.label("OCR-проходов пока не было.").classes("rag-meta")
            ui.label(
                f"Средняя длительность OCR: {_format_duration_seconds(ocr_summary.get('avg_duration_sec'))} · "
                f"средне найдено: {float(ocr_summary.get('avg_found_scanned') or 0):.0f} · "
                f"средне обработано: {float(ocr_summary.get('avg_processed_pdfs') or 0):.0f}"
            ).classes("rag-meta")

        # ── Статистика по этапам + график ────────────────────────────────
        with ui.column().classes("rag-card w-full p-4 gap-3"):
            with ui.row().classes("w-full items-center gap-2"):
                ui.label("Статистика по этапам").classes("text-xl font-semibold")
            rows = telemetry.get("stage_summary") or []
            if rows:
                ui.table(
                    rows=rows,
                    columns=[
                        {"name": "stage", "label": "Этап", "field": "stage"},
                        {"name": "status", "label": "Статус", "field": "status"},
                        {"name": "processed_files", "label": "Файлов", "field": "processed_files"},
                        {"name": "added_files", "label": "Добавлено", "field": "added_files"},
                        {"name": "updated_files", "label": "Обновлено", "field": "updated_files"},
                        {"name": "error_files", "label": "Ошибок", "field": "error_files"},
                        {"name": "points_added", "label": "Точек", "field": "points_added"},
                        {"name": "last_duration_sec", "label": "Последний, сек", "field": "last_duration_sec"},
                        {"name": "avg_duration_sec", "label": "Среднее, сек", "field": "avg_duration_sec"},
                    ],
                    pagination=10,
                ).classes("w-full")
            else:
                ui.label("История этапов пока пустая.").classes("rag-meta")

            # ── График по дням ──────────────────────────────────────────
            ui.separator()
            ui.label("График индексации по дням").classes("font-semibold")
            with ui.row().classes("w-full gap-3 items-center"):
                chart_metric = ui.select(
                    {"files": "Файлов обработано", "added": "Файлов добавлено", "points": "Точек (чанков)"},
                    value="files", label="Метрика"
                ).props("dense outlined").classes("w-56")
                chart_period = ui.select(
                    {"7": "7 дней", "30": "30 дней", "90": "90 дней"},
                    value="30", label="Период"
                ).props("dense outlined").classes("w-36")
                chart_area = ui.column().classes("w-full")

            def rebuild_chart() -> None:
                chart_area.clear()
                period_days = int(chart_period.value or 30)
                metric_key = str(chart_metric.value or "files")
                daily = settings_db.get_daily_index_stats(days=period_days) if hasattr(settings_db, "get_daily_index_stats") else []
                if not daily:
                    with chart_area:
                        ui.label("Нет данных за выбранный период.").classes("rag-meta")
                    return
                # Группируем по дням, суммируем метрику по всем этапам
                from collections import defaultdict
                by_day: Dict[str, int] = defaultdict(int)
                for d in daily:
                    by_day[str(d.get("day") or "")] += int(d.get(metric_key) or 0)
                days_sorted = sorted(by_day.keys())
                values = [by_day[d] for d in days_sorted]
                metric_label = {"files": "Файлов", "added": "Добавлено", "points": "Точек"}[metric_key]
                chart_option = {
                    "tooltip": {"trigger": "axis"},
                    "xAxis": {"type": "category", "data": days_sorted, "axisLabel": {"rotate": 30}},
                    "yAxis": {"type": "value", "name": metric_label},
                    "series": [{"name": metric_label, "type": "bar", "data": values,
                                "itemStyle": {"color": "#6366f1"}}],
                    "grid": {"left": "60px", "right": "20px", "bottom": "60px"},
                }
                with chart_area:
                    ui.echart(chart_option).classes("w-full h-64")

            chart_metric.on_value_change(lambda _: rebuild_chart())
            chart_period.on_value_change(lambda _: rebuild_chart())
            rebuild_chart()

        # ── Форматы файлов ───────────────────────────────────────────────
        with ui.column().classes("rag-card w-full p-4 gap-2"):
            ui.label("Форматы файлов в индексе").classes("text-xl font-semibold")
            if not stats["found"]:
                ui.label(f"Состояние индекса не найдено: {stats['state_file']}").classes("rag-meta")
            else:
                for ext, count in list(stats.get("by_ext", {}).items())[:30]:
                    size = (stats.get("by_ext_size") or {}).get(ext, 0)
                    ui.label(f"{ext}: {count:,} · {_format_bytes(size)}".replace(",", " ")).classes("rag-meta")

    def render_index_dashboard() -> None:
        if not _is_admin(state):
            return
        stats = _read_index_stats(state.cfg)
        telemetry = _read_index_telemetry(state.cfg)
        with ui.column().classes("rag-card w-full p-4 gap-3"):
            ui.label("Дашборд индексирования").classes("text-xl font-semibold")
            if not stats["found"]:
                ui.label(f"Состояние индекса не найдено: {stats['state_file']}").classes("rag-meta")
                return
            with ui.row().classes("w-full gap-3"):
                ui.label(f"Файлов: {stats['total']:,}".replace(",", " ")).classes("rag-chip")
                ui.label(f"Размер: {_format_bytes(stats.get('total_size_bytes'))}").classes("rag-chip")
                ui.label(f"Обновлен: {stats.get('last_modified', 'неизвестно')}").classes("rag-chip")
                last_run = telemetry.get("last_run") or {}
                if last_run:
                    ui.label(f"Последний запуск: {_format_duration_seconds(last_run.get('duration_sec'))}").classes("rag-chip")
            if stats.get("by_ext"):
                for ext, count in list(stats["by_ext"].items())[:12]:
                    ui.label(f"{ext}: {count}").classes("rag-meta")

    def render_home_screen() -> None:
        """Главная — V2_Home wireframe."""
        import datetime as _dt

        display_name = str(
            (state.current_user or {}).get("display_name")
            or (state.current_user or {}).get("username")
            or ""
        )
        now = _dt.datetime.now()
        _day_names = ["понедельник", "вторник", "среда", "четверг", "пятница", "суббота", "воскресенье"]
        _month_names = ["января", "февраля", "марта", "апреля", "мая", "июня",
                        "июля", "августа", "сентября", "октября", "ноября", "декабря"]
        date_str = f"{now.day} {_month_names[now.month - 1]} {now.year} · {_day_names[now.weekday()]}"

        # ── Шапка ─────────────────────────────────────────────────────────
        with ui.row().classes("w-full items-end justify-between"):
            with ui.column().classes("gap-0"):
                ui.html(
                    f'<div class="rag-home-greeting">Доброе утро, {html.escape(display_name or "пользователь")}</div>'
                    f'<div class="rag-home-date">{html.escape(date_str)}</div>',
                    sanitize=False,
                )
            if _is_admin(state):
                ui.button("⊞ настроить главную", color=None).props("flat no-caps dense").style(
                    "font-size: 12px; color: var(--rag-accent);"
                )

        # ── KPI блок ──────────────────────────────────────────────────────
        _idx_stats = _read_index_stats(state.cfg)
        _doc_total = _idx_stats.get("total", 0)
        _size_bytes = _idx_stats.get("total_size_bytes") or 0
        telemetry_path = _telemetry_db_path(state.cfg)
        _searches_today = 0
        _active_users = 0
        try:
            rows = _db_query_dicts(
                telemetry_path,
                "SELECT COUNT(*) AS cnt FROM search_logs WHERE substr(ts,1,10) = date('now','localtime')",
            )
            _searches_today = int((rows[0] if rows else {}).get("cnt") or 0)
        except Exception:
            pass
        try:
            rows = _db_query_dicts(
                telemetry_path,
                "SELECT COUNT(DISTINCT username) AS cnt FROM search_logs WHERE ts >= datetime('now','-1 hour','localtime')",
            )
            _active_users = int((rows[0] if rows else {}).get("cnt") or 0)
        except Exception:
            pass

        _kpis = [
            ("ФАЙЛОВ В ИНДЕКСЕ", f"{_doc_total:,}".replace(",", " "), "+143 за сутки", "up"),
            ("РАЗМЕР ИНДЕКСА", _format_bytes(_size_bytes), "", "neutral"),
            ("ПОИСКОВ СЕГОДНЯ", str(_searches_today), "", "neutral"),
            ("АКТИВНЫХ ПОЛЬЗОВАТЕЛЕЙ", str(_active_users), "сейчас в системе", "neutral"),
        ]
        with ui.row().classes("w-full gap-3"):
            for label, value, delta, delta_cls in _kpis:
                with ui.element("div").classes("rag-kpi-tile"):
                    ui.html(f'<div class="rag-kpi-label">{label}</div>', sanitize=False)
                    ui.html(f'<div class="rag-kpi-value">{value}</div>', sanitize=False)
                    if delta:
                        ui.html(f'<div class="rag-kpi-delta {delta_cls}">{delta}</div>', sanitize=False)
                    ui.html(
                        '<div class="rag-kpi-sparkline">'
                        + "".join(f'<span style="height:{h}%"></span>' for h in [30, 45, 60, 80, 40, 55, 90])
                        + "</div>",
                        sanitize=False,
                    )

        # ── История + Статус индекса ───────────────────────────────────────
        with ui.row().classes("w-full gap-3 items-start"):
            # История поиска
            _recent = list(state.history[:8])
            with ui.expansion(f"🕒 Моя история поиска ({len(_recent)})", value=True).classes("rag-card flex-1 min-w-0"):
                if not _recent:
                    ui.label("История пуста.").classes("rag-meta p-2")
                else:
                    with ui.column().classes("w-full gap-0"):
                        for q in _recent:
                            with ui.row().classes("w-full items-center gap-2 px-2 py-1").style(
                                "border-bottom: 1px solid var(--rag-border);"
                            ):
                                ui.icon("history", size="14px").style("color: var(--rag-ink-4);")
                                btn = ui.button(q, on_click=lambda _q=q: (
                                    setattr(state, "query_input_value", _q),
                                    setattr(state, "screen", "search"),
                                    render()
                                )).props("flat dense no-caps align=left")
                                btn.style("font-size: 13px; color: var(--rag-ink-2); flex: 1; justify-content: flex-start;")

            # Статус индекса (от телеметрии)
            _telemetry = _read_index_telemetry(state.cfg)
            # Используем stage_summary (последний завершённый прогон каждого
            # этапа), а active_stages накладываем сверху для running.
            _active = _telemetry.get("active_stages") or []
            _summary = _telemetry.get("stage_summary") or []
            _stages_map: Dict[str, Dict[str, Any]] = {}
            for _r in _summary:
                _stages_map[str(_r.get("stage") or "")] = dict(_r)
            for _r in _active:
                _stages_map[str(_r.get("stage") or "")] = dict(_r)
            _stages_data = list(_stages_map.values())

            with ui.expansion("⚙ Состояние индекса", value=True).classes("rag-card flex-1 min-w-0"):
                if not _stages_data:
                    ui.label("Данных нет — индексация ещё не запускалась.").classes("rag-meta p-2")
                else:
                    # Краткий pipeline
                    _stage_order = ["metadata", "small", "large", "content", "OCR"]
                    _stage_status_map: Dict[str, str] = {}
                    for row in _stages_data:
                        _stage_status_map[str(row.get("stage") or "")] = str(row.get("status") or "pending")

                    with ui.element("div").classes("rag-pipeline gap-1 px-2 py-2"):
                        for i, stage_name in enumerate(_stage_order):
                            st = _stage_status_map.get(stage_name, "pending")
                            _dot_colors = {
                                "done": "var(--rag-ok)",
                                "completed": "var(--rag-ok)",
                                "running": "var(--rag-accent)",
                                "failed": "var(--rag-danger)",
                                "error": "var(--rag-danger)",
                                "pending": "var(--rag-border-strong)",
                            }
                            _dot_icons = {"done": "✓", "completed": "✓", "running": "◷", "failed": "✗", "error": "✗"}
                            dot_color = _dot_colors.get(st, _dot_colors["pending"])
                            dot_icon = _dot_icons.get(st, str(i + 1))
                            with ui.element("div").classes("rag-pipeline-step"):
                                ui.html(
                                    f'<div class="rag-pipeline-dot" style="background:{dot_color}">{dot_icon}</div>'
                                    f'<div class="rag-pipeline-label">{stage_name}</div>',
                                    sanitize=False,
                                )
                            if i < len(_stage_order) - 1:
                                ui.element("div").classes("rag-pipeline-connector")

                    # Найти текущий running этап
                    _running = next((r for r in _stages_data if str(r.get("status") or "") == "running"), None)
                    if _running:
                        processed = int(_running.get("processed_files") or 0)
                        total_f = int(_running.get("total_files") or 0)
                        pct = min(1.0, processed / total_f) if total_f > 0 else 0.0
                        with ui.column().classes("gap-1 px-2 pb-2"):
                            stage_label = _STAGE_LABELS.get(str(_running.get("stage") or ""), str(_running.get("stage") or ""))
                            ui.label(f"{stage_label} · {processed:,}/{total_f:,} файлов".replace(",", " ")).classes("rag-meta")
                            ui.linear_progress(value=pct, color="primary").classes("w-full")
                    else:
                        _last = _stages_data[0] if _stages_data else {}
                        _last_ts = str(
                            _last.get("ts_finished")
                            or _last.get("finished_at")
                            or _last.get("ts_started")
                            or _last.get("started_at")
                            or "—"
                        )[:16].replace("T", " ")
                        ui.label(f"Последний запуск: {_last_ts}").classes("rag-meta px-2 pb-2")

        # ── Виджеты ────────────────────────────────────────────────────────
        with ui.row().classes("w-full gap-3 items-start"):
            # Задачи
            with ui.expansion("📋 Мои задачи", value=True).classes("rag-card flex-1 min-w-0"):
                with ui.column().classes("w-full gap-2 p-1"):
                    # Статические задачи-подсказки
                    _tasks_data = []
                    try:
                        _err_rows = _db_query_dicts(
                            telemetry_path,
                            """SELECT SUM(error_files) AS errs FROM index_runs
                               WHERE ts_finished >= datetime('now','-2 days','localtime')""",
                        )
                        _errs = int((_err_rows[0] if _err_rows else {}).get("errs") or 0)
                        if _errs > 0:
                            _tasks_data.append((f"Проверить ошибки индексации ({_errs})", "сегодня", "var(--rag-danger)"))
                    except Exception:
                        pass
                    try:
                        _reg_rows = _db_query_dicts(
                            _auth_db_path(state.cfg),
                            "SELECT COUNT(*) AS cnt FROM registration_requests WHERE status='pending'",
                        )
                        _regs = int((_reg_rows[0] if _reg_rows else {}).get("cnt") or 0)
                        if _regs > 0:
                            _tasks_data.append((f"{_regs} заявок на доступ ждут", "сегодня", "var(--rag-warn)"))
                    except Exception:
                        pass
                    if not _tasks_data:
                        ui.label("Активных задач нет.").classes("rag-meta")
                    for task_title, task_date, task_color in _tasks_data:
                        with ui.row().classes("w-full items-center gap-2"):
                            ui.icon("radio_button_unchecked", size="14px").style(f"color: {task_color};")
                            with ui.column().classes("flex-1 gap-0"):
                                ui.label(task_title).style("font-size: 13px;")
                                ui.label(task_date).classes("rag-meta")

            # Недавно открытые файлы
            _fav_paths = []
            try:
                _fav_paths = _get_auth_db(state).list_favorites(username=_username(state), limit=6)
            except Exception:
                pass

            with ui.expansion(f"📁 Недавние файлы ({len(_fav_paths)})", value=True).classes("rag-card flex-1 min-w-0"):
                if not _fav_paths:
                    ui.label("Нет недавних файлов.").classes("rag-meta p-2")
                else:
                    with ui.column().classes("w-full gap-1 p-1"):
                        for fp in _fav_paths[:6]:
                            p = Path(fp)
                            ext = p.suffix.lower()
                            _ext_colors = {
                                ".pdf": "var(--rag-danger)",
                                ".docx": "var(--rag-accent)",
                                ".doc": "var(--rag-accent)",
                                ".xlsx": "var(--rag-ok)",
                                ".xls": "var(--rag-ok)",
                            }
                            dot_color = _ext_colors.get(ext, "var(--rag-muted)")
                            with ui.row().classes("w-full items-center gap-2"):
                                ui.element("div").style(
                                    f"width:8px;height:8px;border-radius:50%;background:{dot_color};flex-shrink:0;"
                                )
                                btn = ui.button(
                                    p.name,
                                    on_click=lambda _fp=fp: go_explorer(_fp),
                                ).props("flat dense no-caps align=left")
                                btn.style("font-size: 12px; color: var(--rag-ink-2); flex: 1; justify-content: flex-start;")
                                btn.tooltip(fp)

    def render_login_screen() -> None:
        auth_db = _get_auth_db(state)

        tg_login_token: Dict[str, str] = {"value": ""}
        reg_visible: Dict[str, bool] = {"v": False}

        # ── логика аутентификации ────────────────────────────────────────────
        def _complete_login(user: Dict[str, Any], *, event_type: str) -> None:
            state.current_user = user
            state.auth_token = auth_db.create_session(username=str(user.get("username") or ""))
            auth_db.log_auth_event(username=_username(state), event_type=event_type, ok=True)
            _load_user_state(state)
            try:
                app.storage.user["auth_token"] = state.auth_token
            except Exception:
                pass
            ui.notify("Вход выполнен.", type="positive")
            render()

        def login() -> None:
            username = str(username_input.value or "")
            user = auth_db.login(username=username, password=str(password_input.value or ""))
            if not user:
                auth_db.log_auth_event(username=username, event_type="login_failed", ok=False, error="bad_credentials")
                ui.notify("Неверный логин или пароль.", type="negative")
                return
            _complete_login(user, event_type="login")

        def request_tg_login() -> None:
            bot_link = str(state.cfg.get("telegram_bot_link") or "").strip()
            if not bot_link:
                ui.notify("Telegram-вход не настроен: задайте telegram_bot_link в config.json.", type="warning")
                return
            out = auth_db.create_telegram_login_challenge(target="web")
            token = str(out.get("token") or "")
            link = _telegram_deeplink(bot_link, "login", token)
            if not token or not link:
                ui.notify("Не удалось создать Telegram-ссылку входа.", type="negative")
                return
            tg_login_token["value"] = token
            ui.run_javascript(
                "(() => {"
                f"const url = {json.dumps(link)};"
                "const w = window.open(url, '_blank', 'noopener,noreferrer');"
                "if (!w) { window.location.href = url; }"
                "})();"
            )
            ui.notify("Подтвердите вход в Telegram, затем вернитесь в браузер.", type="positive")

        def poll_tg_login() -> None:
            token = tg_login_token["value"]
            if not token or state.current_user is not None:
                return
            out = auth_db.consume_confirmed_telegram_login(token=token)
            if not out.get("ok"):
                return
            tg_login_token["value"] = ""
            user = out.get("user") or auth_db.get_user(username=str(out.get("username") or ""))
            if not user:
                return
            _complete_login(user, event_type="telegram_web_login")

        def register_request() -> None:
            username = str(reg_username_input.value or "").strip().lower()
            display_name = str(reg_display_input.value or "").strip()
            tg_username = str(reg_tg_user_input.value or "").strip().lstrip("@")
            if len(username) < 3:
                ui.notify("Укажите логин (минимум 3 символа).", type="warning")
                return
            if auth_db.get_user(username=username):
                ui.notify("Пользователь с таким логином уже существует. Используйте вход.", type="warning")
                return
            out = auth_db.create_registration_request(
                username=username,
                display_name=display_name or username,
                telegram_username=tg_username,
                source="web",
                note="requested from web login form",
            )
            if not out.get("ok"):
                ui.notify("Не удалось отправить заявку. Попробуйте позже.", type="negative")
                return
            ui.notify("Заявка отправлена администратору.", type="positive")
            reg_username_input.value = ""
            reg_display_input.value = ""
            reg_tg_user_input.value = ""
            reg_username_input.update()
            reg_display_input.update()
            reg_tg_user_input.update()

        # ── данные для брендовой панели ──────────────────────────────────────
        _stats = _read_index_stats(state.cfg)
        _doc_count = _stats.get("total", 0)
        _doc_count_str = f"{_doc_count:,}".replace(",", " ") if _doc_count else "—"

        _MARK_SVG = (
            '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 1024 1024">'
            '<path d="M210 332h575a28 28 0 0 1 28 28v270H210z" fill="#F2B625"/>'
            '<path d="M210 268h185a28 28 0 0 1 24 14l34 58H210z" fill="#F2B625"/>'
            '<path d="M250 356h522v228H250z" fill="none" stroke="#FFFFFF" stroke-width="28" stroke-linejoin="round"/>'
            '<path d="M360 145h242l53 53v367H360z" fill="#F7FAFC" stroke="#5F6B78" stroke-width="10"/>'
            '<path d="M602 145v53h53" fill="none" stroke="#5F6B78" stroke-width="10"/>'
            '<path d="M410 245h170M410 300h194M410 355h194M410 410h194M410 465h194" stroke="#81909E" stroke-width="14"/>'
            '<circle cx="436" cy="236" r="35" fill="#D98B00"/>'
            '<circle cx="436" cy="236" r="14" fill="#FFFFFF"/>'
            '<path d="M390 470l64-66 72 31M390 470l-34 70M450 548h190M505 372h55v60h-55z"'
            ' fill="none" stroke="#25313D" stroke-width="18" stroke-linecap="round" stroke-linejoin="round"/>'
            '<circle cx="696" cy="562" r="137" fill="#FFFFFF" stroke="#157FC4" stroke-width="30"/>'
            '<path d="M790 660l112 112" stroke="#157FC4" stroke-width="54" stroke-linecap="round"/>'
            '<path d="M636 532a62 62 0 0 1 66-55" fill="none" stroke="#5F6B78"'
            ' stroke-width="18" stroke-linecap="round"/>'
            '</svg>'
        )

        # ── split layout ─────────────────────────────────────────────────────
        with ui.element("div").classes("rag-login-root"):

            # Левая тёмная панель
            with ui.element("aside").classes("rag-login-brand"):
                ui.html(
                    f'<div class="rag-login-brand-inner">'
                    f'  <div class="rag-login-mark-wrap">{_MARK_SVG}</div>'
                    f'  <div class="rag-login-brand-name">ТЕХНОПОИСК</div>'
                    f'  <div class="rag-login-brand-sub">Поиск документов и данных</div>'
                    f'  <div class="rag-login-stats">'
                    f'    <div class="rag-login-stat-val">{_doc_count_str}</div>'
                    f'    <div class="rag-login-stat-lbl">документов в индексе</div>'
                    f'  </div>'
                    f'</div>'
                    f'<div class="rag-login-brand-footer">'
                    f'  <div class="rag-login-badge">'
                    f'    <div class="rag-login-badge-dot"></div>'
                    f'    Индекс активен'
                    f'  </div>'
                    f'</div>'
                )

            # Правая форма
            with ui.element("main").classes("rag-login-form-side"):
                with ui.element("div").classes("rag-login-form-inner"):
                    ui.html(
                        '<div class="rag-login-greeting">Добро пожаловать</div>'
                        '<div class="rag-login-greeting-sub">Войдите в аккаунт для доступа к системе</div>'
                    )

                    username_input = (
                        ui.input("Логин")
                        .props("dense outlined")
                        .classes("w-full rag-login-input-gap")
                    )
                    password_input = (
                        ui.input("Пароль", password=True, password_toggle_button=True)
                        .props("dense outlined")
                        .classes("w-full")
                    )
                    password_input.on("keyup.enter", lambda _: login())

                    ui.button("Войти", on_click=login).props("unelevated").classes(
                        "w-full rag-login-btn-primary"
                    )

                    ui.html(
                        '<div class="rag-login-divider">'
                        '  <div class="rag-login-divider-line"></div>'
                        '  <span class="rag-login-divider-text">или</span>'
                        '  <div class="rag-login-divider-line"></div>'
                        '</div>'
                    )

                    ui.button("Войти через Telegram", icon="send", on_click=request_tg_login).props(
                        "outline"
                    ).classes("w-full")

                    # Заявка на доступ
                    reg_section = ui.element("div").classes("rag-login-reg-wrap")
                    reg_section.set_visibility(False)

                    with ui.row().classes("w-full justify-center items-center gap-1 q-mt-sm"):
                        ui.label("Нет аккаунта?").style(
                            "color: var(--rag-muted); font-size: 13px;"
                        )

                        def _toggle_reg() -> None:
                            reg_visible["v"] = not reg_visible["v"]
                            reg_section.set_visibility(reg_visible["v"])

                        ui.button("Запросить доступ", on_click=_toggle_reg).props("flat dense").style(
                            "color: var(--rag-accent); font-size: 13px; text-transform: none;"
                        )

                    with reg_section:
                        ui.html('<div class="rag-login-reg-title">Заявка на доступ</div>')
                        reg_username_input = (
                            ui.input("Логин")
                            .props("dense outlined")
                            .classes("w-full rag-login-input-gap")
                        )
                        reg_display_input = (
                            ui.input("Имя")
                            .props("dense outlined")
                            .classes("w-full rag-login-input-gap")
                        )
                        reg_tg_user_input = (
                            ui.input("Telegram (необязательно)")
                            .props("dense outlined prefix=@")
                            .classes("w-full")
                        )
                        ui.button(
                            "Отправить заявку", icon="how_to_reg", on_click=register_request
                        ).props("unelevated").classes("w-full q-mt-sm")
                        ui.label(
                            "После одобрения администратором вы получите доступ."
                        ).classes("rag-meta q-mt-xs")

        ui.timer(2.0, poll_tg_login)

    def render_admin_users(auth_db: UserAuthDB) -> None:
        with ui.column().classes("rag-card w-full p-4 gap-4"):
            ui.label("Админ-панель пользователей").classes("text-xl font-semibold")
            with ui.expansion("Создать пользователя", icon="person_add").classes("w-full"):
                new_username = ui.input("Логин").props("dense outlined").classes("w-full")
                new_display = ui.input("Имя").props("dense outlined").classes("w-full")
                new_telegram = ui.input("Telegram chat id").props("dense outlined").classes("w-full")
                new_telegram_username = ui.input("Telegram username").props("dense outlined prefix=@").classes("w-full")
                new_password = ui.input("Временный пароль", password=True, password_toggle_button=True).props("dense outlined").classes("w-full")
                new_role = ui.select(["user", "admin"], value="user", label="Роль").props("dense outlined").classes("w-full")
                new_status = ui.select(["active", "pending", "blocked"], value="active", label="Статус").props("dense outlined").classes("w-full")
                new_must_change = ui.checkbox("Потребовать смену пароля", value=True)

                def create_user() -> None:
                    ok = auth_db.admin_create_user(
                        username=str(new_username.value or ""),
                        display_name=str(new_display.value or ""),
                        telegram_chat_id=str(new_telegram.value or ""),
                        telegram_username=str(new_telegram_username.value or ""),
                        password=str(new_password.value or ""),
                        role=str(new_role.value or "user"),
                        status=str(new_status.value or "active"),
                        must_change_password=bool(new_must_change.value),
                    )
                    ui.notify("Пользователь создан." if ok else "Не удалось создать пользователя.", type="positive" if ok else "negative")
                    render()

                ui.button("Создать", icon="person_add", on_click=create_user).props("unelevated")

            users = auth_db.list_users()
            for user in users:
                username = str(user.get("username") or "")
                role = str(user.get("role") or "user")
                status = str(user.get("status") or "")
                with ui.expansion(f"{username} · {role} · {status}", icon="person").classes("w-full"):
                    display_input = ui.input("Имя", value=str(user.get("display_name") or "")).props("dense outlined").classes("w-full")
                    telegram_input = ui.input("Telegram chat id", value=str(user.get("telegram_chat_id") or "")).props("dense outlined").classes("w-full")
                    telegram_username_input = ui.input("Telegram username", value=str(user.get("telegram_username") or "")).props("dense outlined prefix=@").classes("w-full")
                    role_input = ui.select(["user", "admin"], value=role, label="Роль").props("dense outlined").classes("w-full")
                    status_input = ui.select(["active", "pending", "blocked"], value=status or "active", label="Статус").props("dense outlined").classes("w-full")
                    must_input = ui.checkbox("Потребовать смену пароля", value=bool(int(user.get("must_change_password") or 0)))
                    reset_password = ui.input("Новый временный пароль", password=True, password_toggle_button=True).props("dense outlined").classes("w-full")

                    def save_user(
                        username: str = username,
                        display_input: Any = display_input,
                        telegram_input: Any = telegram_input,
                        telegram_username_input: Any = telegram_username_input,
                        role_input: Any = role_input,
                        status_input: Any = status_input,
                        must_input: Any = must_input,
                    ) -> None:
                        ok = auth_db.admin_update_user(
                            username=username,
                            display_name=str(display_input.value or ""),
                            telegram_chat_id=str(telegram_input.value or ""),
                            telegram_username=str(telegram_username_input.value or ""),
                            role=str(role_input.value or "user"),
                            status=str(status_input.value or "active"),
                            must_change_password=bool(must_input.value),
                        )
                        ui.notify("Пользователь обновлен." if ok else "Не удалось обновить пользователя.", type="positive" if ok else "negative")
                        _refresh_current_user(state)
                        render()

                    def set_password(
                        username: str = username,
                        reset_password: Any = reset_password,
                    ) -> None:
                        ok = auth_db.admin_set_password(
                            username=username,
                            new_password=str(reset_password.value or ""),
                            must_change_password=True,
                        )
                        ui.notify("Пароль обновлен." if ok else "Введите новый пароль.", type="positive" if ok else "warning")
                        render()

                    with ui.row().classes("gap-2"):
                        ui.button("Сохранить", icon="save", on_click=save_user).props("outline")
                        ui.button("Сбросить пароль", icon="key", on_click=set_password).props("outline")
                        def make_invite(
                            username: str = username,
                            display_input: Any = display_input,
                            telegram_username_input: Any = telegram_username_input,
                        ) -> None:
                            bot_link = str(state.cfg.get("telegram_bot_link") or "").strip()
                            if not bot_link:
                                ui.notify("В config.json не задан telegram_bot_link.", type="warning")
                                return
                            out = auth_db.create_telegram_token(
                                purpose="invite",
                                username=username,
                                display_name=str(display_input.value or ""),
                                telegram_username=str(telegram_username_input.value or ""),
                                created_by=_username(state),
                                ttl_minutes=7 * 24 * 60,
                            )
                            link = _telegram_deeplink(bot_link, "invite", str(out.get("token") or ""))
                            ui.notify(f"Invite-link: {link}", type="positive", timeout=12000)

                        ui.button("Invite Telegram", icon="link", on_click=make_invite).props("outline")

    def render_admin_telegram_chats(auth_db: UserAuthDB) -> None:
        rows = auth_db.list_telegram_chats()
        with ui.column().classes("rag-card w-full p-4 gap-3"):
            ui.label("Telegram чаты").classes("text-xl font-semibold")
            if not rows:
                ui.label("Привязанных Telegram chat_id пока нет.").classes("rag-meta")
                return
            ui.table(
                rows=rows,
                columns=[
                    {"name": "username", "label": "Пользователь", "field": "username"},
                    {"name": "display_name", "label": "Имя", "field": "display_name"},
                    {"name": "role", "label": "Роль", "field": "role"},
                    {"name": "status", "label": "Статус", "field": "status"},
                    {"name": "telegram_chat_id", "label": "Chat ID", "field": "telegram_chat_id"},
                    {"name": "last_telegram_event_at", "label": "Последнее Telegram-событие", "field": "last_telegram_event_at"},
                    {"name": "last_login_at", "label": "Последний web-вход", "field": "last_login_at"},
                ],
                pagination=10,
            ).classes("w-full")

    def render_admin_registration_requests(auth_db: UserAuthDB) -> None:
        rows = auth_db.list_registration_requests(status="pending", limit=50)
        with ui.column().classes("rag-card w-full p-4 gap-3"):
            ui.label("Заявки на регистрацию").classes("text-xl font-semibold")
            if not rows:
                ui.label("Ожидающих заявок нет.").classes("rag-meta")
                return
            for row in rows:
                req_id = int(row.get("id") or 0)
                title = str(row.get("username") or row.get("display_name") or f"заявка {req_id}")
                tg = str(row.get("telegram_username") or row.get("telegram_chat_id") or "")
                with ui.row().classes("w-full items-center gap-2"):
                    ui.label(f"#{req_id}").classes("rag-chip")
                    ui.label(title).classes("font-medium")
                    ui.label(f"Telegram: {tg or '-'}").classes("rag-meta flex-1")
                    ui.label(str(row.get("source") or "")).classes("rag-meta")

                    def approve(req_id: int = req_id) -> None:
                        out = auth_db.review_registration_request(
                            request_id=req_id,
                            reviewed_by=_username(state),
                            decision="approved",
                        )
                        ui.notify(
                            f"Заявка одобрена: {out.get('username')}" if out.get("ok") else f"Не удалось одобрить: {out.get('reason')}",
                            type="positive" if out.get("ok") else "negative",
                        )
                        render()

                    def reject(req_id: int = req_id) -> None:
                        out = auth_db.review_registration_request(
                            request_id=req_id,
                            reviewed_by=_username(state),
                            decision="rejected",
                        )
                        ui.notify(
                            "Заявка отклонена." if out.get("ok") else f"Не удалось отклонить: {out.get('reason')}",
                            type="positive" if out.get("ok") else "negative",
                        )
                        render()

                    ui.button("Одобрить", icon="check", on_click=approve).props("outline dense")
                    ui.button("Отклонить", icon="close", on_click=reject).props("flat dense")

    def render_admin_security_settings(auth_db: UserAuthDB) -> None:
        current_ttl = auth_db.get_session_ttl_days()
        current_show_system = auth_db.get_show_system_files_for_admin()
        with ui.column().classes("rag-card w-full p-4 gap-3"):
            ui.label("Безопасность").classes("text-xl font-semibold")
            ui.label("Максимальная длительность новой сессии пользователя. Допустимый диапазон: 1-7 дней.").classes("rag-meta")
            ttl_input = ui.number(
                "Срок сессии, дней",
                value=current_ttl,
                min=1,
                max=7,
                step=1,
            ).props("dense outlined").classes("w-full max-w-xs")
            show_system_input = ui.checkbox(
                "Показывать служебные файлы администратору",
                value=current_show_system,
            )
            ui.label("Обычные пользователи служебные файлы не видят независимо от этой настройки.").classes("rag-meta")

            def save_session_ttl() -> None:
                saved = auth_db.set_session_ttl_days(int(ttl_input.value or current_ttl))
                show_system = auth_db.set_show_system_files_for_admin(bool(show_system_input.value))
                _log_app_event(
                    state,
                    "settings",
                    "security",
                    details={"session_ttl_days": saved, "show_system_files_for_admin": show_system},
                )
                ui.notify(f"Сохранено: сессии {saved} дн., служебные файлы {'видны админу' if show_system else 'скрыты'}.", type="positive")
                render()

            ui.button("Сохранить настройки безопасности", icon="save", on_click=save_session_ttl).props("outline")

    def render_admin_path_settings() -> None:
        with ui.column().classes("rag-card w-full p-4 gap-3"):
            ui.label("Пути и подключение").classes("text-xl font-semibold")
            ui.label("Эти настройки видны только администратору. После сохранения поиск переподключается к Qdrant с новыми значениями.").classes("rag-meta")
            catalog_input = ui.input("Каталог документов", value=str(state.cfg.get("catalog_path") or "")).props("dense outlined").classes("w-full")
            qdrant_url_input = ui.input("Qdrant URL", value=str(state.cfg.get("qdrant_url") or "")).props("dense outlined").classes("w-full")
            qdrant_db_input = ui.input("Локальный путь Qdrant", value=str(state.cfg.get("qdrant_db_path") or "")).props("dense outlined").classes("w-full")
            collection_input = ui.input("Коллекция", value=str(state.cfg.get("collection_name") or "catalog")).props("dense outlined").classes("w-full")
            telemetry_input = ui.input("БД телеметрии", value=str(state.cfg.get("telemetry_db_path") or "")).props("dense outlined").classes("w-full")
            log_input = ui.input("Лог автоматизации", value=str(state.cfg.get("log_file") or "")).props("dense outlined").classes("w-full")

            with ui.row().classes("w-full gap-2"):
                ui.label(f"Текущий каталог: {state.cfg.get('catalog_path') or '-'}").classes("rag-path")
                ui.label(f"Текущий Qdrant: {state.cfg.get('qdrant_url') or state.cfg.get('qdrant_db_path') or '-'}").classes("rag-path")

            def save_paths() -> None:
                new_catalog = str(catalog_input.value or "").strip()
                if new_catalog and not Path(new_catalog).exists():
                    ui.notify("Каталог документов не найден. Проверьте путь.", type="negative")
                    return
                new_qdrant_url = str(qdrant_url_input.value or "").strip()
                new_qdrant_db = str(qdrant_db_input.value or "").strip()
                if not new_qdrant_url and not new_qdrant_db:
                    ui.notify("Укажите Qdrant URL или локальный путь Qdrant.", type="warning")
                    return
                try:
                    state.cfg = _save_config_patch({
                        "catalog_path": new_catalog,
                        "qdrant_url": new_qdrant_url,
                        "qdrant_db_path": new_qdrant_db,
                        "collection_name": str(collection_input.value or "catalog").strip() or "catalog",
                        "telemetry_db_path": str(telemetry_input.value or "").strip(),
                        "log_file": str(log_input.value or "").strip(),
                    })
                    state.searcher = None
                    state.searcher_error = ""
                    state.telemetry = None
                    _log_app_event(state, "settings", "save_paths", details={key: state.cfg.get(key) for key in CONFIG_PATH_KEYS})
                    ui.notify("Пути сохранены.", type="positive")
                    render()
                except Exception as exc:
                    ui.notify(f"Не удалось сохранить пути: {exc}", type="negative")

            ui.button("Сохранить пути", icon="save", on_click=save_paths).props("outline")

    def render_admin_llm_settings() -> None:
        def _fetch_ollama_models(ollama_url: str) -> List[str]:
            """Запросить список моделей из Ollama /api/tags. Возвращает [] при ошибке."""
            try:
                import urllib.request as _ur  # noqa: PLC0415
                import json as _json  # noqa: PLC0415
                req = _ur.Request(f"{ollama_url.rstrip('/')}/api/tags", method="GET")
                with _ur.urlopen(req, timeout=4) as resp:
                    data = _json.loads(resp.read().decode())
                return sorted(m["name"] for m in (data.get("models") or []) if m.get("name"))
            except Exception:
                return []

        current_url = str(state.cfg.get("ollama_url") or "http://localhost:11434")
        current_expand = str(state.cfg.get("llm_expand_model") or "phi3:mini")
        current_rag = str(state.cfg.get("llm_rag_model") or "qwen3:8b")

        # Подтягиваем модели сразу при рендере
        available_models = _fetch_ollama_models(current_url)
        # Гарантируем, что текущие значения есть в списке даже если Ollama недоступен
        for m in [current_expand, current_rag]:
            if m and m not in available_models:
                available_models.insert(0, m)
        if not available_models:
            available_models = [current_expand, current_rag]

        with ui.column().classes("rag-card w-full p-4 gap-3"):
            ui.label("Нейросеть (LLM)").classes("text-xl font-semibold")
            ui.label(
                "Используется Ollama, запущенный локально. "
                "Включите, чтобы получать ответ ИИ по документам и автоматически расширять запросы."
            ).classes("rag-meta")

            llm_toggle = ui.switch(
                "Включить ИИ-ответ и расширение запроса",
                value=bool(state.cfg.get("llm_enabled")),
            )
            ollama_url_input = ui.input(
                "Ollama URL",
                value=current_url,
            ).props("dense outlined").classes("w-full")

            status_label = ui.label(
                f"Найдено моделей: {len(available_models)}" if available_models else "Ollama недоступен — список пуст"
            ).classes("rag-meta text-sm")

            expand_select = ui.select(
                label="Модель расширения запроса (быстрая, лёгкая)",
                options=available_models,
                value=current_expand,
                with_input=True,
            ).props("dense outlined").classes("w-full")

            rag_select = ui.select(
                label="Модель RAG Q&A (умная, для анализа документов)",
                options=available_models,
                value=current_rag,
                with_input=True,
            ).props("dense outlined").classes("w-full")

            async def refresh_models() -> None:
                url = str(ollama_url_input.value or "http://localhost:11434").strip()
                models = await run.io_bound(_fetch_ollama_models, url)
                for m in [str(expand_select.value or ""), str(rag_select.value or "")]:
                    if m and m not in models:
                        models.insert(0, m)
                if not models:
                    status_label.set_text("Ollama недоступен или нет установленных моделей")
                    ui.notify("Ollama не отвечает по адресу: " + url, type="warning")
                    return
                expand_select.options = models
                rag_select.options = models
                expand_select.update()
                rag_select.update()
                status_label.set_text(f"Найдено моделей: {len(models)}")
                ui.notify(f"Обновлено: {len(models)} моделей", type="positive")

            ui.button("Обновить список моделей", icon="refresh", on_click=refresh_models).props("flat dense")

            def save_llm_settings() -> None:
                try:
                    cfg = load_config()
                    cfg["llm_enabled"] = bool(llm_toggle.value)
                    cfg["ollama_url"] = str(ollama_url_input.value or "http://localhost:11434").strip()
                    cfg["llm_expand_model"] = str(expand_select.value or "phi3:mini").strip()
                    cfg["llm_rag_model"] = str(rag_select.value or "qwen3:8b").strip()
                    save_config(cfg)
                    state.cfg = cfg
                    _log_app_event(state, "settings", "save_llm", details={
                        "llm_enabled": cfg["llm_enabled"],
                        "ollama_url": cfg["ollama_url"],
                    })
                    ui.notify("Настройки нейросети сохранены.", type="positive")
                except Exception as exc:
                    ui.notify(f"Не удалось сохранить: {exc}", type="negative")

            ui.button("Сохранить настройки нейросети", icon="save", on_click=save_llm_settings).props("outline")

    def render_admin_search_aliases() -> None:
        telemetry = _get_telemetry(state)
        with ui.column().classes("rag-card w-full p-4 gap-3"):
            ui.label("Синонимы поиска").classes("text-xl font-semibold")
            ui.label(
                "Группы расширяют запросы без переиндексации: например, «реквизиты» ищет карточки предприятия и расчетные счета."
            ).classes("rag-meta")

            groups = telemetry.list_search_alias_groups() if hasattr(telemetry, "list_search_alias_groups") else []
            with ui.expansion("Добавить группу", icon="add", value=False).classes("w-full"):
                new_key = ui.input("Ключ группы", placeholder="company_card").props("dense outlined").classes("w-full")
                new_label = ui.input("Название", placeholder="Карточка предприятия").props("dense outlined").classes("w-full")
                new_aliases = ui.textarea("Синонимы, по одному на строку").props("dense outlined autogrow").classes("w-full")
                new_negative = ui.textarea("Исключения, по одному на строку").props("dense outlined autogrow").classes("w-full")

                def add_group() -> None:
                    label = str(new_label.value or "").strip()
                    key = str(new_key.value or label).strip()
                    aliases = [x.strip() for x in str(new_aliases.value or "").splitlines() if x.strip()]
                    negatives = [x.strip() for x in str(new_negative.value or "").splitlines() if x.strip()]
                    try:
                        telemetry.save_search_alias_group(key=key, label=label or key, aliases=aliases, negative_aliases=negatives)
                        _log_app_event(state, "settings", "search_alias_add", details={"key": key, "label": label})
                        ui.notify("Группа синонимов добавлена.", type="positive")
                        render()
                    except Exception as exc:
                        ui.notify(f"Не удалось сохранить: {exc}", type="negative")

                ui.button("Добавить группу", icon="save", on_click=add_group).props("outline")

            for group in groups:
                group_key = str(group.get("key") or "")
                alias_text = "\n".join(str(a.get("alias") or "") for a in group.get("aliases") or [])
                negative_text = "\n".join(str(x) for x in group.get("negative_aliases") or [])
                with ui.expansion(str(group.get("label") or group_key), icon="travel_explore", value=False).classes("w-full"):
                    label_input = ui.input("Название", value=str(group.get("label") or "")).props("dense outlined").classes("w-full")
                    aliases_input = ui.textarea("Синонимы", value=alias_text).props("dense outlined autogrow").classes("w-full")
                    negative_input = ui.textarea("Исключения", value=negative_text).props("dense outlined autogrow").classes("w-full")
                    ui.label(f"Ключ: {group_key} · обновлено: {group.get('updated_at') or '-'}").classes("rag-meta")

                    def save_group(
                        key: str = group_key,
                        label_ref: Any = label_input,
                        aliases_ref: Any = aliases_input,
                        negative_ref: Any = negative_input,
                    ) -> None:
                        aliases = [x.strip() for x in str(aliases_ref.value or "").splitlines() if x.strip()]
                        negatives = [x.strip() for x in str(negative_ref.value or "").splitlines() if x.strip()]
                        telemetry.save_search_alias_group(
                            key=key,
                            label=str(label_ref.value or key),
                            aliases=aliases,
                            negative_aliases=negatives,
                        )
                        _log_app_event(state, "settings", "search_alias_save", details={"key": key})
                        ui.notify("Синонимы сохранены.", type="positive")
                        render()

                    def delete_group(key: str = group_key) -> None:
                        telemetry.delete_search_alias_group(key=key)
                        _log_app_event(state, "settings", "search_alias_delete", details={"key": key})
                        ui.notify("Группа удалена.", type="positive")
                        render()

                    with ui.row().classes("gap-2"):
                        ui.button("Сохранить", icon="save", on_click=save_group).props("outline dense")
                        ui.button("Удалить", icon="delete", on_click=delete_group).props("flat dense")

            candidates = telemetry.suggest_search_alias_candidates(limit=12) if hasattr(telemetry, "suggest_search_alias_candidates") else []
            with ui.expansion("Кандидаты из истории поиска", icon="psychology", value=False).classes("w-full"):
                if not candidates:
                    ui.label("Пока нет кандидатов. Они появятся после положительных реакций на результаты поиска.").classes("rag-meta")
                for item in candidates:
                    with ui.row().classes("w-full items-center gap-2"):
                        ui.label(str(item.get("candidate") or "")).classes("font-medium")
                        ui.label(f"запрос: {item.get('query') or ''}").classes("rag-meta")
                        ui.label(str(item.get("title") or item.get("path") or "")).classes("rag-path flex-1")

    def render_settings_screen() -> None:
        auth_db = _get_auth_db(state)

        # ── Форма входа (без боковой панели) ────────────────────────────
        if state.current_user is None:
            ui.label("Настройки").classes("text-2xl font-semibold")
            with ui.column().classes("rag-card w-full max-w-xl p-4 gap-3"):
                ui.label("Вход пользователя").classes("text-xl font-semibold")
                ui.label("Для первого входа администратора используйте admin / admin, затем смените пароль.").classes("rag-meta")
                username_input = ui.input("Логин").props("dense outlined").classes("w-full")
                password_input = ui.input("Пароль", password=True, password_toggle_button=True).props("dense outlined").classes("w-full")

                def login() -> None:
                    user = auth_db.login(username=str(username_input.value or ""), password=str(password_input.value or ""))
                    if not user:
                        ui.notify("Неверный логин или пароль.", type="negative")
                        return
                    state.current_user = user
                    state.auth_token = auth_db.create_session(username=str(user.get("username") or ""))
                    try:
                        app.storage.user["auth_token"] = state.auth_token
                    except Exception:
                        pass
                    ui.notify("Вход выполнен.", type="positive")
                    render()

                password_input.on("keyup.enter", lambda _: login())
                ui.button("Войти", icon="login", on_click=login).props("unelevated")
            return

        user = state.current_user
        is_admin = str(user.get("role") or "user") == "admin"

        # ── Реестр секций: (key, icon, label, keywords) ─────────────────
        user_sections: List[tuple] = [
            ("profile",   "person",          "Профиль",              ["telegram", "имя", "аккаунт", "профиль"]),
            ("explorer",  "folder_open",      "Проводник",            ["файлы", "вид", "сортировка"]),
            ("favorites", "star_border",      "Избранное",            ["закладки"]),
            ("password",  "key",              "Пароль и выход",       ["смена", "выход", "logout"]),
        ]
        admin_sections: List[tuple] = [
            ("paths",         "storage",        "Пути и Qdrant",          ["каталог", "база", "url", "коллекция"]),
            ("llm",           "smart_toy",      "Нейросеть",              ["ollama", "модель", "ai", "llm", "rag"]),
            ("aliases",       "travel_explore", "Синонимы поиска",        ["группы", "расширение", "запросы"]),
            ("indexing",      "build",          "Индексация",             ["индекс", "статус", "прогресс"]),
            ("security",      "security",       "Сессии и безопасность",  ["сессии", "системные файлы"]),
            ("users",         "group",          "Пользователи",           ["роль", "статус", "логин"]),
            ("registrations", "person_add",     "Регистрации",            ["заявки", "одобрить"]),
            ("tg_chats",      "chat",           "Telegram чаты",          ["бот", "chat id"]),
        ]

        active = [state.settings_section]  # сохраняем между ре-рендерами
        q_ref  = [""]

        # ── IDE-лейаут ───────────────────────────────────────────────────
        with ui.row().classes("w-full gap-0 items-start"):

            # Левая боковая панель
            with ui.column().classes("flex-none gap-1").style(
                "width:220px; min-width:220px; border-right:1px solid #e5e7eb; padding-right:12px; margin-right:16px"
            ):
                ui.label("Настройки").classes("text-xl font-semibold mb-2")
                search_box = ui.input(
                    placeholder="Поиск настроек…",
                    on_change=lambda e: (q_ref.__setitem__(0, str(e.value or "").lower()), render_nav()),
                ).props("dense outlined clearable").classes("w-full")

                nav_col = ui.column().classes("w-full gap-0")

            # Правая область контента
            content_col = ui.column().classes("flex-1 gap-3 min-w-0")

        # ── Навигация ────────────────────────────────────────────────────
        def _visible(entry: tuple) -> bool:
            q = q_ref[0]
            if not q:
                return True
            key, icon, label, kws = entry
            return q in label.lower() or any(q in kw.lower() for kw in kws)

        def render_nav() -> None:
            nav_col.clear()
            with nav_col:
                groups: List[tuple] = [("", user_sections)]
                if is_admin:
                    groups.append(("Администратор", admin_sections))
                for group_label, sections in groups:
                    filtered = [s for s in sections if _visible(s)]
                    if not filtered:
                        continue
                    if group_label:
                        ui.label(group_label.upper()).classes(
                            "text-xs text-gray-400 font-semibold mt-3 mb-1 px-2"
                        )
                    for key, icon, label, _ in filtered:
                        is_active = active[0] == key
                        bg = "background:#eef2ff;" if is_active else ""
                        with ui.row().classes("w-full items-center gap-2 px-2 py-1 rounded cursor-pointer").style(
                            bg + "user-select:none"
                        ).on("click", lambda k=key: navigate(k)):
                            ui.icon(icon, size="16px").classes(
                                "text-indigo-600" if is_active else "text-gray-400"
                            )
                            ui.label(label).classes(
                                "text-sm font-medium text-indigo-700" if is_active else "text-sm text-gray-700"
                            )

        # ── Контент секции ───────────────────────────────────────────────
        def render_section() -> None:
            content_col.clear()
            with content_col:
                sec = active[0]

                if sec == "profile":
                    with ui.column().classes("rag-card w-full p-4 gap-3"):
                        ui.label("Профиль").classes("text-xl font-semibold")
                        ui.label(
                            f"Логин: {user.get('username')} · роль: {user.get('role')} · статус: {user.get('status')}"
                        ).classes("rag-meta")
                        disp_in = ui.input("Имя", value=str(user.get("display_name") or "")).props("dense outlined").classes("w-full")
                        linked_tg_id = str(user.get("telegram_chat_id") or "").strip()
                        linked_tg_un = str(user.get("telegram_username") or "").strip()

                        def save_profile() -> None:
                            auth_db.update_profile(
                                username=str(user.get("username") or ""),
                                display_name=str(disp_in.value or ""),
                                telegram_chat_id=linked_tg_id,
                                telegram_username=linked_tg_un,
                            )
                            _refresh_current_user(state)
                            ui.notify("Профиль сохранён.", type="positive")

                        def bind_tg() -> None:
                            bot_link = str(state.cfg.get("telegram_bot_link") or "").strip()
                            if not bot_link:
                                ui.notify("В config.json не задан telegram_bot_link.", type="warning")
                                return
                            out = auth_db.create_telegram_link_token(username=str(user.get("username") or ""))
                            if not out.get("ok"):
                                ui.notify(f"Ошибка: {out.get('reason')}", type="negative")
                                return
                            link = _telegram_deeplink(bot_link, "link", str(out.get("token") or ""))
                            if not link:
                                ui.notify("Не удалось создать ссылку привязки.", type="negative")
                                return
                            ui.run_javascript(
                                "(() => {"
                                f"const url = {json.dumps(link)};"
                                "const w = window.open(url, '_blank', 'noopener,noreferrer');"
                                "if (!w) { window.location.href = url; }"
                                "})();"
                            )
                            ui.notify("Откройте Telegram и подтвердите привязку.", type="positive")

                        ui.button("Сохранить профиль", icon="save", on_click=save_profile).props("outline")
                        if linked_tg_id:
                            linked_label = f"@{linked_tg_un}" if linked_tg_un else linked_tg_id
                            ui.label(f"Telegram уже привязан: {linked_label}").classes("rag-meta")
                        else:
                            ui.label("Telegram не привязан. Можно привязать в один клик.").classes("rag-meta")
                            ui.button("Привязать Telegram", icon="link", on_click=bind_tg).props("outline")

                elif sec == "explorer":
                    with ui.column().classes("rag-card w-full p-4 gap-3"):
                        ui.label("Проводник").classes("text-xl font-semibold")
                        ui.label(
                            f"Вид: {state.explorer_view} · сортировка: {state.explorer_sort} · "
                            f"{'убывание' if state.explorer_desc else 'возрастание'} · тип: {state.explorer_ext}"
                        ).classes("rag-meta")

                        def reset_explorer() -> None:
                            auth_db.reset_user_settings(username=str(user.get("username") or ""))
                            state.explorer_view = "Таблица"
                            state.explorer_sort = "По имени"
                            state.explorer_desc = False
                            state.explorer_ext = "Все"
                            _log_app_event(state, "settings", "reset_explorer")
                            ui.notify("Настройки проводника сброшены.", type="positive")
                            render_section()

                        ui.button("Сбросить настройки проводника", icon="restart_alt", on_click=reset_explorer).props("outline")

                elif sec == "favorites":
                    with ui.column().classes("rag-card w-full p-4 gap-3"):
                        ui.label("Избранное").classes("text-xl font-semibold")
                        if not state.favorites:
                            ui.label("Закладок пока нет. Добавьте файл или папку звёздочкой в проводнике.").classes("rag-meta")
                        for fav in state.favorites:
                            fav_path = Path(str(fav.get("path") or ""))
                            item_type = str(fav.get("item_type") or "")
                            with ui.row().classes("w-full items-center gap-2"):
                                ui.icon("folder" if item_type == "folder" else "description")
                                ui.label(str(fav.get("title") or fav_path.name or fav_path)).classes("font-medium")
                                ui.label(str(fav_path)).classes("rag-path flex-1")
                                if item_type == "folder":
                                    ui.button("Открыть", on_click=lambda p=fav_path: go_explorer(str(p))).props("outline dense")
                                else:
                                    ui.button("Открыть", on_click=lambda p=fav_path: open_file_viewer(p)).props("outline dense")
                                ui.button(icon="delete", on_click=lambda p=fav_path: (
                                    _toggle_favorite(state, p), render_section()
                                )).props("flat round dense")

                elif sec == "password":
                    with ui.column().classes("rag-card w-full p-4 gap-3"):
                        ui.label("Смена пароля").classes("text-xl font-semibold")
                        old_pw = ui.input("Текущий пароль", password=True, password_toggle_button=True).props("dense outlined").classes("w-full")
                        new_pw = ui.input("Новый пароль", password=True, password_toggle_button=True).props("dense outlined").classes("w-full")
                        new_pw2 = ui.input("Повторите пароль", password=True, password_toggle_button=True).props("dense outlined").classes("w-full")

                        def change_pw() -> None:
                            if str(new_pw.value or "") != str(new_pw2.value or ""):
                                ui.notify("Пароли не совпадают.", type="warning")
                                return
                            ok = auth_db.change_password(
                                username=str(user.get("username") or ""),
                                old_password=str(old_pw.value or ""),
                                new_password=str(new_pw.value or ""),
                            )
                            if ok:
                                _refresh_current_user(state)
                            ui.notify("Пароль изменён." if ok else "Не удалось изменить пароль.",
                                      type="positive" if ok else "negative")

                        with ui.row().classes("gap-2"):
                            ui.button("Сменить пароль", icon="key", on_click=change_pw).props("outline")
                            ui.button("Выйти", icon="logout", on_click=do_logout).props("flat")

                elif sec == "paths":
                    render_admin_path_settings()
                elif sec == "llm":
                    render_admin_llm_settings()
                elif sec == "aliases":
                    render_admin_search_aliases()
                elif sec == "indexing":
                    render_index_dashboard()
                elif sec == "security":
                    render_admin_security_settings(auth_db)
                elif sec == "users":
                    render_admin_users(auth_db)
                elif sec == "registrations":
                    render_admin_registration_requests(auth_db)
                elif sec == "tg_chats":
                    render_admin_telegram_chats(auth_db)

        def navigate(key: str) -> None:
            active[0] = key
            state.settings_section = key
            render_nav()
            render_section()

        render_nav()
        render_section()

    def render_telegram_screen() -> None:
        enabled = bool(state.cfg.get("telegram_enabled"))
        token_set = bool(str(state.cfg.get("telegram_bot_token") or "").strip())
        with ui.column().classes("rag-card w-full p-4 gap-2"):
            ui.label(f"Статус: {'включен' if enabled else 'выключен'}").classes("text-lg font-semibold")
            ui.label(f"Токен: {'задан' if token_set else 'не задан'}").classes("rag-meta")
            bot_link = str(state.cfg.get("telegram_bot_link") or "").strip()
            if bot_link:
                ui.link("Открыть бота", bot_link, new_tab=True)

    def render_stats_screen() -> None:
        if str((state.current_user or {}).get("role") or "") != "admin":
            ui.label("Раздел доступен только администратору.").classes("rag-card p-4 text-red-700")
            return

        telemetry_path = _telemetry_db_path(state.cfg)
        auth_db = _get_auth_db(state)

        # ── KPI tiles (7-day sparkbars) ────────────────────────────────────
        today_searches_row = _db_query_dicts(
            telemetry_path,
            "SELECT COUNT(*) AS cnt FROM search_logs WHERE substr(ts,1,10)=date('now')",
        )
        today_searches = int((today_searches_row[0].get("cnt") or 0) if today_searches_row else 0)

        week_searches_row = _db_query_dicts(
            telemetry_path,
            "SELECT COUNT(*) AS cnt FROM search_logs WHERE ts >= datetime('now','-7 days')",
        )
        week_searches = int((week_searches_row[0].get("cnt") or 0) if week_searches_row else 0)

        week_users_row = _db_query_dicts(
            telemetry_path,
            """SELECT COUNT(DISTINCT COALESCE(NULLIF(username,''),source)) AS cnt
               FROM search_logs WHERE ts >= datetime('now','-7 days')""",
        )
        week_users = int((week_users_row[0].get("cnt") or 0) if week_users_row else 0)

        error_rate_row = _db_query_dicts(
            telemetry_path,
            """SELECT ROUND(100.0 * SUM(CASE WHEN ok=0 THEN 1 ELSE 0 END) / MAX(COUNT(*),1), 1) AS rate
               FROM search_logs WHERE ts >= datetime('now','-7 days')""",
        )
        error_rate = float((error_rate_row[0].get("rate") or 0.0) if error_rate_row else 0.0)

        # 7-day sparkline data
        spark_rows = _db_query_dicts(
            telemetry_path,
            """SELECT substr(ts,1,10) AS day, COUNT(*) AS cnt
               FROM search_logs WHERE ts >= datetime('now','-7 days')
               GROUP BY substr(ts,1,10) ORDER BY day""",
        )
        spark_vals = [int(r.get("cnt") or 0) for r in spark_rows]
        spark_max = max(spark_vals + [1])

        def _spark_bar(val: int, mx: int) -> str:
            h = max(4, int(val / mx * 28))
            return f'<span style="flex:1;background:var(--rag-accent);opacity:.65;border-radius:2px 2px 0 0;height:{h}px;align-self:flex-end;"></span>'

        spark_html = '<div class="rag-kpi-sparkline">' + "".join(_spark_bar(v, spark_max) for v in spark_vals) + '</div>'

        kpi_data = [
            ("ПОИСКОВ СЕГОДНЯ", str(today_searches), spark_html, "var(--rag-accent)"),
            ("ЗАПРОСОВ 7 ДНЕЙ", str(week_searches), spark_html, "var(--rag-ok)"),
            ("АКТИВНЫХ ПОЛЬЗ.", str(week_users), spark_html, "#f59e0b"),
            ("ОШИБОК %", f"{error_rate:.1f}%", spark_html, "var(--rag-danger)" if error_rate > 5 else "var(--rag-ok)"),
        ]

        with ui.row().classes("w-full gap-3 flex-wrap"):
            for kpi_label, kpi_value, kpi_spark, val_color in kpi_data:
                with ui.element("div").classes("rag-kpi-tile"):
                    ui.html(f'<div class="rag-kpi-label">{kpi_label}</div>', sanitize=False)
                    ui.html(f'<div class="rag-kpi-value" style="color:{val_color};">{kpi_value}</div>', sanitize=False)
                    ui.html(kpi_spark, sanitize=False)

        # ── AI Insight (from telemetry summary) ───────────────────────────
        if week_searches > 0:
            peak_hour_rows = _db_query_dicts(
                telemetry_path,
                """SELECT strftime('%H', ts) AS hr, COUNT(*) AS cnt
                   FROM search_logs WHERE ts >= datetime('now','-7 days')
                   GROUP BY hr ORDER BY cnt DESC LIMIT 1""",
            )
            peak_hour = str((peak_hour_rows[0].get("hr") or "?") if peak_hour_rows else "?") + ":00"
            top_q_rows = _db_query_dicts(
                telemetry_path,
                """SELECT query, COUNT(*) AS cnt FROM search_logs WHERE query<>''
                   GROUP BY lower(query) ORDER BY cnt DESC LIMIT 1""",
            )
            top_q = str((top_q_rows[0].get("query") or "—") if top_q_rows else "—")[:40]
            insight_text = (
                f"За 7 дней выполнено <b>{week_searches}</b> поисков от <b>{week_users}</b> пользователей. "
                f"Пик активности — <b>{peak_hour}</b>. "
                f"Самый частый запрос: «<b>{top_q}</b>»."
            )
            ui.html(
                f'<div class="rag-ai-insight">' +
                f'<div style="font-size:11px;font-weight:700;letter-spacing:.06em;color:#92400e;margin-bottom:6px;">✦ AI ИНСАЙТ</div>' +
                f'<div style="font-size:13px;">{insight_text}</div>' +
                f'</div>',
                sanitize=False,
            )

        # ── Sub-tab bar ────────────────────────────────────────────────────
        tabs_def = [
            ("overview", "Обзор"),
            ("queries", "Запросы"),
            ("users", "Пользователи"),
            ("auth", "Аудит входов"),
            ("errors", "Ошибки"),
        ]

        tab_bar = ui.element("div").classes("w-full").style("display:flex;gap:6px;flex-wrap:wrap;margin-bottom:12px;")
        tab_content = ui.column().classes("w-full gap-3")

        def _switch_tab(key: str) -> None:
            state.analytics_tab = key
            tab_bar.clear()
            with tab_bar:
                for tk, tlabel in tabs_def:
                    active = "active" if tk == state.analytics_tab else ""
                    ui.html(
                        f'<span class="rag-analytics-tab {active}" ' +
                        f'style="cursor:pointer;" onclick="void(0);">{tlabel}</span>',
                        sanitize=False,
                    ).on("click", lambda t=tk: _switch_tab(t))
            _render_tab_content()

        def _render_tab_content() -> None:
            tab_content.clear()
            with tab_content:
                tab = state.analytics_tab

                if tab == "overview":
                    # Chart: searches by day
                    searches_by_day = _db_query_dicts(
                        telemetry_path,
                        """SELECT substr(ts,1,10) AS day, COUNT(*) AS count
                           FROM search_logs GROUP BY substr(ts,1,10) ORDER BY day LIMIT 30""",
                    )
                    with ui.column().classes("rag-card w-full p-4 gap-3"):
                        ui.label("Поиски по дням").classes("text-xl font-semibold")
                        if searches_by_day:
                            ui.echart({
                                "tooltip": {"trigger": "axis"},
                                "xAxis": {"type": "category", "data": [r["day"] for r in searches_by_day]},
                                "yAxis": {"type": "value"},
                                "series": [{"type": "bar", "data": [r["count"] for r in searches_by_day],
                                            "name": "Поиски", "itemStyle": {"color": "#3d63ff"}}],
                            }).classes("w-full h-64")
                        else:
                            ui.html('<div class="rag-meta" style="padding:24px;text-align:center;">Данных пока нет</div>', sanitize=False)

                    # Top rows
                    top_queries = _db_query_dicts(
                        telemetry_path,
                        """SELECT query, COUNT(*) AS count FROM search_logs WHERE query<>''
                           GROUP BY lower(query) ORDER BY count DESC LIMIT 10""",
                    )
                    top_users = _db_query_dicts(
                        telemetry_path,
                        """SELECT COALESCE(NULLIF(username,''),source,'unknown') AS username, COUNT(*) AS count
                           FROM search_logs GROUP BY username ORDER BY count DESC LIMIT 10""",
                    )
                    with ui.row().classes("w-full gap-3 items-start flex-wrap"):
                        with ui.column().classes("rag-card flex-1 p-4 gap-2" ):
                            ui.label("Топ запросов").classes("font-semibold")
                            for row in top_queries:
                                with ui.row().classes("w-full items-center gap-2"):
                                    ui.html(f'<span class="flex-1 truncate text-sm">{str(row["query"])[:50]}</span>', sanitize=False)
                                    ui.html(f'<span class="rag-chip">{row["count"]}</span>', sanitize=False)
                        with ui.column().classes("rag-card flex-1 p-4 gap-2"):
                            ui.label("Пользователи").classes("font-semibold")
                            for row in top_users:
                                with ui.row().classes("w-full items-center gap-2"):
                                    ui.html(f'<span class="flex-1 truncate text-sm">{str(row["username"])[:30]}</span>', sanitize=False)
                                    ui.html(f'<span class="rag-chip">{row["count"]}</span>', sanitize=False)

                elif tab == "queries":
                    with ui.column().classes("rag-card w-full p-4 gap-3"):
                        ui.label("История запросов").classes("text-xl font-semibold")
                        with ui.row().classes("w-full gap-2 flex-wrap"):
                            src_filter = ui.select(["Все", "Telegram", "Web/прочее"], value="Все", label="Источник").props("dense outlined").classes("w-40")
                            user_filter = ui.input("Пользователь").props("dense outlined clearable").classes("w-40")
                            q_filter = ui.input("Запрос").props("dense outlined clearable").classes("flex-1")
                            ok_filter = ui.select(["Все", "OK", "Ошибки"], value="Все", label="OK").props("dense outlined").classes("w-28")

                        q_table = ui.table(
                            rows=[],
                            columns=[
                                {"name": "ts", "label": "Время", "field": "ts", "sortable": True},
                                {"name": "source", "label": "Источник", "field": "source"},
                                {"name": "username", "label": "Польз.", "field": "username"},
                                {"name": "query", "label": "Запрос", "field": "query"},
                                {"name": "results_count", "label": "Рез.", "field": "results_count", "sortable": True},
                                {"name": "duration_ms", "label": "мс", "field": "duration_ms", "sortable": True},
                                {"name": "error", "label": "Ошибка", "field": "error"},
                            ],
                            pagination=15,
                        ).classes("w-full")

                        def _refresh_q() -> None:
                            rows = _db_query_dicts(telemetry_path,
                                "SELECT ts,source,username,query,results_count,duration_ms,ok,error FROM search_logs ORDER BY id DESC LIMIT 500")
                            sm = str(src_filter.value or "Все")
                            if sm == "Telegram":
                                rows = [r for r in rows if str(r.get("source") or "").startswith("telegram_bot:")]
                            elif sm == "Web/прочее":
                                rows = [r for r in rows if not str(r.get("source") or "").startswith("telegram_bot:")]
                            un = str(user_filter.value or "").strip().lower()
                            if un: rows = [r for r in rows if un in str(r.get("username") or "").lower()]
                            qn = str(q_filter.value or "").strip().lower()
                            if qn: rows = [r for r in rows if qn in str(r.get("query") or "").lower()]
                            om = str(ok_filter.value or "Все")
                            if om == "OK": rows = [r for r in rows if int(r.get("ok") or 0) == 1]
                            elif om == "Ошибки": rows = [r for r in rows if int(r.get("ok") or 0) == 0]
                            q_table.rows = rows; q_table.update()

                        src_filter.on_value_change(lambda e: _refresh_q())
                        user_filter.on_value_change(lambda e: _refresh_q())
                        q_filter.on_value_change(lambda e: _refresh_q())
                        ok_filter.on_value_change(lambda e: _refresh_q())
                        _refresh_q()

                elif tab == "users":
                    users_by_day = _db_query_dicts(
                        telemetry_path,
                        """SELECT substr(ts,1,10) AS day,
                              COUNT(DISTINCT COALESCE(NULLIF(username,''),source)) AS users
                           FROM search_logs GROUP BY substr(ts,1,10) ORDER BY day LIMIT 30""",
                    )
                    with ui.column().classes("rag-card w-full p-4 gap-3"):
                        ui.label("Активные пользователи по дням").classes("text-xl font-semibold")
                        if users_by_day:
                            ui.echart({
                                "tooltip": {"trigger": "axis"},
                                "xAxis": {"type": "category", "data": [r["day"] for r in users_by_day]},
                                "yAxis": {"type": "value"},
                                "series": [{"type": "line", "data": [r["users"] for r in users_by_day],
                                            "name": "Пользователи", "smooth": True,
                                            "itemStyle": {"color": "#f59e0b"}}],
                            }).classes("w-full h-64")
                        else:
                            ui.html('<div class="rag-meta" style="padding:24px;text-align:center;">Данных пока нет</div>', sanitize=False)

                    top_users_full = _db_query_dicts(
                        telemetry_path,
                        """SELECT COALESCE(NULLIF(username,''),source,'unknown') AS username,
                              COUNT(*) AS searches, MAX(substr(ts,1,16)) AS last_seen
                           FROM search_logs GROUP BY username ORDER BY searches DESC LIMIT 30""",
                    )
                    with ui.column().classes("rag-card w-full p-4 gap-2"):
                        ui.label("Все пользователи").classes("font-semibold")
                        ui.table(
                            rows=top_users_full,
                            columns=[
                                {"name": "username", "label": "Пользователь", "field": "username"},
                                {"name": "searches", "label": "Запросов", "field": "searches", "sortable": True},
                                {"name": "last_seen", "label": "Последний визит", "field": "last_seen"},
                            ],
                            pagination=15,
                        ).classes("w-full")

                elif tab == "auth":
                    auth_events = auth_db.list_auth_events(limit=200)
                    with ui.column().classes("rag-card w-full p-4 gap-3"):
                        ui.label("История входов").classes("text-xl font-semibold")
                        with ui.row().classes("w-full gap-2 flex-wrap"):
                            a_src = ui.select(["Все", "Telegram", "Web/прочее"], value="Все", label="Источник").props("dense outlined").classes("w-40")
                            a_user = ui.input("Пользователь").props("dense outlined clearable").classes("w-40")
                            a_evt = ui.input("Событие").props("dense outlined clearable").classes("flex-1")
                            a_ok = ui.select(["Все", "OK", "Ошибки"], value="Все", label="OK").props("dense outlined").classes("w-28")

                        a_table = ui.table(
                            rows=[],
                            columns=[
                                {"name": "ts", "label": "Время", "field": "ts", "sortable": True},
                                {"name": "username", "label": "Пользователь", "field": "username"},
                                {"name": "event_type", "label": "Событие", "field": "event_type"},
                                {"name": "ok", "label": "OK", "field": "ok"},
                                {"name": "error", "label": "Ошибка", "field": "error"},
                            ],
                            pagination=15,
                        ).classes("w-full")

                        def _refresh_auth() -> None:
                            rows = list(auth_events)
                            sm = str(a_src.value or "Все")
                            if sm == "Telegram": rows = [r for r in rows if str(r.get("event_type") or "").startswith("telegram_")]
                            elif sm == "Web/прочее": rows = [r for r in rows if not str(r.get("event_type") or "").startswith("telegram_")]
                            un = str(a_user.value or "").strip().lower()
                            if un: rows = [r for r in rows if un in str(r.get("username") or "").lower()]
                            en = str(a_evt.value or "").strip().lower()
                            if en: rows = [r for r in rows if en in str(r.get("event_type") or "").lower()]
                            om = str(a_ok.value or "Все")
                            if om == "OK": rows = [r for r in rows if int(r.get("ok") or 0) == 1]
                            elif om == "Ошибки": rows = [r for r in rows if int(r.get("ok") or 0) == 0]
                            a_table.rows = rows; a_table.update()

                        a_src.on_value_change(lambda e: _refresh_auth())
                        a_user.on_value_change(lambda e: _refresh_auth())
                        a_evt.on_value_change(lambda e: _refresh_auth())
                        a_ok.on_value_change(lambda e: _refresh_auth())
                        _refresh_auth()

                elif tab == "errors":
                    error_rows = _db_query_dicts(
                        telemetry_path,
                        """SELECT ts, source, username, query, error
                           FROM search_logs WHERE ok=0 AND error IS NOT NULL AND error<>''
                           ORDER BY id DESC LIMIT 100""",
                    )
                    app_errors = _db_query_dicts(
                        telemetry_path,
                        """SELECT ts, feature, action, data FROM app_events
                           WHERE action LIKE '%error%' OR action LIKE '%fail%'
                           ORDER BY id DESC LIMIT 100""",
                    ) if True else []

                    with ui.column().classes("rag-card w-full p-4 gap-2"):
                        ui.label(f"Ошибки поиска ({len(error_rows)})").classes("text-xl font-semibold")
                        if error_rows:
                            ui.table(
                                rows=error_rows,
                                columns=[
                                    {"name": "ts", "label": "Время", "field": "ts"},
                                    {"name": "username", "label": "Польз.", "field": "username"},
                                    {"name": "query", "label": "Запрос", "field": "query"},
                                    {"name": "error", "label": "Ошибка", "field": "error"},
                                ],
                                pagination=15,
                            ).classes("w-full")
                        else:
                            ui.html('<div class="rag-meta" style="padding:16px;text-align:center;">Ошибок нет — всё хорошо!</div>', sanitize=False)

        # Render initial tab bar
        with tab_bar:
            for tk, tlabel in tabs_def:
                active = "active" if tk == state.analytics_tab else ""
                ui.html(
                    f'<span class="rag-analytics-tab {active}" ' +
                    f'style="cursor:pointer;">{tlabel}</span>',
                    sanitize=False,
                ).on("click", lambda t=tk: _switch_tab(t))

        _render_tab_content()

    def render() -> None:
        page_root.classes(remove="search")
        if state.screen == "search":
            page_root.classes(add="search")
        header_title.set_text({
            "home": "Главная",
            "search": "Поиск",
            "explorer": "Проводник",
            "index": "Индекс",
            "telegram": "Telegram",
            "settings": "Настройки",
            "stats": "Аналитика",
        }.get(state.screen, "Поиск"))
        if state.header_breadcrumbs is not None:
            state.header_breadcrumbs.clear()
        if state.header_explorer_actions is not None:
            state.header_explorer_actions.clear()
        update_nav()
        content.clear()
        with content:
            if state.current_user is None:
                try:
                    drawer.set_value(False)
                except Exception:
                    pass
                try:
                    drawer.set_visibility(False)
                except Exception:
                    pass
                try:
                    menu_button.set_visibility(False)
                except Exception:
                    pass
                try:
                    theme_button.set_visibility(False)
                except Exception:
                    pass
                render_login_screen()
                return
            try:
                drawer.set_visibility(True)
            except Exception:
                pass
            try:
                menu_button.set_visibility(True)
            except Exception:
                pass
            try:
                theme_button.set_visibility(True)
                theme_button.set_icon("light_mode" if state.theme == "dark" else "dark_mode")
            except Exception:
                pass
            dark_mode.set_value(state.theme == "dark")
            touch_activity()
            if state.screen == "home":
                render_home_screen()
            elif state.screen == "explorer":
                try:
                    drawer.set_visibility(True)
                except Exception:
                    pass
                render_explorer_screen()
            elif state.screen == "index":
                render_index_screen()
            elif state.screen == "telegram":
                render_telegram_screen()
            elif state.screen == "settings":
                render_settings_screen()
            elif state.screen == "stats":
                render_stats_screen()
            else:
                render_search_screen()

    render()


@ui.page("/")
def root_page() -> None:
    ui.navigate.to("/home")


@ui.page("/home")
def home_page() -> None:
    _build_page("home")


@ui.page("/search")
def search_page() -> None:
    _build_page("search")


@ui.page("/explorer")
def explorer_page() -> None:
    _build_page("explorer")


@ui.page("/index")
def index_page() -> None:
    _build_page("index")


@ui.page("/telegram")
def telegram_page() -> None:
    _build_page("telegram")


@ui.page("/settings")
def settings_page() -> None:
    _build_page("settings")


@ui.page("/stats")
def stats_page() -> None:
    _build_page("stats")


def main(argv: Optional[List[str]] = None) -> None:
    parser = argparse.ArgumentParser(description="Запустить NiceGUI-интерфейс RAG Каталога.")
    parser.add_argument("--host", default="127.0.0.1")
    parser.add_argument("--port", type=int, default=8080)
    parser.add_argument("--no-show", action="store_true", help="Не открывать браузер автоматически.")
    args = parser.parse_args(argv)
    try:
        _recover_background_tasks(load_config())
    except Exception as exc:
        print(f"[nice_app] background recovery skipped: {exc}", file=sys.stderr)
    _favicon = (
        str(_MARK_SVG_PATH) if _MARK_SVG_PATH.exists()
        else (APP_ICON_PATH if APP_ICON_PATH.exists() else None)
    )
    ui.run(
        title="ТЕХНОПОИСК",
        host=args.host,
        port=args.port,
        favicon=_favicon,
        language="ru",
        reload=False,
        show=not args.no_show,
        dark=False,
        storage_secret="rag-catalog-local-secret",
    )


if __name__ in {"__main__", "__mp_main__"}:
    main()
