"""Document and image text extraction helpers.

These functions deliberately return an empty string on recoverable extraction
errors. The indexer owns telemetry/stage accounting; extractors only read files
and log the concrete failure.
"""

from __future__ import annotations

import csv
import logging
import os
import re
import shutil
import subprocess
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Any, Callable
from xml.etree import ElementTree
from zipfile import ZipFile

from docx import Document
from openpyxl import load_workbook

from rag_catalog.core.ocr_runtime import apply_tesseract_runtime

from .contract import ExtractedDocument, TextBlock, document_from_legacy_text

logger = logging.getLogger(__name__)

_XLSX_MAIN_NS = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}"
_PROJECT_ROOT = Path(__file__).resolve().parents[4]
_TOOLS_ROOT = _PROJECT_ROOT / "tools"


def _load_xlsx_workbook(filepath: Path, *, read_only: bool = True, data_only: bool = True) -> Any:
    """Load XLSX, tolerating archives with incorrectly cased sharedStrings path."""
    try:
        return load_workbook(filepath, read_only=read_only, data_only=data_only)
    except KeyError as exc:
        message = str(exc)
        if "xl/sharedStrings.xml" not in message:
            raise

    buffer = BytesIO()
    with ZipFile(filepath, "r") as src, ZipFile(buffer, "w") as dst:
        names = set(src.namelist())
        has_expected = "xl/sharedStrings.xml" in names
        for info in src.infolist():
            name = info.filename
            next_name = name
            if not has_expected and name.lower() == "xl/sharedstrings.xml":
                next_name = "xl/sharedStrings.xml"
            dst.writestr(next_name, src.read(name))
    buffer.seek(0)
    return load_workbook(buffer, read_only=read_only, data_only=data_only)


def _xlsx_cell_text(cell: ElementTree.Element, shared_strings: list[str]) -> str:
    cell_type = cell.attrib.get("t", "")
    if cell_type == "inlineStr":
        return "".join(node.text or "" for node in cell.iter(f"{_XLSX_MAIN_NS}t")).strip()

    value = cell.find(f"{_XLSX_MAIN_NS}v")
    if value is None or value.text is None:
        return ""

    text = value.text.strip()
    if cell_type == "s":
        try:
            return shared_strings[int(text)]
        except (IndexError, ValueError):
            return ""
    if cell_type == "b":
        return "TRUE" if text == "1" else "FALSE"
    return text


def _read_xlsx_shared_strings(zf: ZipFile) -> list[str]:
    try:
        with zf.open("xl/sharedStrings.xml") as fh:
            root = ElementTree.parse(fh).getroot()
    except KeyError:
        return []

    strings: list[str] = []
    for item in root.findall(f"{_XLSX_MAIN_NS}si"):
        strings.append("".join(node.text or "" for node in item.iter(f"{_XLSX_MAIN_NS}t")).strip())
    return strings


def _extract_xlsx_zip_fallback(filepath: Path, *, max_chars: int = 0) -> str:
    """Best-effort XLSX parser for damaged archives openpyxl refuses to load."""
    parts: list[str] = []
    total_chars = 0
    done = False
    with ZipFile(filepath, "r") as zf:
        shared_strings = _read_xlsx_shared_strings(zf)
        worksheet_names = sorted(name for name in zf.namelist() if name.startswith("xl/worksheets/sheet") and name.endswith(".xml"))
        for idx, name in enumerate(worksheet_names, start=1):
            if done:
                break
            parts.append(f"Лист: sheet{idx}")
            with zf.open(name) as fh:
                root = ElementTree.parse(fh).getroot()
            for row in root.iter(f"{_XLSX_MAIN_NS}row"):
                values = [_xlsx_cell_text(cell, shared_strings) for cell in row.findall(f"{_XLSX_MAIN_NS}c")]
                row_text = " | ".join(values)
                if row_text.strip():
                    parts.append(row_text)
                    total_chars += len(row_text)
                    if max_chars and total_chars >= max_chars:
                        done = True
                        break
    return "\n".join(parts)


def extract_docx(filepath: Path) -> str:
    """Extract text from DOCX paragraphs and table cells."""
    try:
        doc = Document(filepath)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    except Exception as exc:
        logger.warning("Ошибка чтения DOCX %s: %s", filepath, exc)
        return ""


def extract_rtf(filepath: Path, *, max_chars: int = 0) -> str:
    """Extract text from RTF without requiring optional dependencies."""
    try:
        raw = _read_text_file(filepath, max_chars=0)
        try:
            from striprtf.striprtf import rtf_to_text  # type: ignore  # noqa: PLC0415

            text = rtf_to_text(raw)
            return text[:max_chars] if max_chars else text
        except Exception:
            pass

        text = re.sub(r"\\'[0-9a-fA-F]{2}", " ", raw)
        text = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", text)
        text = text.replace(r"\~", " ").replace(r"\_", "-")
        text = text.replace("{", " ").replace("}", " ").replace("\\", " ")
        text = re.sub(r"\s+", " ", text).strip()
        return text[:max_chars] if max_chars else text
    except Exception as exc:
        logger.warning("Ошибка чтения RTF %s: %s", filepath, exc)
        return ""


def extract_pptx(filepath: Path, *, max_chars: int = 0) -> str:
    """Extract slide text from PPTX by reading slide XML."""
    try:
        parts: list[str] = []
        total = 0
        done = False
        with ZipFile(filepath, "r") as zf:
            slide_names = sorted(
                name for name in zf.namelist()
                if name.startswith("ppt/slides/slide") and name.endswith(".xml")
            )
            for idx, name in enumerate(slide_names, start=1):
                if done:
                    break
                root = ElementTree.parse(zf.open(name)).getroot()
                texts = [node.text or "" for node in root.iter() if node.tag.endswith("}t") and (node.text or "").strip()]
                slide_text = " ".join(text.strip() for text in texts if text.strip())
                if slide_text:
                    row = f"Слайд: {idx}\n{slide_text}"
                    parts.append(row)
                    total += len(row)
                    if max_chars and total >= max_chars:
                        done = True
        return "\n".join(parts)
    except Exception as exc:
        logger.warning("Ошибка чтения PPTX %s: %s", filepath, exc)
        return ""


def extract_pptx_document(filepath: Path, *, max_chars: int = 0) -> ExtractedDocument:
    """Extract slide text from PPTX as structured slide blocks."""
    try:
        blocks: list[TextBlock] = []
        total = 0
        done = False
        with ZipFile(filepath, "r") as zf:
            slide_names = sorted(
                name for name in zf.namelist()
                if name.startswith("ppt/slides/slide") and name.endswith(".xml")
            )
            for idx, name in enumerate(slide_names, start=1):
                if done:
                    break
                root = ElementTree.parse(zf.open(name)).getroot()
                texts = [node.text or "" for node in root.iter() if node.tag.endswith("}t") and (node.text or "").strip()]
                slide_text = " ".join(text.strip() for text in texts if text.strip())
                if slide_text:
                    blocks.append(TextBlock(text=slide_text, slide=idx))
                    total += len(slide_text)
                    if max_chars and total >= max_chars:
                        done = True
        return ExtractedDocument(blocks=tuple(blocks))
    except Exception as exc:
        logger.warning("Ошибка чтения PPTX %s: %s", filepath, exc)
        return ExtractedDocument(blocks=())


def _first_existing_tool(*paths: Path) -> str:
    for path in paths:
        if path.exists():
            return str(path)
    return ""


def _resolve_env_tool(*env_names: str) -> str:
    for name in env_names:
        value = os.environ.get(name, "").strip()
        if value and Path(value).exists():
            return value
    return ""


def _resolve_antiword() -> str:
    return (
        _resolve_env_tool("RAG_ANTIWORD_CMD", "ANTIWORD")
        or _first_existing_tool(
            _TOOLS_ROOT / "antiword" / "antiword.exe",
            _TOOLS_ROOT / "antiword" / "bin" / "antiword.exe",
            _TOOLS_ROOT / "antiword" / "antiword",
            _TOOLS_ROOT / "antiword" / "bin" / "antiword",
        )
        or shutil.which("antiword")
        or ""
    )


def _resolve_soffice() -> str:
    return (
        _resolve_env_tool("RAG_SOFFICE_CMD", "RAG_LIBREOFFICE_CMD", "SOFFICE")
        or _first_existing_tool(
            _TOOLS_ROOT / "libreoffice" / "program" / "soffice.exe",
            _TOOLS_ROOT / "LibreOffice" / "program" / "soffice.exe",
            _TOOLS_ROOT / "LibreOfficePortable" / "App" / "libreoffice" / "program" / "soffice.exe",
            _TOOLS_ROOT / "LibreOfficePortable" / "App" / "LibreOffice" / "program" / "soffice.exe",
            _TOOLS_ROOT / "libreoffice" / "program" / "soffice",
            _TOOLS_ROOT / "LibreOffice" / "program" / "soffice",
        )
        or shutil.which("soffice")
        or shutil.which("libreoffice")
        or ""
    )


_DOC_TEXT_RUN_RE = re.compile(
    r"[A-Za-zА-Яа-яЁё0-9][A-Za-zА-Яа-яЁё0-9\s,.;:!?()\[\]{}№\"'«»%+\-_/\\]{5,}"
)


def _clean_doc_text_run(text: str) -> str:
    cleaned = re.sub(r"\s+", " ", text.replace("\x00", " ")).strip()
    return cleaned.strip(" ,.;:-_")


def _looks_like_doc_text(text: str) -> bool:
    if len(text) < 8:
        return False
    letters = sum(1 for char in text if char.isalpha())
    allowed = sum(1 for char in text if char.isalnum() or char.isspace() or char in ".,;:!?()[]{}№\"'«»%+-_/\\")
    return letters >= 3 and allowed / max(1, len(text)) >= 0.85


def _extract_doc_binary_fallback(filepath: Path, *, max_chars: int = 0) -> str:
    """Last-resort DOC text extraction without external binaries.

    Legacy .doc is an OLE binary container. A full parser is better, but many
    files still keep human text as UTF-16LE or Windows-1251 runs. This fallback
    extracts only conservative printable runs so indexing is not completely
    blind when bundled converters are absent.
    """
    try:
        data = filepath.read_bytes()
    except Exception as exc:
        logger.warning("DOC fallback: ошибка чтения %s: %s", filepath, exc)
        return ""

    candidates: list[str] = []
    seen: set[str] = set()
    for encoding in ("utf-16le", "cp1251"):
        try:
            decoded = data.decode(encoding, errors="ignore")
        except Exception:
            continue
        for match in _DOC_TEXT_RUN_RE.finditer(decoded):
            run = _clean_doc_text_run(match.group(0))
            if not _looks_like_doc_text(run):
                continue
            key = run.casefold()
            if key in seen:
                continue
            seen.add(key)
            candidates.append(run)
            if max_chars and sum(len(item) + 1 for item in candidates) >= max_chars:
                break
        if max_chars and sum(len(item) + 1 for item in candidates) >= max_chars:
            break

    text = "\n".join(candidates).strip()
    if max_chars:
        text = text[:max_chars]
    if text:
        logger.warning("DOC %s прочитан бинарным fallback без antiword/LibreOffice", filepath)
    return text


def extract_doc(filepath: Path, *, max_chars: int = 0) -> str:
    """Best-effort extraction for legacy binary DOC via bundled tools or fallback."""
    antiword = _resolve_antiword()
    if antiword:
        try:
            proc = subprocess.run(
                [antiword, str(filepath)],
                capture_output=True,
                text=True,
                timeout=120,
                encoding="utf-8",
                errors="replace",
                **_windows_hidden_popen_kwargs(),
            )
            if proc.returncode == 0 and proc.stdout.strip():
                return proc.stdout[:max_chars] if max_chars else proc.stdout
            logger.warning("antiword не извлёк DOC %s: %s", filepath, (proc.stderr or "").strip())
        except Exception as exc:
            logger.warning("antiword: ошибка чтения DOC %s: %s", filepath, exc)

    soffice = _resolve_soffice()
    if soffice:
        try:
            with tempfile.TemporaryDirectory(prefix="rag_doc_") as tmp:
                proc = subprocess.run(
                    [
                        soffice,
                        "--headless",
                        "--convert-to",
                        "txt:Text",
                        "--outdir",
                        tmp,
                        str(filepath),
                    ],
                    capture_output=True,
                    text=True,
                    timeout=180,
                    encoding="utf-8",
                    errors="replace",
                    **_windows_hidden_popen_kwargs(),
                )
                out_path = Path(tmp) / (filepath.stem + ".txt")
                if proc.returncode == 0 and out_path.exists():
                    text = _read_text_file(out_path, max_chars=max_chars)
                    if text.strip():
                        return text
                logger.warning("LibreOffice не извлёк DOC %s: %s", filepath, (proc.stderr or "").strip())
        except Exception as exc:
            logger.warning("LibreOffice: ошибка чтения DOC %s: %s", filepath, exc)

    fallback = _extract_doc_binary_fallback(filepath, max_chars=max_chars)
    if fallback:
        return fallback

    logger.warning(
        "DOC %s не прочитан: добавьте antiword или LibreOffice в tools/ либо задайте RAG_ANTIWORD_CMD/RAG_SOFFICE_CMD",
        filepath,
    )
    return ""


def extract_xlsx(filepath: Path, *, max_chars: int = 0) -> str:
    """Extract text from XLSX with optional early stop by accumulated chars."""
    wb: Any | None = None
    try:
        wb = _load_xlsx_workbook(filepath, read_only=True, data_only=True)
        parts: list[str] = []
        total_chars = 0
        done = False
        for ws in wb.worksheets:
            if done:
                break
            sheet_name = ws.title
            parts.append(f"Лист: {sheet_name}")
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join(str(c) if c is not None else "" for c in row)
                if row_text.strip():
                    parts.append(row_text)
                    total_chars += len(row_text)
                    if max_chars and total_chars >= max_chars:
                        done = True
                        break
        return "\n".join(parts)
    except KeyError as exc:
        if "xl/sharedStrings.xml" in str(exc):
            try:
                text = _extract_xlsx_zip_fallback(filepath, max_chars=max_chars)
                logger.warning("XLSX %s прочитан через fallback без sharedStrings.xml", filepath)
                return text
            except Exception as fallback_exc:
                logger.warning("Ошибка fallback-чтения XLSX %s: %s", filepath, fallback_exc)
                return ""
        logger.warning("Ошибка чтения XLSX %s: %s", filepath, exc)
        return ""
    except Exception as exc:
        logger.warning("Ошибка чтения XLSX %s: %s", filepath, exc)
        return ""
    finally:
        if wb is not None:
            try:
                wb.close()
            except Exception:
                pass


def extract_xlsx_document(filepath: Path, *, max_chars: int = 0) -> ExtractedDocument:
    """Extract XLSX rows as structured sheet/row blocks."""
    wb: Any | None = None
    try:
        wb = _load_xlsx_workbook(filepath, read_only=True, data_only=True)
        blocks: list[TextBlock] = []
        total_chars = 0
        done = False
        for ws in wb.worksheets:
            if done:
                break
            sheet_name = str(ws.title)
            for row_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
                row_text = " | ".join(str(c) if c is not None else "" for c in row)
                if not row_text.strip():
                    continue
                blocks.append(TextBlock(text=row_text, sheet=sheet_name, row_start=row_idx, row_end=row_idx))
                total_chars += len(row_text)
                if max_chars and total_chars >= max_chars:
                    done = True
                    break
        return ExtractedDocument(blocks=tuple(blocks))
    except KeyError as exc:
        if "xl/sharedStrings.xml" in str(exc):
            try:
                text = _extract_xlsx_zip_fallback(filepath, max_chars=max_chars)
                logger.warning("XLSX %s прочитан через fallback без sharedStrings.xml", filepath)
                return document_from_legacy_text(text)
            except Exception as fallback_exc:
                logger.warning("Ошибка fallback-чтения XLSX %s: %s", filepath, fallback_exc)
                return ExtractedDocument(blocks=())
        logger.warning("Ошибка чтения XLSX %s: %s", filepath, exc)
        return ExtractedDocument(blocks=())
    except Exception as exc:
        logger.warning("Ошибка чтения XLSX %s: %s", filepath, exc)
        return ExtractedDocument(blocks=())
    finally:
        if wb is not None:
            try:
                wb.close()
            except Exception:
                pass


def extract_xls(filepath: Path, *, max_chars: int = 0) -> str:
    """Extract text from legacy XLS files via xlrd."""
    try:
        import xlrd  # type: ignore
    except ImportError:
        logger.warning("xlrd не установлен. Установите: pip install xlrd")
        return ""
    try:
        wb = xlrd.open_workbook(str(filepath))
        parts: list[str] = []
        total_chars = 0
        done = False
        for sheet in wb.sheets():
            if done:
                break
            parts.append(f"Лист: {sheet.name}")
            for row_idx in range(sheet.nrows):
                row = sheet.row_values(row_idx)
                row_text = " | ".join(str(v) if v not in ("", None) else "" for v in row)
                if row_text.strip():
                    parts.append(row_text)
                    total_chars += len(row_text)
                    if max_chars and total_chars >= max_chars:
                        done = True
                        break
        return "\n".join(parts)
    except Exception as exc:
        logger.warning("Ошибка чтения XLS %s: %s", filepath, exc)
        return ""


def extract_xls_document(filepath: Path, *, max_chars: int = 0) -> ExtractedDocument:
    """Extract XLS rows as structured sheet/row blocks."""
    try:
        import xlrd  # type: ignore
    except ImportError:
        logger.warning("xlrd не установлен. Установите: pip install xlrd")
        return ExtractedDocument(blocks=())
    try:
        wb = xlrd.open_workbook(str(filepath))
        blocks: list[TextBlock] = []
        total_chars = 0
        done = False
        for sheet in wb.sheets():
            if done:
                break
            for row_idx in range(sheet.nrows):
                row = sheet.row_values(row_idx)
                row_text = " | ".join(str(v) if v not in ("", None) else "" for v in row)
                if not row_text.strip():
                    continue
                row_number = row_idx + 1
                blocks.append(TextBlock(text=row_text, sheet=str(sheet.name), row_start=row_number, row_end=row_number))
                total_chars += len(row_text)
                if max_chars and total_chars >= max_chars:
                    done = True
                    break
        return ExtractedDocument(blocks=tuple(blocks))
    except Exception as exc:
        logger.warning("Ошибка чтения XLS %s: %s", filepath, exc)
        return ExtractedDocument(blocks=())


def extract_spreadsheet(filepath: Path, *, max_chars: int = 0) -> str:
    """Route spreadsheet extraction by extension."""
    ext = filepath.suffix.lower()
    if ext == ".xls":
        logger.debug("Формат XLS — использую xlrd: %s", filepath.name)
        return extract_xls(filepath, max_chars=max_chars)
    if ext == ".xlsx":
        logger.debug("Формат XLSX — использую openpyxl: %s", filepath.name)
        return extract_xlsx(filepath, max_chars=max_chars)
    logger.warning("Неизвестное табличное расширение: %s", ext)
    return ""


def extract_spreadsheet_document(filepath: Path, *, max_chars: int = 0) -> ExtractedDocument:
    """Route structured spreadsheet extraction by extension."""
    ext = filepath.suffix.lower()
    if ext == ".xls":
        logger.debug("Формат XLS — использую xlrd: %s", filepath.name)
        return extract_xls_document(filepath, max_chars=max_chars)
    if ext == ".xlsx":
        logger.debug("Формат XLSX — использую openpyxl: %s", filepath.name)
        return extract_xlsx_document(filepath, max_chars=max_chars)
    logger.warning("Неизвестное табличное расширение: %s", ext)
    return ExtractedDocument(blocks=())


def _read_text_file(filepath: Path, *, max_chars: int = 0) -> str:
    raw = filepath.read_bytes()
    if max_chars:
        raw = raw[: max(max_chars * 4, 4096)]
    for encoding in ("utf-8-sig", "utf-8", "cp1251", "cp866", "latin-1"):
        try:
            text = raw.decode(encoding)
            return text[:max_chars] if max_chars else text
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")[:max_chars] if max_chars else raw.decode("utf-8", errors="replace")


def extract_text(filepath: Path, *, max_chars: int = 0) -> str:
    """Extract plain text files with common Windows/Russian encodings."""
    try:
        return _read_text_file(filepath, max_chars=max_chars)
    except Exception as exc:
        logger.warning("Ошибка чтения TXT %s: %s", filepath, exc)
        return ""


def extract_csv(filepath: Path, *, max_chars: int = 0) -> str:
    """Extract CSV as readable row text, preserving delimiter-separated values."""
    try:
        text = _read_text_file(filepath, max_chars=0)
        sample = text[:4096]
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        except csv.Error:
            dialect = csv.excel
        parts: list[str] = []
        total = 0
        reader = csv.reader(text.splitlines(), dialect)
        for row in reader:
            row_text = " | ".join(str(cell).strip() for cell in row if str(cell).strip())
            if not row_text:
                continue
            parts.append(row_text)
            total += len(row_text)
            if max_chars and total >= max_chars:
                break
        return "\n".join(parts)
    except Exception as exc:
        logger.warning("Ошибка чтения CSV %s: %s", filepath, exc)
        return ""


def extract_pdf(filepath: Path, *, skip_ocr: bool = False, ocr: Callable[[Path], str] | None = None) -> str:
    """Extract PDF text via pymupdf/pdfplumber, then OCR when text layer is empty."""
    try:
        import fitz  # pymupdf

        parts: list[str] = []
        with fitz.open(str(filepath)) as doc:
            for page_idx, page in enumerate(doc, start=1):
                text = page.get_text()
                if text and text.strip():
                    parts.append(f"Страница: {page_idx}\n{text}")
        full_text = "\n".join(parts).strip()
        if full_text:
            return full_text
        if skip_ocr:
            logger.debug("Нет текстового слоя, OCR пропущен (--no-ocr): %s", filepath.name)
            return ""
        logger.info("Нет текстового слоя в %s — запуск OCR…", filepath.name)
        return ocr(filepath) if ocr else ""
    except ImportError:
        logger.debug("pymupdf не установлен, использую pdfplumber")
    except Exception as exc:
        logger.warning("pymupdf: ошибка чтения %s: %s", filepath.name, exc)
        return ""

    try:
        import pdfplumber  # type: ignore
    except ImportError:
        logger.warning("Ни pymupdf, ни pdfplumber не установлены. pip install pymupdf")
        return ""
    try:
        parts = []
        with pdfplumber.open(filepath) as pdf:
            for page_idx, page in enumerate(pdf.pages, start=1):
                text = page.extract_text()
                if text and text.strip():
                    parts.append(f"Страница: {page_idx}\n{text}")
        full_text = "\n".join(parts).strip()
        if full_text:
            return full_text
        if skip_ocr:
            logger.debug("Нет текстового слоя, OCR пропущен (--no-ocr): %s", filepath.name)
            return ""
        logger.info("Нет текстового слоя в %s — запуск OCR…", filepath.name)
        return ocr(filepath) if ocr else ""
    except Exception as exc:
        logger.warning("pdfplumber: ошибка чтения %s: %s", filepath.name, exc)
        return ""


def extract_pdf_document(
    filepath: Path,
    *,
    skip_ocr: bool = False,
    ocr: Callable[[Path], str] | None = None,
) -> ExtractedDocument:
    """Extract PDF text as page-level blocks, falling back to OCR legacy text."""
    try:
        import fitz  # pymupdf

        blocks: list[TextBlock] = []
        with fitz.open(str(filepath)) as doc:
            for page_idx, page in enumerate(doc, start=1):
                text = page.get_text()
                if text and text.strip():
                    blocks.append(TextBlock(text=text.strip(), page=page_idx))
        if blocks:
            return ExtractedDocument(blocks=tuple(blocks))
        if skip_ocr:
            logger.debug("Нет текстового слоя, OCR пропущен (--no-ocr): %s", filepath.name)
            return ExtractedDocument(blocks=())
        logger.info("Нет текстового слоя в %s — запуск OCR…", filepath.name)
        return document_from_legacy_text(ocr(filepath) if ocr else "")
    except ImportError:
        logger.debug("pymupdf не установлен, использую pdfplumber")
    except Exception as exc:
        logger.warning("pymupdf: ошибка чтения %s: %s", filepath.name, exc)
        return ExtractedDocument(blocks=())

    try:
        import pdfplumber  # type: ignore
    except ImportError:
        logger.warning("Ни pymupdf, ни pdfplumber не установлены. pip install pymupdf")
        return ExtractedDocument(blocks=())
    try:
        blocks = []
        with pdfplumber.open(filepath) as pdf:
            for page_idx, page in enumerate(pdf.pages, start=1):
                text = page.extract_text()
                if text and text.strip():
                    blocks.append(TextBlock(text=text.strip(), page=page_idx))
        if blocks:
            return ExtractedDocument(blocks=tuple(blocks))
        if skip_ocr:
            logger.debug("Нет текстового слоя, OCR пропущен (--no-ocr): %s", filepath.name)
            return ExtractedDocument(blocks=())
        logger.info("Нет текстового слоя в %s — запуск OCR…", filepath.name)
        return document_from_legacy_text(ocr(filepath) if ocr else "")
    except Exception as exc:
        logger.warning("pdfplumber: ошибка чтения %s: %s", filepath.name, exc)
        return ExtractedDocument(blocks=())


def _windows_hidden_popen_kwargs() -> dict[str, Any]:
    """kwargs for subprocess calls without visible console window on Windows."""
    if os.name != "nt":
        return {}
    kwargs: dict[str, Any] = {"creationflags": int(getattr(subprocess, "CREATE_NO_WINDOW", 0) or 0)}
    if hasattr(subprocess, "STARTUPINFO"):
        si = subprocess.STARTUPINFO()
        si.dwFlags |= subprocess.STARTF_USESHOWWINDOW
        si.wShowWindow = getattr(subprocess, "SW_HIDE", 0)
        kwargs["startupinfo"] = si
    return kwargs


def _patch_pdf2image_popen_for_windows(pdf2image_module: Any) -> None:
    """Patch pdf2image internal Popen to suppress console flicker on Windows."""
    if os.name != "nt":
        return
    if getattr(pdf2image_module, "_rag_hidden_popen_patched", False):
        return
    original_popen = getattr(pdf2image_module, "Popen", None)
    if original_popen is None:
        return

    def _hidden_popen(*args: Any, **kwargs: Any) -> Any:
        hidden = _windows_hidden_popen_kwargs()
        for key, value in hidden.items():
            kwargs.setdefault(key, value)
        return original_popen(*args, **kwargs)

    pdf2image_module.Popen = _hidden_popen
    pdf2image_module._rag_hidden_popen_patched = True


def ocr_pdf(
    filepath: Path,
    *,
    tesseract_cmd: str = "",
    poppler_bin: str = "",
    use_rapid: bool = False,
) -> str:
    """OCR scanned PDF.

    use_rapid=True  → RapidOCR + DirectML/GPU (AMD/Intel/NVIDIA без CUDA)
    use_rapid=False → Tesseract (default, CPU)
    """
    if use_rapid:
        try:
            from rag_catalog.core.extractors.ocr_rapid import ocr_pdf_rapid  # noqa: PLC0415
            return ocr_pdf_rapid(filepath, poppler_bin=poppler_bin)
        except Exception as exc:
            logger.warning("RapidOCR PDF не удался, fallback на Tesseract: %s", exc)
    try:
        import pdf2image.pdf2image as pdf2image_impl  # type: ignore
        import pytesseract  # type: ignore
        from pdf2image import convert_from_path  # type: ignore
    except ImportError:
        logger.warning(
            "pytesseract/pdf2image не установлены. "
            "Установите: pip install pytesseract pdf2image"
        )
        return ""

    try:
        apply_tesseract_runtime(pytesseract, tesseract_cmd)
        _patch_pdf2image_popen_for_windows(pdf2image_impl)
        convert_kwargs: dict[str, Any] = {"dpi": 200}
        if str(poppler_bin or "").strip():
            convert_kwargs["poppler_path"] = str(poppler_bin).strip()
        pages = convert_from_path(str(filepath), **convert_kwargs)
        parts: list[str] = []
        for i, page_img in enumerate(pages):
            text = pytesseract.image_to_string(page_img, lang="rus+eng")
            chars = len(text.strip())
            if text.strip():
                parts.append(f"Страница: {i + 1}\n{text}")
            logger.info("OCR страница %d/%d — %d симв. — %s", i + 1, len(pages), chars, filepath.name)
        total_chars = sum(len(p) for p in parts)
        if parts:
            logger.info("OCR завершён: %s — %d стр., %d симв.", filepath.name, len(pages), total_chars)
        else:
            logger.warning("OCR не извлёк текст ни на одной странице: %s", filepath.name)
        return "\n".join(parts)
    except Exception as exc:
        logger.warning("OCR не удался для %s: %s", filepath, exc)
        return ""


def extract_image(
    filepath: Path,
    *,
    tesseract_cmd: str = "",
    max_pages: int = 50,
    use_rapid: bool = False,
) -> str:
    """Extract text from an image.

    use_rapid=True  → RapidOCR + DirectML/GPU
    use_rapid=False → Tesseract (default, CPU)
    """
    if use_rapid:
        try:
            from rag_catalog.core.extractors.ocr_rapid import ocr_image_rapid  # noqa: PLC0415
            return ocr_image_rapid(filepath)
        except Exception as exc:
            logger.warning("RapidOCR image не удался, fallback на Tesseract: %s", exc)
    try:
        import pytesseract  # type: ignore
    except ImportError:
        logger.debug(
            "pytesseract не установлен — OCR изображений недоступен. "
            "Установите: pip install pytesseract"
        )
        return ""
    try:
        from PIL import Image  # type: ignore
    except ImportError:
        logger.debug(
            "Pillow не установлен — OCR изображений недоступен. "
            "Установите: pip install Pillow"
        )
        return ""
    try:
        apply_tesseract_runtime(pytesseract, tesseract_cmd)
        parts: list[str] = []
        with Image.open(filepath) as img:
            n_frames: int = getattr(img, "n_frames", 1)
            if n_frames > max_pages:
                logger.warning(
                    "Изображение %s содержит %d кадров — обрабатываем только первые %d "
                    "(MAX_IMAGE_PAGES). Остальные пропущены.",
                    filepath.name,
                    n_frames,
                    max_pages,
                )
                n_frames = max_pages
            for frame_idx in range(n_frames):
                if n_frames > 1:
                    img.seek(frame_idx)
                    frame = img.copy()
                else:
                    frame = img
                if frame.mode not in ("RGB", "L", "RGBA"):
                    frame = frame.convert("RGB")
                page_text = pytesseract.image_to_string(frame, lang="rus+eng").strip()
                if page_text:
                    parts.append(page_text)
                logger.info(
                    "OCR %s стр.%d/%d: %d симв.",
                    filepath.name,
                    frame_idx + 1,
                    n_frames,
                    len(page_text),
                )
        result = "\n".join(parts).strip()
        if result:
            logger.info("OCR завершён: %s — %d симв., %d стр.", filepath.name, len(result), n_frames)
        else:
            logger.warning("OCR не извлёк текст: %s", filepath.name)
        return result
    except Exception as exc:
        logger.warning("OCR изображения не удался для %s: %s", filepath, exc)
        return ""


# ─── document metadata extraction ────────────────────────────────────────────

_OOXML_CORE_NS = {
    "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
    "dc": "http://purl.org/dc/elements/1.1/",
    "dcterms": "http://purl.org/dc/terms/",
}
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _read_ooxml_core(zf: ZipFile) -> dict[str, str]:
    """Parse docProps/core.xml from an Office Open XML ZIP archive."""
    try:
        with zf.open("docProps/core.xml") as fh:
            root = ElementTree.parse(fh).getroot()
    except (KeyError, ElementTree.ParseError):
        return {}

    def _t(tag: str) -> str:
        el = root.find(tag, _OOXML_CORE_NS)
        return (el.text or "").strip() if el is not None else ""

    return {
        "creator": _t("dc:creator"),
        "last_editor": _t("cp:lastModifiedBy"),
        "doc_created": _t("dcterms:created")[:10],
    }


def _docx_top_editor(zf: ZipFile) -> str:
    """Return the author with the most tracked-change events (w:ins / w:del) in document.xml."""
    try:
        with zf.open("word/document.xml") as fh:
            counts: dict[str, int] = {}
            for _ev, elem in ElementTree.iterparse(fh, events=("start",)):
                if elem.tag in (f"{{{_W_NS}}}ins", f"{{{_W_NS}}}del"):
                    author = elem.attrib.get(f"{{{_W_NS}}}author", "").strip()
                    if author:
                        counts[author] = counts.get(author, 0) + 1
                elem.clear()
        return max(counts, key=lambda k: counts[k]) if counts else ""
    except Exception:
        return ""


def _parse_pdf_date(raw: str) -> str:
    """Convert PDF date string D:YYYYMMDDHHmmss... → YYYY-MM-DD."""
    s = raw.lstrip("Dd:")
    if len(s) >= 8:
        try:
            return f"{s[0:4]}-{s[4:6]}-{s[6:8]}"
        except Exception:
            pass
    return ""


def extract_doc_meta(filepath: Path) -> dict[str, str]:
    """Extract author metadata from document properties.

    Returned keys (all str, empty when unavailable):
      doc_author      — original creator (dc:creator / PDF Author)
      doc_last_editor — last person to save (cp:lastModifiedBy)
      doc_top_editor  — most-frequent tracked-change author (DOCX only)
      doc_created     — document internal creation date YYYY-MM-DD
    """
    ext = filepath.suffix.lower()
    out: dict[str, str] = {
        "doc_author": "",
        "doc_last_editor": "",
        "doc_top_editor": "",
        "doc_created": "",
    }
    if ext in {".docx", ".xlsx", ".xlsm", ".pptx", ".ppsx"}:
        try:
            with ZipFile(filepath, "r") as zf:
                core = _read_ooxml_core(zf)
                out["doc_author"] = core.get("creator", "")
                out["doc_last_editor"] = core.get("last_editor", "")
                out["doc_created"] = core.get("doc_created", "")
                if ext == ".docx":
                    out["doc_top_editor"] = _docx_top_editor(zf)
        except Exception as exc:
            logger.debug("Метаданные %s: %s", filepath.name, exc)
    elif ext == ".pdf":
        try:
            import fitz  # noqa: PLC0415
            with fitz.open(str(filepath)) as doc:
                meta = doc.metadata or {}
                out["doc_author"] = (meta.get("author") or "").strip()
                out["doc_created"] = _parse_pdf_date(meta.get("creationDate") or "")
        except Exception as exc:
            logger.debug("Метаданные PDF %s: %s", filepath.name, exc)
    return out
